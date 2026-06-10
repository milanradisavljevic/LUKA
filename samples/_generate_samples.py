#!/usr/bin/env python3
"""Erzeugt synthetische Beispiel-Abgaben (.docx) zum gefahrlosen Testen von LUKA.

Alle Texte sind frei erfunden (keine echten Schülerdaten) und enthalten absichtlich
Fehler (Rechtschreibung/Grammatik/Zeichensetzung/Ausdruck), damit die Korrektur etwas
zu finden hat. Erneut ausführbar:  python3 _generate_samples.py
"""
from pathlib import Path
from docx import Document

HERE = Path(__file__).resolve().parent

# (Dateiname, Titel, Absätze) — bewusst mit Fehlern gespickt.
SAMPLES = [
    (
        "Beispiel_Kommentar_AB.docx",
        "Soziale Medien — Fluch oder Segen?",
        [
            "In der heutigen zeit sind soziale Medien aus unserem Alltag nicht mehr "
            "wegzudenken. Viele Jugendliche verbringen täglich mehrere stunden auf "
            "Plattformen wie Instagram oder TikTok, was durchaus problematisch sein kann.",
            "Einerseits bieten soziale Medien viele Vorteile, man kann mit Freunden in "
            "Kontakt bleiben und sich schnell informieren. Andererseits führt die ständige "
            "Nutzung oft dazu das man sich mit anderen vergleicht und unzufrieden wird.",
            "Meiner Meinung nach kommt es darauf an wie man die Medien benutzt. Wenn man "
            "bewusst damit umgeht sind sie ein nützliches Werkzeug, wenn nicht können sie "
            "einem schaden. Abschließend lässt sich sagen das ein gesundes Maß wichtig ist.",
        ],
    ),
    (
        "Beispiel_Eroerterung_CD.docx",
        "Sollte das Handy in der Schule verboten werden?",
        [
            "Das Thema Handyverbot an Schulen wird seit jahren heftig diskutiert. "
            "Befürworter argumentieren das Handys vom Unterricht ablenken und die "
            "Konzentration der Schüler stören.",
            "Es ist sicherlich richtig das ein Smartphone ablenken kann. Trotzdem finde "
            "ich ein komplettes verbot übertrieben, denn Handys können auch im Unterricht "
            "sinvoll eingesetzt werden, zum Beispiel zum recherchieren.",
            "Zusammenfassend bin ich der Meinung, dass man Schülern beibringen sollte einen "
            "verantwortungsvollen Umgang mit dem Handy zu lernen, anstatt es einfach zu "
            "verbieten. Ein verbot löst das eigentliche Problem nämlich nicht.",
        ],
    ),
]


def build(filename: str, title: str, paragraphs: list[str]) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    out = HERE / filename
    doc.save(out)
    print(f"geschrieben: {out.name}")


if __name__ == "__main__":
    for fn, title, paras in SAMPLES:
        build(fn, title, paras)
    print("Fertig. Diese Dateien dienen nur zum Testen (frei erfunden).")
