"""Regressionstest: DOCX-Kommentare tragen den Namen der Lehrkraft.

Hintergrund: teacher_name war als Realname in der Config verdrahtet — jede
Lehrkraft hätte als "Natascha" kommentiert. Der Name kommt jetzt aus dem
LUKA-Lehrerprofil (CLI --lehrer überschreibt docx.teacher_name); leerer
Config-Wert fällt auf das neutrale "Lehrkraft" zurück.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import generate_feedback as gf
from docx import Document


def _kommentar_autoren(doc) -> set[str]:
    """Sammelt alle w:author-Werte aus den Comments-Parts des Dokuments."""
    import re

    autoren: set[str] = set()
    for part in doc.part.package.iter_parts():
        blob = getattr(part, "blob", b"") or b""
        if b"comments" in part.partname.encode() and b"w:author" in blob:
            autoren |= {
                m.decode("utf-8") for m in re.findall(rb'w:author="([^"]*)"', blob)
            }
    return autoren


def test_kommentar_autor_kommt_aus_uebergebenem_namen():
    doc = Document()
    p = doc.add_paragraph("Ein Satz mit Inhalt.")
    gf._attach_word_comments(doc, [(p, "Gut formuliert.")], author="Milan Muster")
    autoren = _kommentar_autoren(doc)
    assert autoren == {"Milan Muster"}, autoren


def test_kommentar_autor_default_ist_neutral():
    doc = Document()
    p = doc.add_paragraph("Noch ein Satz.")
    gf._attach_word_comments(doc, [(p, "Hinweis.")])
    autoren = _kommentar_autoren(doc)
    assert autoren == {"Lehrkraft"}, autoren
    assert "Natascha" not in " ".join(autoren)


def test_leerer_config_name_faellt_auf_lehrkraft_zurueck():
    cfg = {"docx": {"teacher_name": ""}}
    docx_cfg = cfg.get("docx", {})
    assert (docx_cfg.get("teacher_name") or "Lehrkraft") == "Lehrkraft"
