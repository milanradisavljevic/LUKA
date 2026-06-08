#!/usr/bin/env python3
"""Seed-Skript für synthetische Testdaten.

Nur für Testdaten. Legt synthetische Schüler mit 3–4 Arbeiten an.
Mehrfach ausführbar (idempotent).
"""

from __future__ import annotations

import hashlib
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

import natascha_core as nc
import natascha_db as db

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore[misc,assignment]


KLASSE = "TEST-7a"

SCHUELER = [
    {
        "vorname": "Testschueler",
        "nachname": "Mona",
        # Aufsteigende Tendenz: Kriterien steigen, Noten sinken (Verbesserung)
        "arbeiten": [
            {
                "aufgabe": "SA1",
                "note_app": 5.0,
                "kriterien": {"inhalt": 2, "textstruktur": 2, "ausdruck": 2, "sprachrichtigkeit": 1},
                "fehler": [
                    ("das Haus", "das Haus,", "Z", "Komma fehlt"),
                    ("im Garten", "im Garten,", "Z", "Komma fehlt"),
                    ("der Hund", "der Hund,", "Z", "Komma fehlt"),
                    ("wahr", "war", "R", "Rechtschreibung"),
                ],
                "note_lehrer": None,
            },
            {
                "aufgabe": "SA2",
                "note_app": 4.0,
                "kriterien": {"inhalt": 3, "textstruktur": 2, "ausdruck": 3, "sprachrichtigkeit": 2},
                "fehler": [
                    ("das Haus", "das Haus,", "Z", "Komma fehlt"),
                    ("im Garten", "im Garten,", "Z", "Komma fehlt"),
                    ("wahr", "war", "R", "Rechtschreibung"),
                    ("geht", "ging", "G", "Zeitform"),
                ],
                "note_lehrer": 3.0,
            },
            {
                "aufgabe": "SA3",
                "note_app": 3.0,
                "kriterien": {"inhalt": 3, "textstruktur": 3, "ausdruck": 3, "sprachrichtigkeit": 3},
                "fehler": [
                    ("das Haus", "das Haus,", "Z", "Komma fehlt"),
                    ("wahr", "war", "R", "Rechtschreibung"),
                ],
                "note_lehrer": None,
            },
            {
                "aufgabe": "SA4",
                "note_app": 2.0,
                "kriterien": {"inhalt": 4, "textstruktur": 3, "ausdruck": 4, "sprachrichtigkeit": 3},
                "fehler": [
                    ("das Haus", "das Haus,", "Z", "Komma fehlt"),
                ],
                "note_lehrer": 1.0,
            },
        ],
    },
    {
        "vorname": "Testschueler",
        "nachname": "Max",
        # Stabil mittel, andere Fehlerschwerpunkte
        "arbeiten": [
            {
                "aufgabe": "SA1",
                "note_app": 3.0,
                "kriterien": {"inhalt": 3, "textstruktur": 3, "ausdruck": 3, "sprachrichtigkeit": 3},
                "fehler": [
                    ("er geht", "er ging", "G", "Zeitform"),
                    ("die Frau", "die Frau,", "Z", "Komma fehlt"),
                ],
                "note_lehrer": None,
            },
            {
                "aufgabe": "SA2",
                "note_app": 3.0,
                "kriterien": {"inhalt": 3, "textstruktur": 3, "ausdruck": 3, "sprachrichtigkeit": 3},
                "fehler": [
                    ("er geht", "er ging", "G", "Zeitform"),
                    ("die Frau", "die Frau,", "Z", "Komma fehlt"),
                    ("gut", "gut,", "A", "Ausdruck"),
                ],
                "note_lehrer": None,
            },
            {
                "aufgabe": "SA3",
                "note_app": 3.0,
                "kriterien": {"inhalt": 3, "textstruktur": 3, "ausdruck": 3, "sprachrichtigkeit": 3},
                "fehler": [
                    ("er geht", "er ging", "G", "Zeitform"),
                ],
                "note_lehrer": None,
            },
            {
                "aufgabe": "SA4",
                "note_app": 3.0,
                "kriterien": {"inhalt": 3, "textstruktur": 3, "ausdruck": 3, "sprachrichtigkeit": 3},
                "fehler": [
                    ("er geht", "er ging", "G", "Zeitform"),
                ],
                "note_lehrer": None,
            },
        ],
    },
    {
        "vorname": "Testschueler",
        "nachname": "Mia",
        # Absteigende Tendenz (Verschlechterung)
        "arbeiten": [
            {
                "aufgabe": "SA1",
                "note_app": 2.0,
                "kriterien": {"inhalt": 4, "textstruktur": 4, "ausdruck": 4, "sprachrichtigkeit": 4},
                "fehler": [
                    ("schön", "schön,", "A", "Ausdruck"),
                ],
                "note_lehrer": None,
            },
            {
                "aufgabe": "SA2",
                "note_app": 3.0,
                "kriterien": {"inhalt": 3, "textstruktur": 3, "ausdruck": 3, "sprachrichtigkeit": 3},
                "fehler": [
                    ("schön", "schön,", "A", "Ausdruck"),
                    ("der Baum", "der Baum,", "Z", "Komma fehlt"),
                ],
                "note_lehrer": None,
            },
            {
                "aufgabe": "SA3",
                "note_app": 4.0,
                "kriterien": {"inhalt": 2, "textstruktur": 2, "ausdruck": 2, "sprachrichtigkeit": 2},
                "fehler": [
                    ("schön", "schön,", "A", "Ausdruck"),
                    ("der Baum", "der Baum,", "Z", "Komma fehlt"),
                    ("haus", "Haus", "R", "Rechtschreibung"),
                ],
                "note_lehrer": None,
            },
            {
                "aufgabe": "SA4",
                "note_app": 5.0,
                "kriterien": {"inhalt": 2, "textstruktur": 1, "ausdruck": 2, "sprachrichtigkeit": 1},
                "fehler": [
                    ("schön", "schön,", "A", "Ausdruck"),
                    ("der Baum", "der Baum,", "Z", "Komma fehlt"),
                    ("haus", "Haus", "R", "Rechtschreibung"),
                    ("geht", "ging", "G", "Zeitform"),
                ],
                "note_lehrer": None,
            },
        ],
    },
]


def _hash(vorname: str, nachname: str, aufgabe: str) -> str:
    """Deterministischer Hash pro Schüler+Aufgabe für idempotentes Insert.

    Für den ursprünglichen Testschüler (Mona) wird der alte Hash beibehalten,
    damit bestehende Testdaten nicht dupliziert werden.
    """
    if vorname == "Testschueler" and nachname == "Mona":
        return hashlib.sha256(f"{KLASSE}-{aufgabe}".encode()).hexdigest()
    return hashlib.sha256(f"{KLASSE}-{vorname}-{nachname}-{aufgabe}".encode()).hexdigest()


def main() -> int:
    config = nc.load_config()
    db_path = db.get_db_path(config)
    db.init_db(db_path)

    for schueler_data in SCHUELER:
        vorname = schueler_data["vorname"]
        nachname = schueler_data["nachname"]

        # Schüler idempotent anlegen oder wiederverwenden
        schueler = db.get_schueler_by_name(db_path, KLASSE, vorname, nachname)
        if schueler is None:
            schueler_id = db.insert_schueler(db_path, KLASSE, vorname, nachname)
            print(f"Schüler angelegt: {vorname} {nachname} (ID: {schueler_id})")
        else:
            schueler_id = schueler["id"]
            print(f"Schüler existiert bereits: {vorname} {nachname} (ID: {schueler_id})")

        for arbeit in schueler_data["arbeiten"]:
            aufgabe = arbeit["aufgabe"]
            datei_hash = _hash(vorname, nachname, aufgabe)

            # Idempotenz: Hash prüfen
            existing = db.get_abgabe_by_hash(db_path, datei_hash)
            if existing:
                print(f"  {aufgabe}: bereits vorhanden (ID: {existing['id']}) — überspringe")
                continue

            aid = db.insert_abgabe(
                db_path,
                schueler_id=schueler_id,
                klasse=KLASSE,
                aufgabe=aufgabe,
                dateiname=f"{aufgabe}.docx",
                datei_hash=datei_hash,
                note=arbeit["note_app"],
                schulstufe="Oberstufe",
            )
            print(f"  {aufgabe}: Abgabe angelegt (ID: {aid})")

            # Kriterien
            for name, stufe in arbeit["kriterien"].items():
                db.insert_kriterium(db_path, aid, name, float(stufe), 0.25)

            # Fehler
            for zitat, korr, typ, erkl in arbeit["fehler"]:
                db.insert_fehler(db_path, aid, zitat, korr, typ, erkl)

            # Lehrer-Feedback (optional)
            if arbeit["note_lehrer"] is not None:
                db.upsert_lehrer_feedback(
                    db_path,
                    abgabe_id=aid,
                    klasse=KLASSE,
                    aufgabe=aufgabe,
                    note_final=arbeit["note_lehrer"],
                    note_app_snapshot=arbeit["note_app"],
                    lehrer_kommentar=f"Test-Kommentar {aufgabe}",
                    schueler_id=schueler_id,
                )
                print(f"    → Lehrer-Note: {arbeit['note_lehrer']} (App: {arbeit['note_app']})")

        # Dummy-DOCX erstellen
        input_dir = Path(__file__).resolve().parent / "input" / KLASSE
        input_dir.mkdir(parents=True, exist_ok=True)
        docx_path = input_dir / f"{vorname}_{nachname}.docx"
        if Document is not None:
            if not docx_path.exists():
                doc = Document()
                doc.add_paragraph(f"{vorname} {nachname}")
                doc.add_paragraph(
                    "Dies ist ein automatisch erzeugter Dummy-Text für Testzwecke. "
                    "Die Schülerin zeigt eine stabile Leistung mit Potenzial zur Verbesserung."
                )
                doc.save(str(docx_path))
                print(f"  Dummy-DOCX erstellt: {docx_path}")
            else:
                print(f"  Dummy-DOCX existiert bereits: {docx_path}")

    print(f"\nFertig. {len(SCHUELER)} Schüler in Klasse '{KLASSE}'.")
    print(f"In der App: Klasse '{KLASSE}' auswählen → Dateien erscheinen in der Liste")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
