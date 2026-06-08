#!/usr/bin/env python3
"""
Korrektur-Feedback Generator

Erzeugt Feedback-DOCX aus fertigen Analyse-JSONs in /output/feedback_data.
"""

from __future__ import annotations

import argparse
import json
import logging
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

try:
    from lxml import etree as _etree
    _LXML_AVAILABLE = True
except ImportError:
    _LXML_AVAILABLE = False

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_COMMENTS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
_COMMENTS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)

# Farbpalette
C_PRIMARY = RGBColor(0x1F, 0x49, 0x7D)  # Dunkelblau  – Header, Standard
C_HEADER = RGBColor(0x2E, 0x4D, 0x8A)  # Dunkelblau  – Abschnittstitel
C_DIVIDER = RGBColor(0x99, 0x99, 0x99)  # Grau        – Trennlinien
C_STRENGTH = RGBColor(0x00, 0x70, 0x00)  # Dunkelgrün  – Stärken
C_WEAKNESS = RGBColor(0xC0, 0x00, 0x00)  # Dunkelrot   – Schwächen
C_SUGGESTION = RGBColor(0x00, 0x46, 0x7F)  # Blau        – Verbesserungsvorschläge
C_GRADE_1 = RGBColor(0x70, 0xC0, 0x70)  # Hellgrün   – Note 1
C_GRADE_2 = RGBColor(0x00, 0x80, 0x00)  # Dunkelgrün – Note 2
C_GRADE_3 = RGBColor(0xC0, 0xA0, 0x00)  # Gelb       – Note 3
C_GRADE_4 = RGBColor(0xE0, 0x60, 0x00)  # Orange     – Note 4
C_GRADE_5 = RGBColor(0xC0, 0x00, 0x00)  # Rot        – Note 5

# Rückwärtskompatible Aliasse (werden in Tests referenziert)
C_GRADE_GOOD = C_GRADE_2
C_GRADE_OK = C_GRADE_3
C_GRADE_FAIL = C_GRADE_5

_GRADE_COLORS = {1: C_GRADE_1, 2: C_GRADE_2, 3: C_GRADE_3, 4: C_GRADE_4, 5: C_GRADE_5}

_FEHLER_TYP_COLORS: dict[str, RGBColor] = {
    "R": RGBColor(0x1F, 0x49, 0x7D),  # Blau        – Rechtschreibung
    "G": RGBColor(0xC0, 0x00, 0x00),  # Rot         – Grammatik
    "Z": RGBColor(0xE0, 0x60, 0x00),  # Orange      – Zeichensetzung
    "A": RGBColor(0x80, 0x80, 0x80),  # Grau        – Ausdruck
}

SRDP_WORTLAUT: dict[str, dict[int, str]] = {
    "schreibhandlung": {
        0: "nicht erfüllt",
        1: "Schreibhandlung(en) im Sinne der geforderten Textsorte überwiegend realisiert",
        2: "Schreibhandlung(en) im Sinne der geforderten Textsorte weitgehend realisiert",
        3: "Schreibhandlung(en) im Sinne der geforderten Textsorte durchgehend realisiert",
        4: "Schreibhandlung(en) im Sinne der geforderten Textsorte umfassend realisiert",
    },
    "arbeitsauftraege": {
        0: "nicht erfüllt",
        1: "Arbeitsaufträge überwiegend erfüllt",
        2: "Arbeitsaufträge weitgehend erfüllt",
        3: "alle Arbeitsaufträge erfüllt",
        4: "alle Arbeitsaufträge umfassend erfüllt",
    },
    "textbeilage": {
        0: "nicht erfüllt",
        1: "Textbeilage(n) im Sinne der Arbeitsaufträge überwiegend erfasst",
        2: "Textbeilage(n) im Sinne der Arbeitsaufträge weitgehend erfasst",
        3: "Textbeilage(n) im Sinne der Arbeitsaufträge vollständig erfasst",
        4: "Textbeilage(n) im Sinne der Arbeitsaufträge vollständig erfasst",
    },
    "sachlich": {
        0: "nicht erfüllt",
        1: "sachlich überwiegend richtig",
        2: "sachlich weitgehend richtig",
        3: "sachlich richtig",
        4: "sachlich durchgehend richtig",
    },
    "qualitaet": {
        0: "nicht erfüllt",
        1: "oberflächlich / wenig treffsicher / reproduzierend",
        2: "ansatzweise komplex / weitgehend treffsicher / Ansätze zur Eigenständigkeit",
        3: "komplex / treffsicher / merklich eigenständig",
        4: "in hohem Maße komplex / treffsicher / eigenständig; ggf. ideenreich",
    },
    "kohaerenz": {
        0: "nicht erfüllt",
        1: "Text gedanklich und formal überwiegend der Textsorte angemessen strukturiert",
        2: "Text gedanklich und formal weitgehend der Textsorte angemessen strukturiert",
        3: "Text gedanklich und formal durchgehend der Textsorte angemessen und klar strukturiert",
        4: "Text gedanklich und formal durchgehend der Textsorte angemessen, klar, zielgerichtet und ggf. eigenständig strukturiert",
    },
    "bezugnahme": {
        0: "nicht erfüllt",
        1: "Bezugnahme auf die Textbeilage(n) im Sinne der geforderten Textsorte überwiegend erkennbar",
        2: "Bezugnahme auf die Textbeilage(n) im Sinne der geforderten Textsorte realisiert",
        3: "gelungene Verknüpfung mit der/den Textbeilage(n) im Sinne der geforderten Textsorte",
        4: "besonders gelungene Verknüpfung mit der/den Textbeilage(n) im Sinne der geforderten Textsorte",
    },
    "kohaesion": {
        0: "nicht erfüllt",
        1: "Einsatz passender Kohäsionsmittel überwiegend erkennbar",
        2: "Einsatz passender Kohäsionsmittel weitgehend erkennbar",
        3: "nahezu durchgehender Einsatz passender Kohäsionsmittel",
        4: "durchgehender Einsatz passender Kohäsionsmittel",
    },
    "situationsadaequat": {
        0: "nicht erfüllt",
        1: "überwiegend schreibhandlungs- und situationsadäquate Sprachverwendung",
        2: "weitgehend schreibhandlungs- und situationsadäquate Sprachverwendung",
        3: "nahezu durchgehend schreibhandlungs- und situationsadäquate Sprachverwendung",
        4: "durchgehend schreibhandlungs- und situationsadäquate Sprachverwendung",
    },
    "wortwahl": {
        0: "nicht erfüllt",
        1: "überwiegend angemessene und semantisch korrekte Ausdrucksweise sowie geringe Varianz in der Wortwahl",
        2: "weitgehend angemessene und semantisch korrekte Ausdrucksweise sowie variantenreiche Wortwahl",
        3: "durchgehend angemessene und semantisch korrekte Ausdrucksweise sowie präzise und variantenreiche Wortwahl",
        4: "durchgehend angemessene und semantisch korrekte Ausdrucksweise sowie besonders präzise, differenzierte und variantenreiche Wortwahl",
    },
    "satzstrukturen": {
        0: "nicht erfüllt",
        1: "überwiegend gut verständliche bzw. nur wenig variierende Satzstrukturen",
        2: "weitgehend gut verständliche und variantenreiche Satzstrukturen",
        3: "durchgehend variantenreiche und komplexe bzw. der Textsorte angemessene Satzstrukturen",
        4: "besonders variantenreiche und komplexe bzw. der Textsorte angemessene Satzstrukturen",
    },
    "eigenstaendigkeit": {
        0: "nicht erfüllt",
        1: "viele an die Textbeilage(n) angelehnte oder wörtlich übernommene Formulierungen",
        2: "weitgehend eigenständige Formulierungen",
        3: "nahezu durchgehend eigenständige Formulierungen",
        4: "durchgehend eigenständige Formulierungen",
    },
    "orthografie": {
        0: "nicht erfüllt",
        1: "überwiegend richtige Anwendung der Regeln der Orthografie",
        2: "weitgehend richtige Anwendung der Regeln der Orthografie",
        3: "richtige Anwendung der Regeln der Orthografie; wenige Fehler",
        4: "orthografisch (nahezu) fehlerfrei",
    },
    "zeichensetzung": {
        0: "nicht erfüllt",
        1: "überwiegend richtige Anwendung der Regeln der Zeichensetzung",
        2: "weitgehend richtige Anwendung der Regeln der Zeichensetzung",
        3: "richtige Anwendung der Regeln der Zeichensetzung; wenige Fehler",
        4: "Zeichensetzung (nahezu) fehlerfrei",
    },
    "grammatik": {
        0: "nicht erfüllt",
        1: "überwiegend richtige Anwendung der Regeln der Grammatik",
        2: "weitgehend richtige Anwendung der Regeln der Grammatik",
        3: "richtige Anwendung der Regeln der Grammatik; wenige Fehler",
        4: "grammatikalisch (nahezu) fehlerfrei",
    },
}


def _note_color(note: int) -> RGBColor:
    """Notenfarbe: 1=hellgrün, 2=dunkelgrün, 3=gelb, 4=orange, 5=rot."""
    return _GRADE_COLORS.get(note, C_GRADE_5)


def _set_cell_bg(cell, hex_color: str) -> None:
    """Setzt die Hintergrundfarbe einer Tabellenzelle via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


GERMAN_ORDER = ["inhalt", "textstruktur", "ausdruck", "sprachrichtigkeit"]
ENGLISH_ORDER = [
    "task_achievement",
    "organisation_layout",
    "lexical_range_accuracy",
    "grammatical_range_accuracy",
]

ALIASES = {
    "stil_ausdruck": "ausdruck",
    "normative_sprachrichtigkeit": "sprachrichtigkeit",
    "erfuellung_aufgabenstellung": "task_achievement",
    "aufbau_layout": "organisation_layout",
    "wortschatz": "lexical_range_accuracy",
    "grammatik": "grammatical_range_accuracy",
}


@dataclass(slots=True)
class CriterionFeedback:
    """Bewertungsdaten fuer ein einzelnes Kriterium."""

    key: str
    stufe: str
    punkte: float
    staerken: list[str]
    schwaechen: list[str]
    vorschlaege: list[str]
    gewicht: float | None = None
    fehler_detail: list[str] | None = None
    fehlerschwerpunkte: list[str] | None = None
    rhetorische_figuren: list[str] | None = None


@dataclass(slots=True)
class WordLevelHinweis:
    zitat: str
    kommentar: str


@dataclass(slots=True)
class SprachFehler:
    zitat: str
    korrektur: str
    typ: str  # "R" | "G" | "Z" | "A"
    erklaerung: str = ""


@dataclass(slots=True)
class GradeRecommendation:
    """Zusammenfassung der Notenempfehlung."""

    durchschnitt: float
    note: int
    bezeichnung: str
    begruendung: str
    k1_note: int | None = None
    k3_note: int | None = None
    k1_schnitt: float | None = None
    k3_schnitt: float | None = None
    sonderregel: str | None = None


@dataclass(slots=True)
class FeedbackData:
    """Vollstaendige Daten fuer ein Feedback-Dokument."""

    datei: str
    schueler: str | None
    klasse: str | None
    textsorte: str
    fach: str
    schulstufe: str
    rubrik: str
    bewertung: list[CriterionFeedback]
    notenempfehlung: GradeRecommendation | None
    hinweise: list[WordLevelHinweis]
    fehler: list[SprachFehler] | None = None
    modell: str | None = None
    zusammenfassung: str | None = None
    staerken_global: list[str] | None = None
    verbesserungsbereiche: list[str] | None = None
    srdp_detail: dict[str, Any] | None = None


@dataclass(slots=True)
class ProjectPaths:
    """Projektpfade relativ zum Skriptstandort."""

    root: Path
    input_dir: Path
    output_dir: Path
    feedback_data_dir: Path
    fehlerlog: Path


def project_paths() -> ProjectPaths:
    """Ermittelt die Projektpfade relativ zum Skript."""

    root = Path(__file__).resolve().parent
    output_dir = root / "output"
    return ProjectPaths(
        root=root,
        input_dir=root / "input",
        output_dir=output_dir,
        feedback_data_dir=output_dir / "feedback_data",
        fehlerlog=output_dir / "fehlerlog.txt",
    )


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Fuegt eine Ueberschrift ein."""

    doc.add_heading(text, level=level)


def add_divider(doc: Document) -> None:
    """Fuegt einen dezenten Trenner ein."""

    paragraph = doc.add_paragraph("-" * 56)
    if paragraph.runs:
        paragraph.runs[0].font.color.rgb = C_DIVIDER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)


def add_section_header(doc: Document, text: str, color: RGBColor = C_HEADER) -> None:
    """Fuegt einen Abschnittstitel im bestaenden Stil ein."""

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(2)


def add_label(doc: Document, label: str, value: str) -> None:
    """Fuegt eine Label-Wert-Zeile ein."""

    paragraph = doc.add_paragraph()
    key_run = paragraph.add_run(f"{label}: ")
    key_run.bold = True
    key_run.font.size = Pt(10)
    value_run = paragraph.add_run(value)
    value_run.font.size = Pt(10)
    paragraph.paragraph_format.space_after = Pt(1)


def add_bullet(
    doc: Document, text: str, indent: int = 1, color: RGBColor | None = None
) -> None:
    """Fuegt einen Listenpunkt ein."""

    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    if color is not None:
        run.font.color.rgb = color
    paragraph.paragraph_format.left_indent = Inches(0.3 * indent)
    paragraph.paragraph_format.space_after = Pt(1)


def add_body(doc: Document, text: str) -> None:
    """Fuegt einen normalen Absatz ein."""

    paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        run.font.size = Pt(10)
    paragraph.paragraph_format.space_after = Pt(2)


def setup_page(doc: Document, config: dict[str, Any] | None = None) -> None:
    """Setzt A4-Seitenformat, Raender, Fusszeile und optionale Kopfzeile."""

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = fp.add_run(
        f"Erstellt mit NATASCHA — {date.today().strftime('%d.%m.%Y')}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = C_DIVIDER

    if config:
        docx_cfg = config.get("docx", {})
        teacher = docx_cfg.get("teacher_name", "")
        school = docx_cfg.get("school_name", "")
        if teacher or school:
            header = section.header
            header.is_linked_to_previous = False
            parts = [p for p in [teacher, school] if p]
            hp = header.paragraphs[0]
            header_run = hp.add_run(" | ".join(parts))
            header_run.font.size = Pt(8)
            header_run.font.color.rgb = C_DIVIDER


def add_lehrer_kommentar_block(
    doc: Document, data: FeedbackData, bewertungsmodus: str = "benotet"
) -> None:
    """Fügt den holistischen LEHRKRAFT-KOMMENTAR-Block ein.

    Enthält ZUSAMMENFASSUNG, STÄRKEN, VERBESSERUNGSBEREICHE und NOTENEMPFEHLUNG
    als kompakten Überblick vor der detaillierten Kriterien-Aufschlüsselung.
    Wird nur gerendert wenn mindestens ein Feld vorhanden ist.
    """
    if not any([data.zusammenfassung, data.staerken_global, data.verbesserungsbereiche]):
        return

    add_section_header(doc, "LEHRKRAFT-KOMMENTAR")

    if data.zusammenfassung:
        add_section_header(doc, "ZUSAMMENFASSUNG", color=C_PRIMARY)
        p = doc.add_paragraph(data.zusammenfassung)
        p.runs[0].font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    if data.staerken_global:
        add_section_header(doc, "STÄRKEN", color=C_STRENGTH)
        for item in data.staerken_global:
            add_bullet(doc, item, color=C_STRENGTH)

    if data.verbesserungsbereiche:
        add_section_header(doc, "VERBESSERUNGSBEREICHE", color=C_WEAKNESS)
        for item in data.verbesserungsbereiche:
            add_bullet(doc, item, color=C_WEAKNESS)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    if data.notenempfehlung and bewertungsmodus == "benotet":
        grade_color = _note_color(data.notenempfehlung.note)
        label_run = p.add_run(
            f"NOTENEMPFEHLUNG: {data.notenempfehlung.note} – {data.notenempfehlung.bezeichnung}"
        )
        label_run.bold = True
        label_run.font.size = Pt(11)
        label_run.font.color.rgb = grade_color

        if data.notenempfehlung.k1_note is not None:
            k1_schnitt_str = (
                f" [Stufe {data.notenempfehlung.k1_schnitt:.1f}]"
                if data.notenempfehlung.k1_schnitt is not None
                else ""
            )
            add_body(
                doc,
                f"  K1 (Inhalt + Textstruktur): Note {data.notenempfehlung.k1_note}{k1_schnitt_str}",
            )
        if data.notenempfehlung.k3_note is not None:
            k3_schnitt_str = (
                f" [Stufe {data.notenempfehlung.k3_schnitt:.1f}]"
                if data.notenempfehlung.k3_schnitt is not None
                else ""
            )
            add_body(
                doc,
                f"  K3/1 (Stil + Sprachnormen): Note {data.notenempfehlung.k3_note}{k3_schnitt_str}",
            )
        if data.notenempfehlung.sonderregel:
            add_body(
                doc,
                f"  Hinweis: Sonderregel – {data.notenempfehlung.sonderregel}",
            )
    else:
        label_run = p.add_run("Hausaufgabe – keine Benotung")
        label_run.bold = True
        label_run.font.size = Pt(11)
        label_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    add_divider(doc)


def add_korrektur_header(
    doc: Document,
    data: FeedbackData,
    config: dict[str, Any] | None = None,
    word_count: int = 0,
    aufgabe_label: str = "",
) -> None:
    """Fügt den Titel-Header für das Korrigierte-Schülerarbeit-Format ein."""
    docx_cfg = (config or {}).get("docx", {})

    logo_path = docx_cfg.get("logo_path", "")
    if logo_path:
        lp = Path(logo_path)
        if not lp.is_absolute():
            lp = project_paths().root / lp
        if lp.exists():
            doc.add_picture(str(lp), width=Cm(4))
            doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("KORRIGIERTE SCHÜLERARBEIT")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = C_PRIMARY

    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.autofit = True
    rows_data: list[tuple[str, str]] = []
    if data.schueler:
        val = data.schueler if not data.klasse else f"{data.schueler}, {data.klasse}"
        rows_data.append(("Schüler/in", val))
    rows_data.append(("Textsorte", data.textsorte))
    rows_data.append(("Fach", data.fach))
    if word_count:
        rows_data.append(("Wortanzahl", str(word_count)))
    rows_data.append(("Korrekturdatum", date.today().strftime("%d.%m.%Y")))
    if aufgabe_label:
        rows_data.append(("Aufgabe", aufgabe_label))
    teacher = docx_cfg.get("teacher_name", "")
    school = docx_cfg.get("school_name", "")
    if teacher:
        rows_data.append(("Lehrkraft", teacher))
    if school:
        rows_data.append(("Schule", school))
    for label, value in rows_data:
        row = meta_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    add_divider(doc)


def add_legende(doc: Document) -> None:
    """Fügt die Korrektur-Legende ein."""
    p = doc.add_paragraph()
    run = p.add_run("LEGENDE:")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = C_DIVIDER
    p.paragraph_format.space_after = Pt(1)

    items = [
        ("durchgestrichen", "Fehler im Original", C_WEAKNESS),
        ("grün + fett", "Korrektur/Verbesserung", C_STRENGTH),
        ("Balloon-Kommentare rechts", "Hinweise der Lehrkraft", C_PRIMARY),
    ]
    for mark, meaning, color in items:
        leg_p = doc.add_paragraph()
        leg_p.paragraph_format.space_after = Pt(1)
        leg_p.paragraph_format.left_indent = Inches(0.2)
        mark_run = leg_p.add_run(f"• {mark}")
        mark_run.font.size = Pt(9)
        mark_run.font.color.rgb = color
        sep_run = leg_p.add_run(f" = {meaning}")
        sep_run.font.size = Pt(9)
        sep_run.font.color.rgb = C_DIVIDER

    # Fehler-Kommentar-Legende (Track-Changes-Stil)
    tc_p = doc.add_paragraph()
    tc_p.paragraph_format.space_after = Pt(1)
    tc_p.paragraph_format.left_indent = Inches(0.2)
    tc_hdr = tc_p.add_run("• Fehler-Kommentare: ")
    tc_hdr.font.size = Pt(9)
    tc_hdr.font.color.rgb = C_DIVIDER
    for typ, label, color in [
        ("R", "Rechtschreibung", _FEHLER_TYP_COLORS["R"]),
        ("G", "Grammatik", _FEHLER_TYP_COLORS["G"]),
        ("Z", "Zeichensetzung", _FEHLER_TYP_COLORS["Z"]),
        ("A", "Ausdruck", _FEHLER_TYP_COLORS["A"]),
    ]:
        r = tc_p.add_run(f"[{typ}]={label}  ")
        r.font.size = Pt(9)
        r.font.color.rgb = color

    add_divider(doc)


def _parse_fehler_detail(entry: str) -> tuple[str, str] | None:
    """Parst einen fehler_detail-String 'Fehler: X -> Korrektur: Y' in (fehler, korrektur)."""
    # Normalisiere Pfeil-Varianten
    normalized = entry.replace("->", "→").replace("–>", "→")
    if "→" not in normalized:
        return None

    left, right = normalized.split("→", 1)

    def _extract(part: str, keyword: str) -> str:
        """Entfernt 'Fehler:' / 'Korrektur:'-Prefix und Anführungszeichen."""
        # Entferne bekannte Prefixe
        for kw in (keyword, "Fehler", "Korrektur", "Error", "Correction"):
            idx = part.lower().find(kw.lower())
            if idx != -1:
                part = part[idx + len(kw):]
        # Entferne führende Sonderzeichen und Anführungszeichen
        return part.strip().strip(":").strip().strip("'\"").strip().rstrip("'\".,")

    fehler = _extract(left, "Fehler")
    korrektur = _extract(right, "Korrektur")
    if fehler and korrektur:
        return fehler, korrektur
    return None


def add_student_text_section(
    doc: Document, original_text: str, data: FeedbackData
) -> list[Any]:
    """Fügt den Originaltext mit Inline-Korrekturen ein. Gibt die Text-Paragraphen zurück."""
    add_section_header(doc, "ORIGINALTEXT")

    paragraphs_added: list[Any] = []
    raw_paragraphs = [p for p in original_text.split("\n") if p.strip()]

    for para_idx, raw_para in enumerate(raw_paragraphs):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)

        is_title = (
            para_idx == 0
            and len(raw_para) <= 80
            and not raw_para.rstrip().endswith((".", "!", "?", ",", ";", ":"))
        )

        r = para.add_run(raw_para)
        r.font.size = Pt(12 if is_title else 10)
        r.bold = is_title

        paragraphs_added.append(para)

    add_divider(doc)
    return paragraphs_added


def _attach_word_comments(
    doc: Document,
    paragraph_comment_pairs: list[tuple[Any, str]],
    author: str = "Natascha",
) -> None:
    """Hängt Word-Balloon-Kommentare an Absätze via direkter XML-Manipulation."""
    if not _LXML_AVAILABLE or not paragraph_comment_pairs:
        return

    try:
        from docx.opc.packuri import PackURI
        from docx.opc.part import Part
    except ImportError:
        return

    root = _etree.Element(f"{{{_W_NS}}}comments", nsmap={"w": _W_NS})

    for idx, (para, comment_text) in enumerate(paragraph_comment_pairs):
        c_el = _etree.SubElement(root, f"{{{_W_NS}}}comment")
        c_el.set(f"{{{_W_NS}}}id", str(idx))
        c_el.set(f"{{{_W_NS}}}author", author)
        c_el.set(f"{{{_W_NS}}}date", f"{date.today().isoformat()}T00:00:00Z")
        c_el.set(f"{{{_W_NS}}}initials", "".join(w[0] for w in author.split()[:2] if w))

        c_p = _etree.SubElement(c_el, f"{{{_W_NS}}}p")
        c_pPr = _etree.SubElement(c_p, f"{{{_W_NS}}}pPr")
        c_pStyle = _etree.SubElement(c_pPr, f"{{{_W_NS}}}pStyle")
        c_pStyle.set(f"{{{_W_NS}}}val", "CommentText")
        c_r = _etree.SubElement(c_p, f"{{{_W_NS}}}r")
        c_rPr = _etree.SubElement(c_r, f"{{{_W_NS}}}rPr")
        c_rStyle = _etree.SubElement(c_rPr, f"{{{_W_NS}}}rStyle")
        c_rStyle.set(f"{{{_W_NS}}}val", "CommentReference")
        c_t = _etree.SubElement(c_r, f"{{{_W_NS}}}t")
        c_t.text = comment_text
        c_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        p_xml = para._p
        # commentRangeStart vor dem ersten Child
        crs = OxmlElement("w:commentRangeStart")
        crs.set(qn("w:id"), str(idx))
        p_xml.insert(0, crs)
        # commentRangeEnd und commentReference am Ende
        cre = OxmlElement("w:commentRangeEnd")
        cre.set(qn("w:id"), str(idx))
        p_xml.append(cre)
        r_ref = OxmlElement("w:r")
        rPr_ref = OxmlElement("w:rPr")
        rStyle_ref = OxmlElement("w:rStyle")
        rStyle_ref.set(qn("w:val"), "CommentReference")
        rPr_ref.append(rStyle_ref)
        r_ref.append(rPr_ref)
        cr = OxmlElement("w:commentReference")
        cr.set(qn("w:id"), str(idx))
        r_ref.append(cr)
        p_xml.append(r_ref)

    xml_bytes = _etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    comments_part = Part(
        PackURI("/word/comments.xml"),
        _COMMENTS_CONTENT_TYPE,
        xml_bytes,
    )
    doc.part.relate_to(comments_part, _COMMENTS_REL)


def is_punctuation_only_fix(original: str, korrektur: str) -> bool:
    """True wenn sich Original und Korrektur nur durch Satzzeichen unterscheiden."""

    def strip_punct(s: str) -> str:
        return re.sub(r"[,;:\.!\?\-–]", "", s).strip()

    return bool(original) and strip_punct(original) == strip_punct(korrektur)


def _find_inserted_punct(original: str, korrektur: str) -> str:
    """Gibt die in korrektur neu eingefügten Satzzeichen zurück (vs. original)."""
    orig_chars = re.findall(r"[,;:\.!\?\-–]", original)
    korr_chars = re.findall(r"[,;:\.!\?\-–]", korrektur)
    extra = list(korr_chars)
    for c in orig_chars:
        try:
            extra.remove(c)
        except ValueError:
            pass
    return "".join(extra) or ","


def _find_zitat_in_paragraph(para, zitat_text: str):
    """Find a quote within a paragraph, handling cross-run boundaries.

    Returns (start_run_idx, end_run_idx, char_offset_in_start, char_end_in_end) or None.
    """
    runs = para.runs
    if not runs:
        return None

    full_text = "".join(r.text for r in runs)
    pos = full_text.find(zitat_text)
    if pos == -1:
        return None

    end_pos = pos + len(zitat_text)

    start_run_idx = 0
    char_offset = 0
    accumulated = 0
    for i, run in enumerate(runs):
        if accumulated + len(run.text) > pos:
            start_run_idx = i
            char_offset = pos - accumulated
            break
        accumulated += len(run.text)

    end_run_idx = start_run_idx
    accumulated = 0
    for i, run in enumerate(runs):
        if accumulated + len(run.text) >= end_pos:
            end_run_idx = i
            break
        accumulated += len(run.text)

    char_end_in_end = end_pos - sum(len(runs[j].text) for j in range(end_run_idx))

    return (start_run_idx, end_run_idx, char_offset, char_end_in_end)


def _split_run_at(run, offset: int):
    """Split a run at the given character offset. Returns (before_run, after_run)."""
    if offset <= 0 or offset >= len(run.text):
        return None, None

    before_text = run.text[:offset]
    after_text = run.text[offset:]

    run.text = before_text

    new_run = OxmlElement("w:r")
    rpr = run._r.find(f"{{{_W_NS}}}rPr")
    if rpr is not None:
        import copy
        new_run.append(copy.deepcopy(rpr))

    t_el = OxmlElement("w:t")
    t_el.text = after_text
    t_el.set(qn("xml:space"), "preserve")
    new_run.append(t_el)

    run._r.addnext(new_run)

    from docx.text.run import Run
    after_r = Run(new_run, run._parent)

    return run, after_r


def _make_comment_reference_element(comment_id: int) -> OxmlElement:
    """Create a w:commentReference element wrapped in a w:r with CommentReference style."""
    r_ref = OxmlElement("w:r")
    rPr_ref = OxmlElement("w:rPr")
    rStyle_ref = OxmlElement("w:rStyle")
    rStyle_ref.set(qn("w:val"), "CommentReference")
    rPr_ref.append(rStyle_ref)
    r_ref.append(rPr_ref)
    cr = OxmlElement("w:commentReference")
    cr.set(qn("w:id"), str(comment_id))
    r_ref.append(cr)
    return r_ref


def _attach_word_level_comments(
    doc: Document,
    hinweise_list: list[WordLevelHinweis],
    extra_strings: list[str] | None = None,
    author: str = "Natascha",
    student_paragraphs: list | None = None,
    fehler_list: list[SprachFehler] | None = None,
) -> None:
    """Anchors Word comments at the Run level (individual words/phrases).

    Correct Word XML structure for a word-level comment:
      <w:p>
        <w:r><w:t>before </w:t></w:r>
        <w:commentRangeStart w:id="0"/>
        <w:r><w:t>quote</w:t><w:commentReference w:id="0"/></w:r>
        <w:commentRangeEnd w:id="0"/>
        <w:r><w:t> after</w:t></w:r>
      </w:p>
    """
    if not _LXML_AVAILABLE:
        return

    try:
        from docx.opc.packuri import PackURI
        from docx.opc.part import Part
    except ImportError:
        return


    root = _etree.Element(f"{{{_W_NS}}}comments", nsmap={"w": _W_NS})

    comment_id = 0
    paragraph_fallbacks: list[tuple[Any, str]] = []
    _used_para_ids: set[int] = set()

    text_paragraphs = [p for p in doc.paragraphs if p.text.strip()]

    for hinweis in hinweise_list:
        if hinweis.zitat.strip():
            found = False
            for para in text_paragraphs:
                result = _find_zitat_in_paragraph(para, hinweis.zitat)
                if result is None:
                    continue

                start_idx, end_idx, start_off, end_off = result

                c_el = _etree.SubElement(root, f"{{{_W_NS}}}comment")
                c_el.set(f"{{{_W_NS}}}id", str(comment_id))
                c_el.set(f"{{{_W_NS}}}author", author)
                c_el.set(f"{{{_W_NS}}}date", f"{date.today().isoformat()}T00:00:00Z")
                c_el.set(f"{{{_W_NS}}}initials", "".join(w[0] for w in author.split()[:2] if w))

                c_p = _etree.SubElement(c_el, f"{{{_W_NS}}}p")
                c_pPr = _etree.SubElement(c_p, f"{{{_W_NS}}}pPr")
                c_pStyle = _etree.SubElement(c_pPr, f"{{{_W_NS}}}pStyle")
                c_pStyle.set(f"{{{_W_NS}}}val", "CommentText")
                c_r = _etree.SubElement(c_p, f"{{{_W_NS}}}r")
                c_rPr = _etree.SubElement(c_r, f"{{{_W_NS}}}rPr")
                c_rStyle = _etree.SubElement(c_rPr, f"{{{_W_NS}}}rStyle")
                c_rStyle.set(f"{{{_W_NS}}}val", "CommentReference")
                c_t = _etree.SubElement(c_r, f"{{{_W_NS}}}t")
                c_t.text = hinweis.kommentar
                c_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

                if start_idx == end_idx:
                    # Zitat liegt komplett in einem Run
                    start_run = para.runs[start_idx]
                    end_run = start_run

                    if start_off > 0:
                        before, after = _split_run_at(start_run, start_off)
                        start_run = after
                        end_run = after

                    adjusted_end = end_off - start_off
                    if adjusted_end < len(end_run.text):
                        before, after = _split_run_at(end_run, adjusted_end)
                        end_run = before

                    crs = OxmlElement("w:commentRangeStart")
                    crs.set(qn("w:id"), str(comment_id))
                    start_run._r.addprevious(crs)

                    cre = OxmlElement("w:commentRangeEnd")
                    cre.set(qn("w:id"), str(comment_id))
                    end_run._r.addnext(cre)

                    r_ref = _make_comment_reference_element(comment_id)
                    cre.addnext(r_ref)
                else:
                    # Zitat spannt über mehrere Runs
                    start_run = para.runs[start_idx]
                    end_run = para.runs[end_idx]

                    if start_off > 0:
                        before, after = _split_run_at(start_run, start_off)
                        start_run = after

                    adjusted_end = end_off  # bereits relativ zum End-Run
                    if adjusted_end < len(end_run.text):
                        before, after = _split_run_at(end_run, adjusted_end)
                        end_run = before

                    crs = OxmlElement("w:commentRangeStart")
                    crs.set(qn("w:id"), str(comment_id))
                    start_run._r.addprevious(crs)

                    cre = OxmlElement("w:commentRangeEnd")
                    cre.set(qn("w:id"), str(comment_id))
                    end_run._r.addnext(cre)

                    r_ref = _make_comment_reference_element(comment_id)
                    cre.addnext(r_ref)

                comment_id += 1
                found = True
                _used_para_ids.add(id(para))
                break

            if not found:
                paragraph_fallbacks.append((None, hinweis.kommentar))
        else:
            paragraph_fallbacks.append((None, hinweis.kommentar))

    for s in (extra_strings or []):
        paragraph_fallbacks.append((None, s))

    fallback_pool = student_paragraphs if student_paragraphs is not None else text_paragraphs
    for _, comment_text in paragraph_fallbacks:
        para = next(
            (p for p in fallback_pool if id(p) not in _used_para_ids),
            None,
        )
        if para is None:
            break
        _used_para_ids.add(id(para))

        c_el = _etree.SubElement(root, f"{{{_W_NS}}}comment")
        c_el.set(f"{{{_W_NS}}}id", str(comment_id))
        c_el.set(f"{{{_W_NS}}}author", author)
        c_el.set(f"{{{_W_NS}}}date", f"{date.today().isoformat()}T00:00:00Z")
        c_el.set(f"{{{_W_NS}}}initials", "".join(w[0] for w in author.split()[:2] if w))

        c_p = _etree.SubElement(c_el, f"{{{_W_NS}}}p")
        c_pPr = _etree.SubElement(c_p, f"{{{_W_NS}}}pPr")
        c_pStyle = _etree.SubElement(c_pPr, f"{{{_W_NS}}}pStyle")
        c_pStyle.set(f"{{{_W_NS}}}val", "CommentText")
        c_r = _etree.SubElement(c_p, f"{{{_W_NS}}}r")
        c_rPr = _etree.SubElement(c_r, f"{{{_W_NS}}}rPr")
        c_rStyle = _etree.SubElement(c_rPr, f"{{{_W_NS}}}rStyle")
        c_rStyle.set(f"{{{_W_NS}}}val", "CommentReference")
        c_t = _etree.SubElement(c_r, f"{{{_W_NS}}}t")
        c_t.text = comment_text
        c_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        p_xml = para._p
        crs = OxmlElement("w:commentRangeStart")
        crs.set(qn("w:id"), str(comment_id))
        pPr = p_xml.find(f"{{{_W_NS}}}pPr")
        if pPr is not None:
            pPr.addnext(crs)
        else:
            p_xml.insert(0, crs)
        cre = OxmlElement("w:commentRangeEnd")
        cre.set(qn("w:id"), str(comment_id))
        p_xml.append(cre)
        r_ref = _make_comment_reference_element(comment_id)
        p_xml.append(r_ref)

        comment_id += 1

    # --- Fehler: farbiges Highlight + Balloon-Kommentar ---
    fehler_search_pool = student_paragraphs if student_paragraphs else text_paragraphs
    for fehler in (fehler_list or []):
        try:
            zitat = fehler.zitat.strip()
            if not zitat:
                continue

            for para in fehler_search_pool:
                # Case-insensitive zitat search: find position in lowercased text,
                # then retrieve actual text slice for exact matching
                runs = para.runs
                if not runs:
                    continue
                full_text = "".join(r.text for r in runs)
                low_pos = full_text.lower().find(zitat.lower())
                if low_pos == -1:
                    continue
                actual_zitat = full_text[low_pos: low_pos + len(zitat)]
                result = _find_zitat_in_paragraph(para, actual_zitat)
                if result is None:
                    continue

                start_idx, end_idx, start_off, end_off = result

                kommentar_text = (
                    f"[{fehler.typ}] {fehler.erklaerung}" if fehler.erklaerung
                    else f"[{fehler.typ}]"
                )

                c_el = _etree.SubElement(root, f"{{{_W_NS}}}comment")
                c_el.set(f"{{{_W_NS}}}id", str(comment_id))
                c_el.set(f"{{{_W_NS}}}author", author)
                c_el.set(f"{{{_W_NS}}}date", f"{date.today().isoformat()}T00:00:00Z")
                c_el.set(f"{{{_W_NS}}}initials", "".join(w[0] for w in author.split()[:2] if w))
                c_p = _etree.SubElement(c_el, f"{{{_W_NS}}}p")
                c_pPr = _etree.SubElement(c_p, f"{{{_W_NS}}}pPr")
                c_pStyle = _etree.SubElement(c_pPr, f"{{{_W_NS}}}pStyle")
                c_pStyle.set(f"{{{_W_NS}}}val", "CommentText")
                c_r = _etree.SubElement(c_p, f"{{{_W_NS}}}r")
                c_rPr = _etree.SubElement(c_r, f"{{{_W_NS}}}rPr")
                c_rStyle = _etree.SubElement(c_rPr, f"{{{_W_NS}}}rStyle")
                c_rStyle.set(f"{{{_W_NS}}}val", "CommentReference")
                c_t = _etree.SubElement(c_r, f"{{{_W_NS}}}t")
                c_t.text = kommentar_text
                c_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

                if fehler.typ == "Z":
                    # Zeichensetzungsfehler: nur Randkommentar, kein Inline-Markup
                    comment_id += 1
                    break

                if start_idx == end_idx:
                    start_run = para.runs[start_idx]
                    end_run = start_run
                    if start_off > 0:
                        before, after = _split_run_at(start_run, start_off)
                        if after is not None:
                            start_run = after
                        end_run = start_run
                    adjusted_end = end_off - start_off
                    if adjusted_end > 0 and adjusted_end < len(end_run.text):
                        before, after = _split_run_at(end_run, adjusted_end)
                        if before is not None:
                            end_run = before
                else:
                    start_run = para.runs[start_idx]
                    end_run = para.runs[end_idx]
                    if start_off > 0:
                        before, after = _split_run_at(start_run, start_off)
                        if after is not None:
                            start_run = after
                    if end_off > 0 and end_off < len(end_run.text):
                        before, after = _split_run_at(end_run, end_off)
                        if before is not None:
                            end_run = before

                # Entscheide ob Punkt/Komma-Fehler oder vollständige Korrektur
                punct_only = is_punctuation_only_fix(fehler.zitat, fehler.korrektur)

                if not punct_only:
                    # Track-Changes-Stil: Fehler rot + durchgestrichen
                    p_elem = para._p
                    in_range = False
                    for child in p_elem:
                        if child is start_run._r:
                            in_range = True
                        if in_range and child.tag == f"{{{_W_NS}}}r":
                            rPr = child.find(f"{{{_W_NS}}}rPr")
                            if rPr is None:
                                rPr = OxmlElement("w:rPr")
                                child.insert(0, rPr)
                            for old_hl in rPr.findall(f"{{{_W_NS}}}highlight"):
                                rPr.remove(old_hl)
                            c_xml = rPr.find(f"{{{_W_NS}}}color")
                            if c_xml is None:
                                c_xml = OxmlElement("w:color")
                                rPr.append(c_xml)
                            c_xml.set(qn("w:val"), "CC0000")
                            if rPr.find(f"{{{_W_NS}}}strike") is None:
                                rPr.append(OxmlElement("w:strike"))
                        if child is end_run._r:
                            in_range = False
                            break

                # Grüner Korrektur-Run: bei Satzzeichen nur das Zeichen, sonst volle Korrektur
                korr_text = (
                    _find_inserted_punct(fehler.zitat, fehler.korrektur)
                    if punct_only
                    else f" {fehler.korrektur}"
                )
                korr_r = OxmlElement("w:r")
                korr_rPr = OxmlElement("w:rPr")
                korr_color = OxmlElement("w:color")
                korr_color.set(qn("w:val"), "007000")
                korr_rPr.append(korr_color)
                korr_sz = OxmlElement("w:sz")
                korr_sz.set(qn("w:val"), "18")  # 9 pt
                korr_rPr.append(korr_sz)
                korr_szCs = OxmlElement("w:szCs")
                korr_szCs.set(qn("w:val"), "18")
                korr_rPr.append(korr_szCs)
                korr_rPr.append(OxmlElement("w:b"))
                korr_r.append(korr_rPr)
                korr_t = OxmlElement("w:t")
                korr_t.set(qn("xml:space"), "preserve")
                korr_t.text = korr_text
                korr_r.append(korr_t)
                end_run._r.addnext(korr_r)

                # Comment range wraps only the strikethrough text
                crs = OxmlElement("w:commentRangeStart")
                crs.set(qn("w:id"), str(comment_id))
                start_run._r.addprevious(crs)

                # commentRangeEnd goes between end_run and the correction run
                cre = OxmlElement("w:commentRangeEnd")
                cre.set(qn("w:id"), str(comment_id))
                end_run._r.addnext(cre)

                r_ref = _make_comment_reference_element(comment_id)
                cre.addnext(r_ref)

                comment_id += 1
                break
        except Exception as exc:
            logging.warning("Fehler-Markierung übersprungen: %r — %s", fehler.zitat, exc)
            continue

    xml_bytes = _etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    comments_part = Part(
        PackURI("/word/comments.xml"),
        _COMMENTS_CONTENT_TYPE,
        xml_bytes,
    )
    doc.part.relate_to(comments_part, _COMMENTS_REL)


def parse_args() -> argparse.Namespace:
    """Liest die CLI-Argumente ein."""

    parser = argparse.ArgumentParser(
        description="Erzeugt Feedback-DOCX aus JSON-Dateien."
    )
    parser.add_argument(
        "--file", help="Verarbeitet nur eine JSON-Datei aus /output/feedback_data."
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Ueberschreibt bestehende _feedback.docx-Dateien.",
    )
    parser.add_argument(
        "--dry-run", action="store_true", help="Zeigt nur, was verarbeitet wuerde."
    )
    return parser.parse_args()


def canonical_key(raw_key: str) -> str:
    """Normalisiert Kriterienschluessel."""

    return ALIASES.get(raw_key, raw_key)


def ensure_list(value: Any, field_name: str) -> list[str]:
    """Stellt sicher, dass ein Feld eine String-Liste ist."""

    if value is None:
        return []
    if not isinstance(value, list) or not all(isinstance(item, str) for item in value):
        raise ValueError(f"Feld '{field_name}' muss eine Liste aus Strings sein.")
    return value


def parse_criterion(key: str, payload: dict[str, Any]) -> CriterionFeedback:
    """Validiert ein einzelnes Kriterium."""

    required = {"stufe", "punkte", "staerken", "schwaechen", "vorschlaege"}
    missing = sorted(required - payload.keys())
    if missing:
        raise ValueError(f"Kriterium '{key}' fehlt: {', '.join(missing)}")

    stufe = payload["stufe"]
    if not isinstance(stufe, (str, int, float)):
        raise ValueError(f"Kriterium '{key}' hat eine ungueltige Stufe.")

    punkte = payload["punkte"]
    if not isinstance(punkte, (int, float)):
        raise ValueError(f"Kriterium '{key}' hat ungueltige Punkte.")

    gewicht = payload.get("gewicht")
    if gewicht is not None and not isinstance(gewicht, (int, float)):
        raise ValueError(f"Kriterium '{key}' hat ein ungueltiges Gewicht.")

    return CriterionFeedback(
        key=canonical_key(key),
        stufe=str(stufe),
        punkte=float(punkte),
        gewicht=float(gewicht) if gewicht is not None else None,
        staerken=ensure_list(payload["staerken"], f"{key}.staerken"),
        schwaechen=ensure_list(payload["schwaechen"], f"{key}.schwaechen"),
        vorschlaege=ensure_list(payload["vorschlaege"], f"{key}.vorschlaege"),
        fehler_detail=ensure_list(payload.get("fehler_detail"), f"{key}.fehler_detail")
        or None,
        fehlerschwerpunkte=ensure_list(
            payload.get("fehlerschwerpunkte"), f"{key}.fehlerschwerpunkte"
        )
        or None,
        rhetorische_figuren=ensure_list(
            payload.get("rhetorische_figuren"), f"{key}.rhetorische_figuren"
        )
        or None,
    )


def parse_feedback_data(payload: dict[str, Any]) -> FeedbackData:
    """Validiert die JSON-Struktur und wandelt sie in Dataklassen um."""

    required = {
        "datei",
        "textsorte",
        "fach",
        "schulstufe",
        "rubrik",
        "bewertung",
    }
    missing = sorted(required - payload.keys())
    if missing:
        raise ValueError(f"Feedback-Datensatz fehlt: {', '.join(missing)}")

    if not isinstance(payload["bewertung"], dict) or not payload["bewertung"]:
        raise ValueError("Feld 'bewertung' muss ein nicht-leeres Objekt sein.")

    notenempfehlung: GradeRecommendation | None = None
    note_raw = payload.get("notenempfehlung")
    if note_raw and isinstance(note_raw, dict):
        note_required = {"durchschnitt", "note", "bezeichnung", "begruendung"}
        note_missing = sorted(note_required - note_raw.keys())
        if note_missing:
            raise ValueError(f"Notenempfehlung fehlt: {', '.join(note_missing)}")
        if not isinstance(note_raw["durchschnitt"], (int, float)):
            raise ValueError("'notenempfehlung.durchschnitt' muss numerisch sein.")
        if not isinstance(note_raw["note"], int):
            raise ValueError("'notenempfehlung.note' muss eine ganze Zahl sein.")
        if not isinstance(note_raw["bezeichnung"], str) or not isinstance(
            note_raw["begruendung"], str
        ):
            raise ValueError(
                "'notenempfehlung.bezeichnung' und 'begruendung' muessen Strings sein."
            )
        notenempfehlung = GradeRecommendation(
            durchschnitt=float(note_raw["durchschnitt"]),
            note=note_raw["note"],
            bezeichnung=note_raw["bezeichnung"],
            begruendung=note_raw["begruendung"],
            k1_note=note_raw.get("k1_note"),
            k3_note=note_raw.get("k3_note"),
            k1_schnitt=note_raw.get("k1_schnitt"),
            k3_schnitt=note_raw.get("k3_schnitt"),
            sonderregel=note_raw.get("sonderregel"),
        )

    criteria = [
        parse_criterion(key, value) for key, value in payload["bewertung"].items()
    ]

    raw_hinweise = payload.get("hinweise", [])
    hinweise: list[WordLevelHinweis] = []
    for item in raw_hinweise:
        if isinstance(item, str):
            hinweise.append(WordLevelHinweis(zitat="", kommentar=item))
        elif isinstance(item, dict):
            hinweise.append(WordLevelHinweis(
                zitat=str(item.get("zitat", "")),
                kommentar=str(item.get("kommentar", "")),
            ))

    raw_fehler = payload.get("fehler", [])
    fehler: list[SprachFehler] = [
        SprachFehler(
            zitat=str(e.get("zitat", "")),
            korrektur=str(e.get("korrektur", "")),
            typ=str(e.get("typ", "G")),
            erklaerung=str(e.get("erklaerung", "")),
        )
        for e in raw_fehler if isinstance(e, dict)
    ]

    return FeedbackData(
        datei=str(payload["datei"]),
        schueler=str(payload["schueler"]) if payload.get("schueler") else None,
        klasse=str(payload["klasse"]) if payload.get("klasse") else None,
        textsorte=str(payload["textsorte"]),
        fach=str(payload["fach"]),
        schulstufe=str(payload["schulstufe"]),
        rubrik=str(payload["rubrik"]),
        bewertung=criteria,
        notenempfehlung=notenempfehlung,
        hinweise=hinweise,
        fehler=fehler or None,
        modell=str(payload["modell"]) if payload.get("modell") else None,
        zusammenfassung=str(payload["zusammenfassung"]) if payload.get("zusammenfassung") else None,
        staerken_global=ensure_list(payload.get("staerken_global"), "staerken_global") or None,
        verbesserungsbereiche=(
            ensure_list(payload.get("verbesserungsbereiche"), "verbesserungsbereiche") or None
        ),
        srdp_detail=(
            payload.get("srdp_detail") if isinstance(payload.get("srdp_detail"), dict) else None
        ),
    )


def load_feedback_json(path: Path) -> FeedbackData:
    """Laedt und validiert eine JSON-Datei."""

    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Die JSON-Wurzel muss ein Objekt sein.")
    return parse_feedback_data(data)


def criterion_label(data: FeedbackData, key: str) -> str:
    """Liefert den Anzeige-Namen eines Kriteriums."""

    german_labels = {
        "inhalt": "INHALT",
        "textstruktur": "TEXTSTRUKTUR",
        "ausdruck": "STIL UND AUSDRUCK"
        if data.schulstufe == "Oberstufe"
        else "AUSDRUCK",
        "sprachrichtigkeit": "NORMATIVE SPRACHRICHTIGKEIT"
        if data.schulstufe == "Oberstufe"
        else "SPRACHRICHTIGKEIT",
    }
    english_labels = {
        "task_achievement": "TASK ACHIEVEMENT",
        "organisation_layout": "ORGANISATION AND LAYOUT",
        "lexical_range_accuracy": "LEXICAL RANGE AND ACCURACY",
        "grammatical_range_accuracy": "GRAMMATICAL RANGE AND ACCURACY",
    }
    labels = german_labels if data.fach == "Deutsch" else english_labels
    return labels.get(key, key.replace("_", " ").upper())


def ordered_criteria(data: FeedbackData) -> list[CriterionFeedback]:
    """Sortiert Kriterien in einer fachlich passenden Reihenfolge."""

    desired = GERMAN_ORDER if data.fach == "Deutsch" else ENGLISH_ORDER
    rank = {name: index for index, name in enumerate(desired)}
    return sorted(data.bewertung, key=lambda item: rank.get(item.key, len(rank)))


def render_list_section(
    doc: Document,
    title: str,
    entries: list[str] | None,
    color: RGBColor | None = None,
) -> None:
    """Rendert eine farbige Ueberschrift plus Liste oder Platzhalter."""

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    if color is not None:
        run.font.color.rgb = color
    paragraph.paragraph_format.space_after = Pt(1)
    if entries:
        for entry in entries:
            add_bullet(doc, entry, color=color)
    else:
        add_bullet(doc, "Keine Angaben.")


def add_summary_table(
    doc: Document, data: FeedbackData, bewertungsmodus: str = "benotet"
) -> None:
    """Fuegt eine Uebersichtstabelle mit allen Kriterien und der Gesamtnote ein."""

    criteria = ordered_criteria(data)
    extra_rows = 1 if data.notenempfehlung and bewertungsmodus == "benotet" else 0

    table = doc.add_table(rows=1 + len(criteria) + extra_rows, cols=3)
    table.style = "Table Grid"

    third_col = "Punkte" if bewertungsmodus == "benotet" else "Feedback"
    for cell, text in zip(table.rows[0].cells, ["Kriterium", "Stufe / Bewertung", third_col]):
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        _set_cell_bg(cell, "1F497D")

    for row_idx, crit in enumerate(criteria, start=1):
        cells = table.rows[row_idx].cells
        label = criterion_label(data, crit.key)
        value_col = f"{crit.punkte:g}" if bewertungsmodus == "benotet" else "-"
        for cell, text in zip(cells, [label, crit.stufe, value_col]):
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        if row_idx % 2 == 0:
            for cell in cells:
                _set_cell_bg(cell, "EEF2FF")

    if data.notenempfehlung and bewertungsmodus == "benotet":
        note = data.notenempfehlung.note
        grade_color = _note_color(note)
        grade_cells = table.rows[-1].cells
        for cell, text in zip(
            grade_cells,
            ["GESAMTNOTE", "", f"{note} – {data.notenempfehlung.bezeichnung}"],
        ):
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            _set_cell_bg(cell, str(grade_color))

    doc.add_paragraph()


def add_fehlerprotokoll(doc: Document, fehler_list: list[SprachFehler]) -> None:
    """Rendert eine Fehlerprotokoll-Tabelle mit allen Sprachfehlern."""
    add_section_header(doc, "FEHLERPROTOKOLL")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, col_title in enumerate(("Nr.", "Typ", "Zitat", "Korrektur", "Erklärung")):
        p = hdr[i].paragraphs[0]
        run = p.add_run(col_title)
        run.bold = True
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:fill"), "DDDDDD")
        hdr[i]._tc.get_or_add_tcPr().append(shading)

    for nr, fehler in enumerate(fehler_list, start=1):
        row = table.add_row().cells
        row[0].text = str(nr)
        typ_p = row[1].paragraphs[0]
        typ_run = typ_p.add_run(fehler.typ)
        typ_run.bold = True
        typ_run.font.color.rgb = _FEHLER_TYP_COLORS.get(fehler.typ, RGBColor(0, 0, 0))
        row[2].text = fehler.zitat
        row[3].text = fehler.korrektur
        row[4].text = fehler.erklaerung

    counts: dict[str, int] = {"R": 0, "G": 0, "Z": 0, "A": 0}
    for f in fehler_list:
        counts[f.typ] = counts.get(f.typ, 0) + 1
    summary = (
        f"R: {counts['R']} | G: {counts['G']} | Z: {counts['Z']} | A: {counts['A']}"
        f" — Gesamt: {len(fehler_list)} Fehler"
    )
    p = doc.add_paragraph()
    run = p.add_run(summary)
    run.bold = True
    run.font.size = Pt(10)


def _add_srdp_table(
    doc: Document,
    data: dict[str, Any],
    kriterien: list[tuple[str, str]],
) -> None:
    """Erstellt eine SRDP-Tabelle mit Kriterium | Einschätzung | Begründung."""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for cell, heading in zip(hdr, ("Kriterium", "Einschätzung", "Begründung")):
        cell.text = heading
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    table.columns[0].width = Cm(3.5)
    table.columns[1].width = Cm(5.0)
    table.columns[2].width = Cm(8.5)

    for key, label in kriterien:
        entry = data.get(key) or {}
        stufe = int(entry.get("stufe", 0))
        begruendung = str(entry.get("begruendung", ""))
        einschaetzung = SRDP_WORTLAUT.get(key, {}).get(stufe, f"Stufe {stufe}")

        row = table.add_row()

        p = row.cells[0].paragraphs[0]
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(9)

        p = row.cells[1].paragraphs[0]
        r = p.add_run(einschaetzung)
        r.italic = True
        r.font.size = Pt(9)

        p = row.cells[2].paragraphs[0]
        r = p.add_run(begruendung)
        r.font.size = Pt(9)


def add_srdp_raster(doc: Document, srdp_detail: dict[str, Any]) -> None:
    """Fügt den ausgefüllten SRDP-Beurteilungsraster als Tabellen ins Dokument ein."""
    doc.add_page_break()
    add_section_header(doc, "BEURTEILUNG NACH SRDP-RASTER")

    gesamteindruck = srdp_detail.get("gesamteindruck", "")
    if gesamteindruck:
        p = doc.add_paragraph()
        r = p.add_run(gesamteindruck)
        r.font.size = Pt(10)
        r.italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("K1 – Inhalt").bold = True
    _add_srdp_table(doc, srdp_detail.get("k1_inhalt") or {}, [
        ("schreibhandlung", "Schreibhandlung(en)"),
        ("arbeitsauftraege", "Arbeitsaufträge"),
        ("textbeilage", "Textbeilage(n)"),
        ("sachlich", "Sachliche Richtigkeit"),
        ("qualitaet", "Qualität der Auseinandersetzung"),
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("K1 – Textstruktur").bold = True
    _add_srdp_table(doc, srdp_detail.get("k1_textstruktur") or {}, [
        ("kohaerenz", "Kohärenz"),
        ("bezugnahme", "Bezugnahme auf Textbeilage(n)"),
        ("kohaesion", "Kohäsionsmittel"),
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("K3/1 – Stil und Ausdruck").bold = True
    _add_srdp_table(doc, srdp_detail.get("k3_stil") or {}, [
        ("situationsadaequat", "Situationsadäquatheit"),
        ("wortwahl", "Wortwahl / Ausdruck"),
        ("satzstrukturen", "Satzstrukturen"),
        ("eigenstaendigkeit", "Eigenständigkeit"),
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("K3/1 – Sprachnormen").bold = True
    _add_srdp_table(doc, srdp_detail.get("k3_sprachnormen") or {}, [
        ("orthografie", "Orthografie"),
        ("zeichensetzung", "Zeichensetzung"),
        ("grammatik", "Grammatik"),
    ])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Zentrale Verbesserungshinweise").bold = True
    for key, label in [
        ("verbesserung_inhaltlich", "Inhaltlich"),
        ("verbesserung_strukturell", "Strukturell"),
        ("verbesserung_sprachlich", "Sprachlich"),
    ]:
        text = srdp_detail.get(key, "")
        if text:
            p = doc.add_paragraph()
            r = p.add_run(f"{label}: ")
            r.bold = True
            r.font.size = Pt(10)
            r2 = p.add_run(text)
            r2.font.size = Pt(10)


def build_feedback_document(
    data: FeedbackData,
    config: dict[str, Any] | None = None,
    original_text: str | None = None,
    source_note: str | None = None,
    aufgabe_label: str = "",
    bewertungsmodus: str = "benotet",
) -> Document:
    """Erstellt das DOCX-Dokument im Stil einer authentischen Lehrkraft-Korrektur."""

    doc = Document()
    setup_page(doc, config)

    cfg = config or {}
    docx_cfg = cfg.get("docx", {})
    teacher_name = docx_cfg.get("teacher_name", "Lehrkraft")

    word_count = len(original_text.split()) if original_text else 0
    add_korrektur_header(doc, data, config, word_count=word_count, aufgabe_label=aufgabe_label)
    add_legende(doc)

    # Originaltext mit Inline-Korrekturen und Word-Balloons
    text_paragraphs: list[Any] = []
    if original_text and original_text.strip():
        text_paragraphs = add_student_text_section(doc, original_text, data)

        # Balloon-Kommentare: hinweise (Wortebene) + vorschläge (Absatz-Ebene)
        extra_comment_strings: list[str] = []
        for crit in ordered_criteria(data):
            for v in (crit.vorschlaege or []):
                extra_comment_strings.append(v)

        _attach_word_level_comments(
            doc, data.hinweise or [], author=teacher_name,
            student_paragraphs=text_paragraphs,
            fehler_list=data.fehler or [],
        )
    else:
        # Kein Originaltext verfügbar → Hinweis (ggf. mit Quelldatei-Info)
        msg = source_note or "(Originaltext nicht verfügbar — Input-DOCX nicht gefunden.)"
        add_body(doc, msg)
        add_divider(doc)

    # LEHRKRAFT-KOMMENTAR (Zusammenfassung, Stärken, Note)
    add_lehrer_kommentar_block(doc, data, bewertungsmodus=bewertungsmodus)

    # Detaillierte Kriterien-Auswertung
    if bewertungsmodus == "benotet":
        add_summary_table(doc, data, bewertungsmodus=bewertungsmodus)
        add_divider(doc)

    for index, criterion in enumerate(ordered_criteria(data), start=1):
        add_section_header(doc, f"{index}. {criterion_label(data, criterion.key)}")
        if bewertungsmodus == "benotet":
            points_text = f"{criterion.punkte:g} Punkte"
            if criterion.gewicht is not None:
                points_text = f"{points_text} | Gewicht: {criterion.gewicht:g} %"
            add_label(doc, "Bewertung", f"{criterion.stufe} [{points_text}]")
        doc.add_paragraph()

        render_list_section(doc, "Stärken:", criterion.staerken, color=C_STRENGTH)
        doc.add_paragraph()

        if criterion.rhetorische_figuren:
            render_list_section(
                doc, "Rhetorische Figuren:", criterion.rhetorische_figuren
            )
            doc.add_paragraph()

        render_list_section(
            doc, "Schwächen / Fehler:", criterion.schwaechen, color=C_WEAKNESS
        )
        doc.add_paragraph()
        render_list_section(
            doc, "Verbesserungsvorschläge:", criterion.vorschlaege, color=C_SUGGESTION
        )
        add_divider(doc)

    if data.fehler:
        add_fehlerprotokoll(doc, data.fehler)
        add_divider(doc)

    if data.srdp_detail and bewertungsmodus == "benotet":
        add_srdp_raster(doc, data.srdp_detail)
        add_divider(doc)

    if data.notenempfehlung and bewertungsmodus == "benotet":
        has_k1k3 = (
            data.notenempfehlung.k1_note is not None
            or data.notenempfehlung.k3_note is not None
        )
        if has_k1k3:
            add_section_header(doc, "NOTENEMPFEHLUNG (SRDP-basiert)")
        else:
            add_section_header(doc, "NOTENEMPFEHLUNG")
        add_divider(doc)

        if has_k1k3:
            if data.notenempfehlung.k1_note is not None:
                k1_schnitt_str = (
                    f" [Stufe {data.notenempfehlung.k1_schnitt:.1f}]"
                    if data.notenempfehlung.k1_schnitt is not None
                    else ""
                )
                add_label(
                    doc, "K1 (Inhalt + Textstruktur)",
                    f"Note {data.notenempfehlung.k1_note}{k1_schnitt_str}",
                )
            if data.notenempfehlung.k3_note is not None:
                k3_schnitt_str = (
                    f" [Stufe {data.notenempfehlung.k3_schnitt:.1f}]"
                    if data.notenempfehlung.k3_schnitt is not None
                    else ""
                )
                add_label(
                    doc, "K3/1 (Stil + Sprachnormen)",
                    f"Note {data.notenempfehlung.k3_note}{k3_schnitt_str}",
                )
        else:
            for criterion in ordered_criteria(data):
                label = criterion_label(data, criterion.key).title()
                value = f"{criterion.stufe} ({criterion.punkte:g} Punkte)"
                if criterion.gewicht is not None:
                    value = f"{value} x {criterion.gewicht:g} %"
                add_label(doc, label, value)

        doc.add_paragraph()
        add_label(doc, "Durchschnitt", f"{data.notenempfehlung.durchschnitt:.2f}")

        note_paragraph = doc.add_paragraph()
        note_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        note_run = note_paragraph.add_run(
            f"Empfohlene Note: {data.notenempfehlung.note} – {data.notenempfehlung.bezeichnung}"
        )
        note_run.bold = True
        note_run.font.size = Pt(12)
        note_run.font.color.rgb = _note_color(data.notenempfehlung.note)

        if data.notenempfehlung.sonderregel:
            doc.add_paragraph()
            add_body(
                doc,
                f"Sonderregel: {data.notenempfehlung.sonderregel} – automatisch Nicht genügend.",
            )

        doc.add_paragraph()
        add_body(doc, "Begründung:")
        add_body(doc, data.notenempfehlung.begruendung)
        doc.add_paragraph()
        add_body(
            doc,
            "Hinweis: Diese Empfehlung wurde nach SRDP-Standard berechnet. "
            "Die endgültige Note liegt im Ermessen der Lehrkraft.",
        )
        add_divider(doc)
    else:
        add_section_header(doc, "HAUSAUFGABE – KEINE BENOTUNG")
        add_divider(doc)
        add_body(
            doc,
            "Gemäß österreichischem Schulrecht dürfen Hausaufgaben nicht benotet werden. "
            "Die obenstehende Rückmeldung dient ausschließlich der Förderung.",
        )
        add_divider(doc)
    return doc


def build_statistics_document(
    stats: dict[str, Any],
    config: dict[str, Any] | None = None,
    klasse_name: str = "",
) -> Document:
    """Erstellt ein A4-DOCX-Dokument mit Klassen-Statistiken.

    Args:
        stats: Rückgabewert von compute_statistics()
        config: natascha_config.toml-Dict (optional, für Header/Footer)
        klasse_name: Anzeigename der Klasse (optional)
    """
    _NOTE_LABELS_MAP = {
        1: "Sehr gut",
        2: "Gut",
        3: "Befriedigend",
        4: "Genügend",
        5: "Nicht gen.",
    }
    doc = Document()
    setup_page(doc, config)
    add_heading(doc, "KLASSENAUSWERTUNG", level=1)

    # Meta-Tabelle
    meta = doc.add_table(rows=0, cols=2)
    meta.autofit = True
    today = date.today().strftime("%d.%m.%Y")
    meta_rows: list[tuple[str, str]] = []
    if klasse_name:
        meta_rows.append(("Klasse", klasse_name))
    meta_rows.append(("Datum", today))
    docx_cfg = (config or {}).get("docx", {})
    teacher = docx_cfg.get("teacher_name", "")
    school = docx_cfg.get("school_name", "")
    if teacher:
        meta_rows.append(("Lehrer/in", teacher))
    if school:
        meta_rows.append(("Schule", school))
    for label, value in meta_rows:
        row = meta.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

    # Übersicht
    total = stats.get("total", 0)
    avg = stats.get("grade_average", 0.0)
    p_total = doc.add_paragraph()
    r_total = p_total.add_run(f"Ausgewertete Schüler/innen: {total}")
    r_total.bold = True
    r_total.font.size = Pt(11)
    if total > 0:
        p_avg = doc.add_paragraph()
        r_avg = p_avg.add_run(f"Gesamtdurchschnitt: {avg:.2f}")
        r_avg.bold = True
        r_avg.font.size = Pt(11)
        r_avg.font.color.rgb = _note_color(max(1, min(5, round(avg))))
    doc.add_paragraph()

    # Notenverteilung
    add_heading(doc, "Notenverteilung", level=2)
    dist = stats.get("grade_distribution", {})
    grade_table = doc.add_table(rows=1 + 5, cols=3)
    grade_table.autofit = True
    hrow = grade_table.rows[0]
    hrow.cells[0].text = "Note"
    hrow.cells[1].text = "Bezeichnung"
    hrow.cells[2].text = f"Anzahl (von {total})"
    for cell in hrow.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = C_PRIMARY
    for i, note in enumerate(range(1, 6), 1):
        count = dist.get(note, 0)
        pct = (count / total * 100) if total > 0 else 0.0
        row = grade_table.rows[i]
        row.cells[0].text = str(note)
        row.cells[1].text = _NOTE_LABELS_MAP.get(note, "")
        row.cells[2].text = f"{count}  ({pct:.1f}%)"
        color = _note_color(note)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = color
    doc.add_paragraph()

    # Kriterien-Auswertung
    crit_avgs = stats.get("criteria_averages", {})
    if crit_avgs:
        add_heading(doc, "Kriterien-Auswertung", level=2)
        weakest = stats.get("weakest_criterion")
        strongest = stats.get("strongest_criterion")
        crit_table = doc.add_table(rows=1 + len(crit_avgs), cols=4)
        crit_table.autofit = True
        hrow = crit_table.rows[0]
        hrow.cells[0].text = "Kriterium"
        hrow.cells[1].text = "Ø Punkte"
        hrow.cells[2].text = "Min"
        hrow.cells[3].text = "Max"
        for cell in hrow.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = C_PRIMARY
        for i, (key, vals) in enumerate(
            sorted(crit_avgs.items(), key=lambda x: x[1]["avg"]), 1
        ):
            row = crit_table.rows[i]
            label = key.replace("_", " ").title()
            row.cells[0].text = label
            row.cells[1].text = f"{vals['avg']:.2f}"
            row.cells[2].text = f"{vals['min']:.1f}"
            row.cells[3].text = f"{vals['max']:.1f}"
            if key == weakest:
                highlight = C_GRADE_5
            elif key == strongest:
                highlight = C_GRADE_1
            else:
                highlight = None
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        if highlight:
                            run.font.color.rgb = highlight
                            run.bold = True

    return doc


def output_filename(original_name: str) -> str:
    """Leitet den Zieldateinamen aus der Originaldatei ab."""

    return f"{Path(original_name).stem}_feedback.docx"


def log_error(log_path: Path, message: str) -> None:
    """Schreibt eine Fehlermeldung ins Fehlerlog."""

    log_path.parent.mkdir(parents=True, exist_ok=True)
    with log_path.open("a", encoding="utf-8") as handle:
        handle.write(f"{message}\n")


def collect_json_files(paths: ProjectPaths, selected_file: str | None) -> list[Path]:
    """Bestimmt die zu verarbeitenden JSON-Dateien."""

    if selected_file:
        candidate = Path(selected_file)
        if not candidate.is_absolute():
            candidate = paths.feedback_data_dir / selected_file
        if not candidate.exists():
            raise FileNotFoundError(f"JSON-Datei nicht gefunden: {candidate}")
        return [candidate]
    return sorted(paths.feedback_data_dir.glob("*.json"))


def process_file(
    json_path: Path, paths: ProjectPaths, force: bool = False, dry_run: bool = False
) -> str:
    """Verarbeitet eine einzelne JSON-Datei."""

    raw_payload = json.loads(json_path.read_text(encoding="utf-8"))
    bewertungsmodus: str = (
        raw_payload.get("bewertungsmodus", "benotet")
        if isinstance(raw_payload, dict)
        else "benotet"
    )
    feedback = parse_feedback_data(raw_payload)
    output_path = paths.output_dir / output_filename(feedback.datei)

    if output_path.exists() and not force:
        return f"Übersprungen (existiert bereits): {output_path.name}"

    if dry_run:
        return f"Dry-run: würde {json_path.name} -> {output_path.name} verarbeiten"

    paths.output_dir.mkdir(parents=True, exist_ok=True)

    # Versuche Originaltext aus Input-DOCX zu lesen
    original_text: str | None = None
    input_docx = paths.input_dir / feedback.datei
    if input_docx.exists():
        try:
            from docx import Document as _DocxDoc
            _d = _DocxDoc(str(input_docx))
            original_text = "\n".join(p.text for p in _d.paragraphs)
        except Exception:
            logging.debug("Original-DOCX-Text nicht lesbar: %s", input_docx, exc_info=True)

    document = build_feedback_document(
        feedback, original_text=original_text, bewertungsmodus=bewertungsmodus
    )
    document.save(output_path)
    return f"Gespeichert: {output_path.name}"


def main() -> int:
    """CLI-Einstiegspunkt."""

    args = parse_args()
    paths = project_paths()
    paths.feedback_data_dir.mkdir(parents=True, exist_ok=True)
    paths.output_dir.mkdir(parents=True, exist_ok=True)

    if not args.dry_run:
        paths.fehlerlog.write_text("", encoding="utf-8")

    try:
        json_files = collect_json_files(paths, args.file)
    except Exception as exc:
        log_error(paths.fehlerlog, str(exc))
        print(str(exc))
        return 1

    if not json_files:
        print("Keine JSON-Dateien in /output/feedback_data gefunden.")
        return 0

    failures = 0
    for json_path in json_files:
        try:
            message = process_file(
                json_path, paths, force=args.force, dry_run=args.dry_run
            )
            print(message)
        except Exception as exc:
            failures += 1
            message = f"{json_path.name} - Fehler: {exc}"
            log_error(paths.fehlerlog, message)
            print(message)

    if failures:
        print(f"Abgeschlossen mit {failures} Fehler(n).")
        return 1

    print("Fertig.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
