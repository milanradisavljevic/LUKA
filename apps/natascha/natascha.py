#!/usr/bin/env python3
"""
NATASCHA Dashboard – Textual-basiertes Korrektur-Dashboard.

Drei-Spalten-Layout mit Dateiliste, Zuordnung und Vorschau.
Tastatur-Navigation, modale Dialoge, asynchrone Analyse.
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import re
import subprocess
import sys
import threading
import time
from dataclasses import dataclass
from enum import Enum
from functools import lru_cache
from pathlib import Path
from typing import Any, Literal

if sys.version_info < (3, 11):
    print("Python 3.11+ wird benoetigt.")
    raise SystemExit(1)

try:
    from textual import work
    from textual.app import App, ComposeResult
    from textual.binding import Binding
    from textual.containers import Container, Horizontal, Vertical, VerticalScroll
    from textual.reactive import reactive
    from textual.screen import ModalScreen
    from textual.widget import Widget
    from textual.widgets import (
        Button,
        Collapsible,
        DataTable,
        Input,
        Label,
        ListItem,
        ListView,
        Markdown,
        ProgressBar,
        RadioButton,
        RadioSet,
        Select,
        Static,
        TabbedContent,
        TabPane,
        TextArea,
    )
except ImportError:
    print("textual fehlt: pip install textual")
    raise SystemExit(1)

from rich.text import Text

sys.path.insert(0, str(Path(__file__).resolve().parent))
import generate_feedback as gf
import natascha_core as nc

try:
    import natascha_db as ndb
except Exception:
    ndb = None  # type: ignore[assignment]

try:
    import natascha_bridge  # Datei-Brücke → Lehrunterlagen-Tool (Phase 1)
except Exception:
    natascha_bridge = None  # type: ignore[assignment]

# Notenfarben (passend zum DOCX-Farbschema)
_NOTE_COLORS: dict[int, str] = {
    1: "#70C070",  # Hellgrün
    2: "#008000",  # Dunkelgrün
    3: "#C0A000",  # Gelb
    4: "#E06000",  # Orange
    5: "#C00000",  # Rot
}


def _safe_select_value(value: str, options: list[tuple[str, str]], default: str = "") -> str:
    """Gibt value zurück, wenn es in options enthalten ist, sonst default oder ersten Wert."""
    values = [v for _, v in options]
    if value in values:
        return value
    return default if default in values else (values[0] if values else "")


def _note_rich_text(note: int | str, label: str) -> Text:
    """Gibt ein farbig formatiertes Rich-Text-Objekt für die Notenzeile zurück."""
    color = _NOTE_COLORS.get(int(note) if str(note).isdigit() else 0, "#C00000")
    t = Text()
    t.append(label, style=f"bold {color}")
    return t


class FileStatus(Enum):
    PENDING = "pending"
    ANALYZED = "analyzed"
    DONE = "done"
    PROGRESS = "progress"
    ERROR = "error"


STATUS_SYMBOLS = {
    FileStatus.DONE: ("●", "status-done"),
    FileStatus.ANALYZED: ("●", "status-analyzed"),
    FileStatus.PROGRESS: ("◐", "status-progress"),
    FileStatus.PENDING: ("○", "status-pending"),
    FileStatus.ERROR: ("✗", "status-error"),
}


@dataclass
class FileInfo:
    path: Path
    word_count: int = 0  # Wörter für DOCX; KB-Größe für PDF/Bild
    status: FileStatus = FileStatus.PENDING
    fach: str = ""
    schulstufe: str = ""
    textsorte: str = ""
    rubric: str = ""
    schueler: str = ""
    analysis: dict[str, Any] | None = None
    marked: bool = False
    file_type: Literal["docx", "pdf", "image"] = "docx"
    bewertungsmodus: Literal["benotet", "unbenotet"] = "benotet"
    erwartungshorizont: str = ""


def _detect_file_type(path: Path) -> Literal["docx", "pdf", "image"]:
    suffix = path.suffix.lower()
    if suffix in (".docx", ".odt"):
        return "docx"
    if suffix == ".pdf":
        return "pdf"
    return "image"


def _load_transkription(fi: "FileInfo", paths: "gf.ProjectPaths") -> str:
    """Liest die Transkription für Vision-Dateien — erst aus fi.analysis, dann direkt vom JSON."""
    transkription = (fi.analysis or {}).get("transkription", "").strip()
    if transkription:
        return transkription
    # Fallback: direkt aus der JSON-Datei lesen (falls fi.analysis veraltet ist)
    analysis_path = paths.feedback_data_dir / (fi.path.stem + "_analysis.json")
    if analysis_path.exists():
        try:
            disk = json.loads(analysis_path.read_text(encoding="utf-8"))
            return disk.get("transkription", "").strip()
        except Exception:
            logging.debug("Transkription aus %s nicht lesbar", analysis_path, exc_info=True)
    return ""


def safe_id(prefix: str, raw: str) -> str:
    digest = hashlib.md5(raw.encode()).hexdigest()[:8]
    return f"{prefix}-{digest}"


# =============================================================================
# Logo constants for NATASCHA TUI v8
# =============================================================================

LOGO_FULL = r"""
███╗   ██╗  █████╗  ████████╗  █████╗  ███████╗  ██████╗ ██╗  ██╗  █████╗ 
████╗  ██║ ██╔══██╗ ╚══██╔══╝ ██╔══██╗ ██╔════╝ ██╔════╝ ██║  ██║ ██╔══██╗
██╔██╗ ██║ ███████║    ██║    ███████║ ███████╗ ██║      ███████║ ███████║
██║╚██╗██║ ██╔══██║    ██║    ██╔══██║ ╚════██║ ██║      ██╔══██║ ██╔══██║
██║ ╚████║ ██║  ██║    ██║    ██║  ██║ ███████║ ╚██████╗ ██║  ██║ ██║  ██║
╚═╝  ╚═══╝ ╚═╝  ╚═╝    ╚═╝    ╚═╝  ╚═╝ ╚══════╝  ╚═════╝ ╚═╝  ╚═╝ ╚═╝  ╚═╝
"""

LOGO_COMPACT = r"""
▒█▄░▒█ ░█▀▀█ ▀▀█▀▀ ░█▀▀█ ▒█▀▀▀█ ▒█▀▀█ ▒█░▒█ ░█▀▀█
▒█▒█▒█ ▒█▄▄█ ░░█░░ ▒█▄▄█ ░▀▀▀▄▄ ▒█░░░ ▒█▀▀█ ▒█▄▄█
▒█░░▀█ ▒█░▒█ ░░█░░ ▒█░▒█ ▒█▄▄▄█ ▒█▄▄█ ▒█░▒█ ▒█░▒█
"""

INIT_IMAGE_PATH = Path(__file__).resolve().parent / "assets" / "natascha-init.png"
_HEADER_BG = (0x05, 0x08, 0x10)


def _rgb_hex(color: tuple[int, int, int]) -> str:
    return "#{:02x}{:02x}{:02x}".format(*color)


def _blend_rgba(
    pixel: tuple[int, int, int] | tuple[int, int, int, int],
    background: tuple[int, int, int] = _HEADER_BG,
) -> tuple[int, int, int]:
    if len(pixel) == 3:
        return pixel

    r, g, b, alpha = pixel
    if alpha <= 0:
        return background
    if alpha >= 255:
        return (r, g, b)

    alpha_ratio = alpha / 255
    return tuple(
        round(channel * alpha_ratio + bg_channel * (1 - alpha_ratio))
        for channel, bg_channel in zip((r, g, b), background)
    )


@lru_cache(maxsize=4)
def render_init_image(width: int = 120) -> Text | None:
    """Render the configured init image as terminal-friendly half blocks."""
    try:
        from PIL import Image
    except ImportError:
        return None

    if not INIT_IMAGE_PATH.exists():
        return None

    char_width = 22 if width >= 140 else 18 if width >= 100 else 14

    try:
        with Image.open(INIT_IMAGE_PATH) as image:
            image = image.convert("RGBA")
            target_height = max(8, round(image.height * char_width / image.width))
            resampling = getattr(getattr(Image, "Resampling", Image), "LANCZOS")
            image = image.resize((char_width, target_height), resampling)
            pixels = list(image.getdata())
    except Exception:
        return None

    art = Text(no_wrap=True)
    for y in range(0, target_height, 2):
        for x in range(char_width):
            top = _blend_rgba(pixels[y * char_width + x])
            bottom = _HEADER_BG
            if y + 1 < target_height:
                bottom = _blend_rgba(pixels[(y + 1) * char_width + x])
            art.append("▀", style=f"{_rgb_hex(top)} on {_rgb_hex(bottom)}")
        art.append("\n")

    return art


def render_logo_gradient(width: int = 120) -> Text:
    """Render NATASCHA logo with vertical gradient coloring."""
    logo = LOGO_FULL if width >= 120 else LOGO_COMPACT
    lines = logo.strip("\n").split("\n")
    colors = ["#ff00ff", "#cc44ff", "#8888ff", "#00aaff", "#00ddff", "#00ffaa"]

    result = Text()
    for line, color in zip(lines, colors):
        result.append(line + "\n", style=f"bold {color}")
    return result


def render_brand_art(width: int = 120) -> Text:
    """Render the ASCII gradient logo (PNG is a small icon, not suitable for header)."""
    return render_logo_gradient(width)


def render_acronym() -> Text:
    """Render NATASCHA acronym with highlighted initial letters."""
    parts = [
        ("N", "ormbasierte"),
        ("A", "nalyse von"),
        ("T", "exten"),
        ("A", "utomatisierte"),
        ("S", "chularbeits-"),
        ("C", "orrection mit"),
        ("H", "ilfe-"),
        ("A", "gents"),
    ]
    result = Text()
    for i, (letter, rest) in enumerate(parts):
        result.append(letter, style="bold #00ddff")
        result.append(rest, style="dim white")
        if i < len(parts) - 1:
            result.append("  ", style="dim white")
    return result


class HelpScreen(ModalScreen[None]):
    BINDINGS = [("escape", "dismiss", "Abbrechen")]

    def compose(self) -> ComposeResult:
        with Container(classes="help-container"):
            yield Label("HILFE – NATASCHA Dashboard", classes="panel-title")
            yield Static(
                "1. NAVIGATION\n"
                "  ↑/↓        Dateiliste navigieren\n"
                "  Tab         Panel-Fokus wechseln\n"
                "  /           Suche in Dateiliste\n"
                "  Esc         Suche beenden / Dialog schliessen\n\n"
                "2. MARKIERUNG & STAPELVERARBEITUNG\n"
                "  Space       Datei markieren / weiter (Batch)\n"
                "  Ctrl+A      Alle Dateien markieren\n"
                "  Ctrl+D      Alle Dateien demarkieren\n"
                "  Shift+A     Analyse fuer alle markierten Dateien\n"
                "  Shift+D     DOCX fuer alle markierten Dateien\n\n"
                "3. ANALYSE & AUSGABE\n"
                "  a           Analyse starten (API)\n"
                "  d           DOCX generieren\n"
                "  r           Review-Dialog öffnen\n"
                "  e           Zuordnung bearbeiten (Fach/Stufe/Textsorte/Rubrik)\n"
                "  g           Erwartungshorizont generieren (LLM)\n\n"
                "4. VORSCHAU\n"
                "  1/2/3       Vorschau auf Text / Bewertung / Rubrik\n"
                "  Enter       Naechster Vorschau-Tab\n"
                "  Del         Datei entfernen\n\n"
                "5. SCHÜLER & KLASSE\n"
                "  Shift+S     Schüler-Verwaltung\n"
                "  Shift+F     Klassen-Ansicht öffnen\n"
                "  F2          Retro-Import (bestehende Analysen)\n\n"
                "  👤 Schüler  Lernverlauf einer einzelnen Person über mehrere Arbeiten,\n"
                "              mit KI-Profil und Kriterien-Längsschnitt.\n\n"
                "  👥 Klasse   Alles Klassenbezogene in drei Reitern:\n"
                "    Feedback  Was als Nächstes im Unterricht zu tun ist, inkl. KI-Briefing.\n"
                "    Heatmap   Welche Fehlertypen häufen sich – mit konkreten Beispielen.\n"
                "    Statistik Notenverteilung und Lernfortschritt der ganzen Klasse.\n\n"
                "6. SONSTIGES\n"
                "  s           Einstellungen\n"
                "  ? / F1      Diese Hilfe\n"
                "  o           Sortierung ändern\n"
                "  q           Beenden"
            )
            with Horizontal(classes="help-buttons"):
                yield Button("📖 Anleitung öffnen", id="help-open-anleitung", variant="primary")
                yield Button("Schließen", id="help-close", variant="error")

    def on_button_pressed(self, event: Button.Pressed) -> None:
        if event.button.id == "help-close":
            self.dismiss(None)
        elif event.button.id == "help-open-anleitung":
            self._open_anleitung()

    def _open_anleitung(self) -> None:
        anleitung = Path(__file__).resolve().parent / "NATASCHA_Benutzeranleitung.docx"
        if anleitung.exists():
            nc.open_file(anleitung)
        else:
            self.app.notify(
                "Anleitung nicht gefunden: NATASCHA_Benutzeranleitung.docx", severity="warning"
            )


class ConfirmScreen(ModalScreen[bool]):
    BINDINGS = [("escape", "dismiss", "Abbrechen")]

    def __init__(self, message: str) -> None:
        super().__init__()
        self.message = message

    def compose(self) -> ComposeResult:
        with Container(classes="confirm-container"):
            yield Label(self.message)
            with Horizontal():
                yield Button("Ja", variant="success", id="confirm-yes")
                yield Button("Nein", variant="error", id="confirm-no")

    def on_button_pressed(self, event: Button.Pressed) -> None:
        self.dismiss(event.button.id == "confirm-yes")


class SettingsScreen(ModalScreen[None]):
    BINDINGS = [("escape", "dismiss", "Abbrechen")]

    _PROVIDERS = [
        ("DeepSeek", "deepseek"),
        ("Mistral", "mistral"),
        ("Qwen / DashScope", "qwen"),
        ("Kimi / Moonshot", "kimi"),
        ("OpenAI", "openai"),
        ("Anthropic / Claude", "anthropic"),
        ("Ollama (lokal)", "ollama"),
    ]
    _MODELS_BY_PROVIDER: dict[str, list[tuple[str, str]]] = {
        "deepseek": [
            ("deepseek-chat / V3 (kein Vision, stabil, günstig)", "deepseek-chat"),
            ("deepseek-reasoner / R1 (kein Vision, Reasoning)", "deepseek-reasoner"),
        ],
        "mistral": [
            ("mistral-small-latest (empfohlen, Text)", "mistral-small-latest"),
            ("mistral-large-latest (stärker, Text)", "mistral-large-latest"),
        ],
        "qwen": [
            ("qwen-plus (stabil, empfohlen)", "qwen-plus"),
            ("qwen3.6-plus (aktuell, wenn freigeschaltet)", "qwen3.6-plus"),
            ("qwen3-max (hohe Qualität, wenn freigeschaltet)", "qwen3-max"),
            ("qwen-turbo (schnell und günstig)", "qwen-turbo"),
        ],
        "kimi": [
            ("moonshot-v1-8k (kein Vision)", "moonshot-v1-8k"),
            ("moonshot-v1-32k (kein Vision)", "moonshot-v1-32k"),
        ],
        "openai": [
            ("gpt-4o (Vision: Bilder, stabil)", "gpt-4o"),
            ("gpt-4o-mini (Vision: Bilder, günstig)", "gpt-4o-mini"),
            ("gpt-4.1 (Vision: Bilder, falls Account freigeschaltet)", "gpt-4.1"),
            ("gpt-4.1-mini (Vision: Bilder, falls Account freigeschaltet)", "gpt-4.1-mini"),
            ("gpt-5.2 (experimentell, Account abhängig)", "gpt-5.2"),
            ("gpt-5.1 (experimentell, Account abhängig)", "gpt-5.1"),
            ("gpt-5 (experimentell, Account abhängig)", "gpt-5"),
        ],
        "anthropic": [
            ("claude-sonnet-4-6 (Vision: Bilder+PDF)", "claude-sonnet-4-6"),
            ("claude-opus-4-6 (Vision: Bilder+PDF)", "claude-opus-4-6"),
            ("claude-haiku-4-5-20251001 (Vision: Bilder+PDF)", "claude-haiku-4-5-20251001"),
        ],
        "ollama": [
            ("qwen3.5:27b (27B, empfohlen, kein Vision)", "qwen3.5:27b"),
            ("qwen3.5:35b (35B, langsam, kein Vision)", "qwen3.5:35b"),
            ("qwen3-vl:8b (8B, Vision-Modell*)", "qwen3-vl:8b"),
            ("qwen2.5-coder:32b (kein Vision)", "qwen2.5-coder:32b"),
        ],
    }

    def __init__(self, config: dict[str, Any]) -> None:
        super().__init__()
        self.config = config

    def compose(self) -> ComposeResult:
        defaults = self.config.get("defaults", {})
        api_cfg = self.config.get("api", {})
        current_provider = api_cfg.get("provider", "openai")
        current_model = api_cfg.get("model", "gpt-4o-mini")
        current_fach = defaults.get("fach", "Deutsch")
        current_schulstufe = defaults.get("schulstufe", "Oberstufe")

        provider_values = [v for _, v in self._PROVIDERS]
        if current_provider not in provider_values:
            current_provider = provider_values[0]

        model_opts = self._MODELS_BY_PROVIDER.get(current_provider, [("", "")])
        model_values = [v for _, v in model_opts]
        if current_model not in model_values:
            current_model = model_values[0] if model_values else ""

        # Key-Status für aktuellen Provider ermitteln
        key_env_map = {
            "anthropic": "ANTHROPIC_API_KEY",
            "deepseek": "DEEPSEEK_API_KEY",
            "mistral": "MISTRAL_API_KEY",
            "qwen": "QWEN_API_KEY",
            "kimi": "KIMI_API_KEY",
            "openai": "OPENAI_API_KEY",
            "ollama": None,
        }
        env_key_name = key_env_map.get(current_provider)
        if current_provider == "qwen":
            api_key_val = os.environ.get("QWEN_API_KEY", "") or os.environ.get(
                "DASHSCOPE_API_KEY", ""
            )
        else:
            api_key_val = os.environ.get(env_key_name, "") if env_key_name else "n/a (lokal)"
        if env_key_name is None:
            key_display = "kein Key nötig"
            has_key = "✓"
        elif len(api_key_val) > 12:
            key_display = f"{api_key_val[:6]}...{api_key_val[-3:]}"
            has_key = "✓"
        else:
            key_display = "(nicht gesetzt)"
            has_key = "✗"

        with VerticalScroll(classes="settings-container"):
            yield Label("EINSTELLUNGEN", classes="panel-title")

            with Container(classes="settings-section"):
                yield Label("Pfade", classes="settings-section-title")
                yield Static(f"  Input:   {nc.resolve_path(self.config, 'input')}")
                yield Static(f"  Output:  {nc.resolve_path(self.config, 'output')}")
                yield Static(f"  Rubrics: {nc.resolve_path(self.config, 'rubrics')}")

            with Container(classes="settings-section"):
                yield Label("API-Provider", classes="settings-section-title")
                yield Static(f"  Key:  {has_key} {key_display}")
                yield Label("  Provider:")
                yield Select(
                    options=self._PROVIDERS,
                    value=current_provider,
                    id="settings-provider",
                )
                yield Label("  Modell:")
                yield Select(
                    options=model_opts,
                    value=current_model,
                    id="settings-model",
                )

            with Container(classes="settings-section"):
                yield Label("CLI-Agents", classes="settings-section-title")
                availability = nc.check_agent_availability(self.config)
                for name, avail in availability.items():
                    sym = "✓" if avail else "✗"
                    yield Static(f"  {sym}  {name}")

            with Container(classes="settings-section"):
                yield Label("Standardwerte (aenderbar)", classes="settings-section-title")
                yield Label("  Standard-Fach:")
                yield Select(
                    options=[("Deutsch", "Deutsch"), ("Englisch", "Englisch")],
                    value=current_fach,
                    id="settings-fach",
                )
                yield Label("  Standard-Schulstufe:")
                yield Select(
                    options=[("Oberstufe", "Oberstufe"), ("Unterstufe", "Unterstufe")],
                    value=current_schulstufe,
                    id="settings-schulstufe",
                )

            with Horizontal():
                yield Button("Speichern", variant="success", id="save-settings-btn")
                yield Button("Abbrechen", variant="error", id="close-settings-btn")

    def on_select_changed(self, event: Select.Changed) -> None:
        """Wenn Provider geändert wird, Model-Optionen aktualisieren."""
        if event.select.id == "settings-provider" and event.value is not Select.BLANK:
            new_provider = str(event.value)
            model_opts = self._MODELS_BY_PROVIDER.get(new_provider, [("", "")])
            model_sel = self.query_one("#settings-model", Select)
            model_sel.set_options(model_opts)
            # Gespeichertes Modell wiederherstellen falls es zu diesem Provider passt
            saved_model = self.config.get("api", {}).get("model", "")
            model_values = [v for _, v in model_opts]
            if saved_model in model_values:
                model_sel.value = saved_model

    def on_button_pressed(self, event: Button.Pressed) -> None:
        if event.button.id == "save-settings-btn":
            self._save_settings()
        elif event.button.id == "close-settings-btn":
            self.dismiss(None)

    @staticmethod
    def _sel_value(sel: Select, fallback: str) -> str:
        """Liest den String-Wert eines Select — gibt fallback zurück wenn BLANK oder None."""
        v = sel.value
        if v is Select.BLANK or v is None:
            return fallback
        s = str(v)
        # Textual-interne Sentinel-Strings abfangen
        if s.startswith("_") or "NoSelection" in s or "NULL" in s or "BLANK" in s:
            return fallback
        return s

    def _save_settings(self) -> None:
        provider_sel = self.query_one("#settings-provider", Select)
        model_sel = self.query_one("#settings-model", Select)
        fach_sel = self.query_one("#settings-fach", Select)
        stufe_sel = self.query_one("#settings-schulstufe", Select)

        current_api = self.config.get("api", {})
        new_provider = self._sel_value(provider_sel, current_api.get("provider", "ollama"))
        # Modell-Fallback: erstes Modell des gewählten Providers
        default_model = self._MODELS_BY_PROVIDER.get(new_provider, [("", "gpt-4o-mini")])[0][1]
        new_model = self._sel_value(model_sel, current_api.get("model", default_model))
        new_fach = self._sel_value(fach_sel, "Deutsch")
        new_stufe = self._sel_value(stufe_sel, "Oberstufe")

        try:
            nc.save_settings(new_fach, new_stufe, new_provider, new_model)
            # Update in-memory config so changes take effect immediately
            self.config.setdefault("defaults", {})["fach"] = new_fach
            self.config.setdefault("defaults", {})["schulstufe"] = new_stufe
            self.config.setdefault("api", {})["provider"] = new_provider
            self.config.setdefault("api", {})["model"] = new_model
            self.app.notify("Einstellungen gespeichert – ab sofort aktiv.")
            self.dismiss(None)
        except Exception as e:
            self.app.notify(f"Fehler beim Speichern: {e}", severity="error")


class ReviewScreen(ModalScreen[None]):
    BINDINGS = [
        ("escape", "dismiss", "Abbrechen"),
        ("e", "edit_json", "JSON editieren"),
        ("d", "generate_docx", "DOCX"),
        ("p", "generate_pdf", "PDF"),
    ]

    def __init__(self, file_info: FileInfo, config: dict[str, Any]) -> None:
        super().__init__()
        self.file_info = file_info
        self.config = config

    def compose(self) -> ComposeResult:
        with VerticalScroll(classes="review-container"):
            fname = self.file_info.path.name
            yield Label(f"Review: {fname}", classes="panel-title")

            if not self.file_info.analysis:
                yield Static("Keine Analysedaten vorhanden.")
                return

            data = self.file_info.analysis
            note_data = data.get("notenempfehlung")
            notendetail = data.get("notendetail")
            if note_data and self.file_info.bewertungsmodus != "unbenotet":
                note = note_data.get("note", "?")
                bez = note_data.get("bezeichnung", "?")
                schnitt = note_data.get("durchschnitt", "?")
                yield Static(
                    _note_rich_text(note, f"Note: {note} – {bez}    Durchschnitt: {schnitt}"),
                    classes="review-note",
                )
                if notendetail:
                    k1_note = notendetail.get("k1_note")
                    k1_schnitt = notendetail.get("k1_schnitt")
                    k3_note = notendetail.get("k3_note")
                    k3_schnitt = notendetail.get("k3_schnitt")
                    detail_parts = []
                    if k1_note is not None:
                        k1_s = f" [Stufe {k1_schnitt:.1f}]" if k1_schnitt is not None else ""
                        detail_parts.append(f"K1 (Inhalt + Textstruktur): Note {k1_note}{k1_s}")
                    if k3_note is not None:
                        k3_s = f" [Stufe {k3_schnitt:.1f}]" if k3_schnitt is not None else ""
                        detail_parts.append(f"K3/1 (Stil + Sprachnormen): Note {k3_note}{k3_s}")
                    if detail_parts:
                        yield Static("  " + "  |  ".join(detail_parts), classes="review-note")
                    sonderregel = notendetail.get("sonderregel")
                    if sonderregel:
                        yield Static(
                            f"  ⚠ Sonderregel: {sonderregel}",
                            classes="review-note",
                        )
                llm_nd = data.get("notenempfehlung_llm")
                if llm_nd and str(llm_nd.get("note")) != str(note):
                    yield Static(
                        f"⚠ LLM empfiehlt: {llm_nd.get('note')} – {llm_nd.get('bezeichnung', '?')} (Abweichung)",
                        classes="review-note",
                    )
            else:
                yield Static("Hausaufgabe – keine Benotung", classes="review-note")

            bewertung = data.get("bewertung", {})
            for key, crit in bewertung.items():
                if not isinstance(crit, dict):
                    continue
                punkte = crit.get("punkte", 0)
                max_pts = 5
                filled = int(punkte) if isinstance(punkte, (int, float)) else 0
                bar = "●" * filled + "○" * (max_pts - filled)

                with Container(classes="criterion-box"):
                    yield Label(
                        f"{key.replace('_', ' ').title()} – {punkte} Punkte ({bar})",
                        classes="criterion-header",
                    )
                    with Vertical(classes="criterion-body"):
                        yield Static(f"Stufe: {crit.get('stufe', '?')}")

                        staerken = crit.get("staerken", [])
                        if staerken:
                            lines = "\n".join(f"  + {s}" for s in staerken)
                            yield Static(f"Staerken:\n{lines}", classes="strength")

                        schwaechen = crit.get("schwaechen", [])
                        if schwaechen:
                            lines = "\n".join(f"  - {s}" for s in schwaechen)
                            yield Static(f"Schwaechen:\n{lines}", classes="weakness")

                        vorschlaege = crit.get("vorschlaege", [])
                        if vorschlaege:
                            lines = "\n".join(f"  > {s}" for s in vorschlaege)
                            yield Static(f"Vorschlaege:\n{lines}", classes="suggestion")

            with Horizontal(classes="stats-buttons"):
                yield Button("📄 DOCX erstellen", id="review-docx", variant="primary")
                yield Button("📄 Als PDF", id="review-pdf", variant="default")
                yield Button("Abbrechen", id="review-close", variant="error")

    def action_edit_json(self) -> None:
        if not self.file_info.analysis:
            return
        paths = nc.build_project_paths(self.config)
        paths.feedback_data_dir.mkdir(parents=True, exist_ok=True)
        tmp_path = paths.feedback_data_dir / (self.file_info.path.stem + "_edit.json")
        tmp_path.write_text(
            json.dumps(self.file_info.analysis, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        editor = os.environ.get("EDITOR", "nano")
        self.app.suspend()
        subprocess.run([editor, str(tmp_path)])
        try:
            edited = json.loads(tmp_path.read_text(encoding="utf-8"))
            self.file_info.analysis = edited
            self._save_analysis()
        except json.JSONDecodeError:
            pass

    def action_generate_docx(self) -> None:
        if not self.file_info.analysis:
            return
        self._generate_single_docx()

    def action_generate_pdf(self) -> None:
        if not self.file_info.analysis:
            return
        self._generate_single_pdf()

    def on_button_pressed(self, event: Button.Pressed) -> None:
        if event.button.id == "review-docx":
            self.action_generate_docx()
        elif event.button.id == "review-pdf":
            self.action_generate_pdf()
        elif event.button.id == "review-close":
            self.dismiss(None)

    def _save_analysis(self) -> None:
        if not self.file_info.analysis:
            return
        paths = nc.build_project_paths(self.config)
        paths.feedback_data_dir.mkdir(parents=True, exist_ok=True)
        out_name = self.file_info.path.stem + "_analysis.json"
        analysis_path = paths.feedback_data_dir / out_name
        nc.archive_existing_analysis(analysis_path)
        analysis_path.write_text(
            json.dumps(self.file_info.analysis, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )

    def _get_original_text_and_note(self) -> tuple[str | None, str | None]:
        """Gibt (original_text, source_note) zurück — abhängig vom Dateityp."""
        fi = self.file_info
        if fi.file_type == "docx" and fi.path.exists():
            if fi.path.suffix.lower() == ".odt":
                return nc.read_odt_text(fi.path), None
            return nc.read_docx_text(fi.path), None
        paths = nc.build_project_paths(self.config)
        transkription = _load_transkription(fi, paths)
        if transkription:
            return transkription, None
        return None, f"Analyse basiert auf eingescannter Datei: {fi.path.name}"

    def _generate_single_docx(self) -> None:
        if not self.file_info.analysis:
            return
        paths = nc.build_project_paths(self.config)
        paths.output_dir.mkdir(parents=True, exist_ok=True)
        try:
            feedback = gf.parse_feedback_data(self.file_info.analysis)
            out_name = gf.output_filename(feedback.datei)
            out_path = paths.output_dir / out_name
            original_text, source_note = self._get_original_text_and_note()
            _klasse = nc.active_klasse(self.config)
            _aufgabe = nc.active_aufgabe(self.config, _klasse) if _klasse else None
            _auf_label = (
                nc.get_aufgabe_cfg(self.config, _klasse or "", _aufgabe or "").get("label", "")
                if _aufgabe
                else ""
            )
            doc = gf.build_feedback_document(
                feedback, config=self.config, original_text=original_text,
                source_note=source_note, aufgabe_label=_auf_label,
                bewertungsmodus=self.file_info.bewertungsmodus,
            )
            doc.save(str(out_path))
            self.file_info.status = FileStatus.DONE
            self.app.notify(f"DOCX gespeichert: {out_path.name}", severity="information")
            nc.open_file(out_path)
        except Exception as e:
            nc.log_tui_error(paths, f"{self.file_info.path.name}: {e}")
            self.file_info.status = FileStatus.ERROR
            self.app.notify(f"Fehler beim DOCX-Erstellen: {e}", severity="error")

    def _generate_single_pdf(self) -> None:
        if not self.file_info.analysis:
            return
        paths = nc.build_project_paths(self.config)
        paths.output_dir.mkdir(parents=True, exist_ok=True)
        try:
            feedback = gf.parse_feedback_data(self.file_info.analysis)
            out_name = gf.output_filename(feedback.datei)
            docx_path = paths.output_dir / out_name
            original_text, source_note = self._get_original_text_and_note()
            _klasse_p = nc.active_klasse(self.config)
            _aufgabe_p = nc.active_aufgabe(self.config, _klasse_p) if _klasse_p else None
            _auf_label_p = (
                nc.get_aufgabe_cfg(self.config, _klasse_p or "", _aufgabe_p or "").get("label", "")
                if _aufgabe_p
                else ""
            )
            doc = gf.build_feedback_document(
                feedback, config=self.config, original_text=original_text,
                source_note=source_note, aufgabe_label=_auf_label_p,
                bewertungsmodus=self.file_info.bewertungsmodus,
            )
            doc.save(str(docx_path))
            self.file_info.status = FileStatus.DONE

            pdf_path = nc.docx_to_pdf(docx_path)
            if pdf_path and pdf_path.exists():
                self.app.notify(f"PDF gespeichert: {pdf_path.name}", severity="information")
                nc.open_file(pdf_path)
            else:
                self.app.notify(
                    "PDF-Konvertierung nicht verfügbar — LibreOffice installieren.",
                    severity="warning",
                )
                nc.open_file(docx_path)
        except Exception as e:
            nc.log_tui_error(paths, f"{self.file_info.path.name}: {e}")
            self.app.notify(f"Fehler beim PDF-Erstellen: {e}", severity="error")


class ProgressScreen(ModalScreen[None]):
    BINDINGS = [("escape", "cancel_analysis", "Abbrechen")]

    def __init__(self) -> None:
        super().__init__()
        self._cancelled = False

    def compose(self) -> ComposeResult:
        with Container(classes="progress-container"):
            yield Label("Batch-Analyse", classes="panel-title")
            yield Static("", id="progress-summary")
            yield ProgressBar(total=1, show_eta=False, id="batch-progress")
            yield Static("", id="progress-current")
            yield Static("", id="progress-queue")
            yield Static("Esc: Batch abbrechen", id="progress-help")

    def update_progress(
        self,
        current_file: str,
        status_text: str,
        queue: list[str],
        done: int,
        total: int,
    ) -> None:
        remaining = max(total - done, 0)
        self.query_one("#progress-summary", Static).update(
            f"Gesamt: {done}/{total} erledigt | {remaining} offen"
        )
        self.query_one("#batch-progress", ProgressBar).update(
            total=max(total, 1),
            progress=min(done, total),
        )
        self.query_one("#progress-current", Static).update(
            f"Aktuell: {current_file}\nStatus:  {status_text}"
        )
        if queue:
            shown = queue[:8]
            q_text = "Warteschlange:\n" + "\n".join(f"  - {f}" for f in shown)
            if len(queue) > len(shown):
                q_text += f"\n  ... {len(queue) - len(shown)} weitere"
            self.query_one("#progress-queue", Static).update(q_text)
        else:
            self.query_one("#progress-queue", Static).update("Warteschlange: leer")

    def action_cancel_analysis(self) -> None:
        self._cancelled = True
        # Signal the worker thread via the app-level event
        try:
            self.app._cancel_event.set()
        except AttributeError:
            pass
        self.dismiss(None)


class EditAssignmentScreen(ModalScreen[bool]):
    BINDINGS = [("escape", "dismiss", "Abbrechen")]

    _ready: bool = False  # guards against cascade resets during initial mount

    _TEXTSORTEN_DEUTSCH = [
        "Erörterung",
        "Kommentar",
        "Leserbrief",
        "Textanalyse",
        "Textinterpretation",
        "Zusammenfassung",
        "Offener Brief",
        "Meinungsrede",
        "Empfehlung/Rezension",
    ]
    _TEXTSORTEN_ENGLISCH = [
        "Article",
        "Report",
        "Essay",
        "Email/Letter",
        "Review",
        "Blog Post",
        "Story",
    ]

    def __init__(
        self,
        file_info: FileInfo,
        config: dict[str, Any],
        seq_idx: int = 1,
        seq_total: int = 1,
    ) -> None:
        super().__init__()
        self.file_info = file_info
        self.config = config
        self._seq_idx = seq_idx
        self._seq_total = seq_total
        self._current_fach = file_info.fach or "Deutsch"
        self._current_schulstufe = file_info.schulstufe or "Oberstufe"

    def _textsorte_options(self, fach: str) -> list[tuple[str, str]]:
        lst = self._TEXTSORTEN_ENGLISCH if fach == "Englisch" else self._TEXTSORTEN_DEUTSCH
        return [(t, t) for t in lst]

    def _rubric_options(self) -> list[tuple[str, str]]:
        opts = nc.rubric_options_for(
            self._current_fach, self._current_schulstufe, self.config,
            current_rubric=self.file_info.rubric,
        )
        return [(o, o) for o in opts] if opts else [("(keine Rubrik gefunden)", "")]

    def _erwartungshorizont_options(self) -> list[tuple[str, str]]:
        rubrics_dir = nc.resolve_path(self.config, "rubrics")
        files = sorted(
            f.name for f in rubrics_dir.iterdir()
            if f.is_file() and f.name.startswith("erwartungshorizont_") and f.suffix == ".md"
        ) if rubrics_dir.exists() else []
        return [("(keiner)", "")] + [(f, f) for f in files]

    def compose(self) -> ComposeResult:
        ts_opts = self._textsorte_options(self._current_fach)
        ts_values = [v for _, v in ts_opts]
        current_ts = (
            self.file_info.textsorte
            if self.file_info.textsorte in ts_values
            else (ts_values[0] if ts_values else "")
        )

        rubric_opts = self._rubric_options()
        rubric_values = [v for _, v in rubric_opts]
        current_rubric = (
            self.file_info.rubric
            if self.file_info.rubric in rubric_values
            else (rubric_values[0] if rubric_values else "")
        )

        eh_opts = self._erwartungshorizont_options()
        eh_values = [v for _, v in eh_opts]
        current_eh = (
            self.file_info.erwartungshorizont
            if self.file_info.erwartungshorizont in eh_values
            else ""
        )

        with VerticalScroll(classes="settings-container"):
            seq_info = f" [{self._seq_idx}/{self._seq_total}]" if self._seq_total > 1 else ""
            yield Label(f"Zuordnung{seq_info}: {self.file_info.path.name}", classes="panel-title")
            yield Label("Fach:")
            yield Select(
                options=[("Deutsch", "Deutsch"), ("Englisch", "Englisch")],
                value=self._current_fach,
                id="edit-fach",
            )
            yield Label("Schulstufe:")
            yield Select(
                options=[("Oberstufe", "Oberstufe"), ("Unterstufe", "Unterstufe")],
                value=self._current_schulstufe,
                id="edit-schulstufe",
            )
            yield Label("Textsorte:")
            yield Select(
                options=ts_opts,
                value=current_ts,
                id="edit-textsorte",
            )
            yield Label("Rubrik:")
            yield Select(
                options=rubric_opts,
                value=current_rubric,
                id="edit-rubric",
            )
            yield Label("Erwartungshorizont (optional):")
            yield Select(
                options=eh_opts,
                value=current_eh if current_eh else "",
                id="edit-erwartungshorizont",
            )
            yield Label("Bewertungsmodus:")
            yield Select(
                options=[
                    ("Schularbeit (benotet)", "benotet"),
                    ("Hausaufgabe (unbenotet)", "unbenotet"),
                ],
                value=self.file_info.bewertungsmodus,
                id="edit-modus",
            )
            yield Label("Schüler/in (optional):")
            schueler_opts = [("(nicht zugeordnet)", "")]
            if ndb is not None:
                try:
                    _klasse = nc.active_klasse(self.config) or ""
                    if _klasse:
                        db_path = ndb.get_db_path(self.config)
                        ndb.init_db(db_path)
                        for s in ndb.get_schueler_by_klasse(db_path, _klasse):
                            name = f"{s['vorname']} {s['nachname']}".strip()
                            schueler_opts.append((name, name))
                except Exception:
                    pass
            current_schueler = self.file_info.schueler or ""
            if current_schueler and not any(v == current_schueler for _, v in schueler_opts):
                schueler_opts.append((current_schueler, current_schueler))
            yield Select(
                options=schueler_opts,
                value=current_schueler,
                id="edit-schueler",
            )
            with Horizontal():
                yield Button("Speichern", variant="success", id="edit-save")
                yield Button("Abbrechen", variant="error", id="edit-cancel")

    def on_mount(self) -> None:
        self._ready = True

    def on_select_changed(self, event: Select.Changed) -> None:
        """Kaskadiert Textsorte- und Rubrik-Optionen wenn Fach oder Schulstufe geaendert wird."""
        if not self._ready:
            return
        if event.select.id == "edit-fach" and event.value is not Select.BLANK:
            self._current_fach = str(event.value)
            # Textsorte-Optionen aktualisieren, aktuellen Wert wenn möglich beibehalten
            new_ts_opts = self._textsorte_options(self._current_fach)
            ts_values = [v for _, v in new_ts_opts]
            ts_select = self.query_one("#edit-textsorte", Select)
            current_ts = (
                str(ts_select.value)
                if ts_select.value is not Select.BLANK
                else self.file_info.textsorte
            )
            ts_select.set_options(new_ts_opts)
            if current_ts in ts_values:
                ts_select.value = current_ts
            elif ts_values:
                ts_select.value = ts_values[0]
            # Rubrik aktualisieren
            self._update_rubric_select()

        elif event.select.id == "edit-schulstufe" and event.value is not Select.BLANK:
            self._current_schulstufe = str(event.value)
            self._update_rubric_select()

    def _update_rubric_select(self) -> None:
        new_rubric_opts = self._rubric_options()
        rubric_select = self.query_one("#edit-rubric", Select)
        current_val = (
            str(rubric_select.value)
            if rubric_select.value is not Select.BLANK
            else self.file_info.rubric
        )
        rubric_select.set_options(new_rubric_opts)
        rubric_values = [v for _, v in new_rubric_opts]
        if current_val in rubric_values:
            rubric_select.value = current_val
        elif rubric_values:
            rubric_select.value = rubric_values[0]

    def on_button_pressed(self, event: Button.Pressed) -> None:
        if event.button.id == "edit-save":
            fach_sel = self.query_one("#edit-fach", Select)
            stufe_sel = self.query_one("#edit-schulstufe", Select)
            ts_sel = self.query_one("#edit-textsorte", Select)
            rubric_sel = self.query_one("#edit-rubric", Select)
            eh_sel = self.query_one("#edit-erwartungshorizont", Select)
            modus_sel = self.query_one("#edit-modus", Select)
            schueler_sel = self.query_one("#edit-schueler", Select)

            if fach_sel.value is not Select.BLANK:
                self.file_info.fach = str(fach_sel.value)
            if stufe_sel.value is not Select.BLANK:
                self.file_info.schulstufe = str(stufe_sel.value)
            if ts_sel.value is not Select.BLANK:
                self.file_info.textsorte = str(ts_sel.value)
            if rubric_sel.value is not Select.BLANK:
                self.file_info.rubric = str(rubric_sel.value)
            self.file_info.erwartungshorizont = (
                str(eh_sel.value) if eh_sel.value is not Select.BLANK else ""
            )
            if modus_sel.value is not Select.BLANK:
                self.file_info.bewertungsmodus = str(modus_sel.value)  # type: ignore[assignment]
            self.file_info.schueler = (
                str(schueler_sel.value) if schueler_sel.value is not Select.BLANK else ""
            )
            self.dismiss(True)
        else:
            self.dismiss(False)


class NataschaHeader(Widget):
    """Enhanced header with init artwork, acronym tagline, and animated status."""

    DEFAULT_CSS = """
    NataschaHeader {
        dock: top;
        height: 12;
        background: #050810;
        padding: 1 2 0 2;
    }
    NataschaHeader #logo-container {
        width: 75%;
        align: left top;
        overflow-x: hidden;
    }
    NataschaHeader #status-container {
        width: 30%;
        align: right top;
        padding: 1 2;
    }
    NataschaHeader #acronym {
        margin-top: 0;
    }
    """

    _pulse_state = reactive(0)

    def __init__(self, *args: Any, **kwargs: Any) -> None:
        super().__init__(*args, **kwargs)
        self._terminal_width = 120
        self._api_status = "?"
        self._n_files = 0
        self._version = "0.4.0"

    def compose(self) -> ComposeResult:
        with Horizontal():
            with Vertical(id="logo-container"):
                yield Static(render_brand_art(self._terminal_width), id="logo")
                yield Static(render_acronym(), id="acronym")
            with Vertical(id="status-container"):
                yield Static(self._render_status(), id="status")

    def on_mount(self) -> None:
        """Start pulse animation for offline API indicator."""
        self.set_interval(0.8, self._pulse)

    def _pulse(self) -> None:
        """Toggle pulse state for animated status indicator."""
        self._pulse_state = (self._pulse_state + 1) % 2
        try:
            self.query_one("#status", Static).update(self._render_status())
        except Exception:
            logging.debug("Status-Widget-Update (Pulse) fehlgeschlagen", exc_info=True)

    def on_resize(self) -> None:
        """Handle terminal resize for responsive brand art."""
        if self.app.size:
            old_width = self._terminal_width
            self._terminal_width = self.app.size.width
            if old_width != self._terminal_width:
                try:
                    self.query_one("#logo", Static).update(render_brand_art(self._terminal_width))
                except Exception:
                    logging.debug("Logo-Update fehlgeschlagen", exc_info=True)

    def update_status(self, api_status: str, n_files: int, version: str) -> None:
        """Update header status information."""
        self._api_status = api_status
        self._n_files = n_files
        self._version = version
        try:
            self.query_one("#status", Static).update(self._render_status())
        except Exception:
            logging.debug("Status-Widget-Update fehlgeschlagen", exc_info=True)

    def _render_status(self) -> Text:
        """Render status text with animated offline indicator."""
        text = Text(justify="right")

        api_status_str = getattr(self, "_api_status", "?")
        n_files = getattr(self, "_n_files", 0)
        version = getattr(self, "_version", "0.4.0")

        if api_status_str == "✓":
            text.append("● API verbunden\n", style="bold #00ffaa")
        else:
            color = "#ff4466" if self._pulse_state else "#882233"
            text.append("● API offline\n", style=f"bold {color}")

        text.append(f"{n_files} Dateien\n", style="dim")
        text.append(f"v{version}", style="dim italic")
        return text


# Note-Beschriftungen für Statistiken
_NOTE_LABELS = {1: "Sehr gut", 2: "Gut", 3: "Befriedigend", 4: "Genügend", 5: "Nicht gen."}
_BAR_WIDTH = 20  # Breite des ASCII-Balkens




class RubrikEditorScreen(ModalScreen[bool]):
    """Modal zum Bearbeiten einer Rubrik-Datei (.md) aus rubrics/."""

    BINDINGS = [
        ("escape", "request_close", "Abbrechen"),
        ("ctrl+s", "save", "Speichern"),
    ]

    def __init__(self, rubric_filename: str, config: dict[str, Any]) -> None:
        super().__init__()
        self._rubric_filename = rubric_filename
        self._config = config
        self._original_text = ""
        self._dirty = False
        self._mode = "edit"

    def compose(self) -> ComposeResult:
        with VerticalScroll(classes="settings-container"):
            yield Label(f"Rubrik-Editor: {self._rubric_filename}", classes="panel-title")
            with Horizontal(classes="stats-buttons"):
                yield Button("Schreiben", id="re-tab-edit", variant="primary")
                yield Button("Vorschau", id="re-tab-preview", variant="default")
                yield Button("Speichern", id="re-save", variant="success")
                yield Button("📂 Extern öffnen", id="re-external", variant="default")
            yield TextArea(id="re-textarea")
            yield Markdown(id="re-preview", classes="review-container")
            with Horizontal(classes="stats-buttons"):
                yield Button("Abbrechen", id="re-close", variant="error")

    def on_mount(self) -> None:
        try:
            self._original_text = nc.load_rubric(self._rubric_filename, self._config)
        except FileNotFoundError:
            self._original_text = ""
        textarea = self.query_one("#re-textarea", TextArea)
        textarea.load_text(self._original_text)
        self._show_edit_tab()

    def _show_edit_tab(self) -> None:
        self._mode = "edit"
        self.query_one("#re-textarea", TextArea).display = True
        self.query_one("#re-preview", Markdown).display = False
        self.query_one("#re-tab-edit", Button).variant = "primary"
        self.query_one("#re-tab-preview", Button).variant = "default"

    def _show_preview_tab(self) -> None:
        self._mode = "preview"
        textarea = self.query_one("#re-textarea", TextArea)
        preview = self.query_one("#re-preview", Markdown)
        preview.update(textarea.text)
        textarea.display = False
        preview.display = True
        self.query_one("#re-tab-edit", Button).variant = "default"
        self.query_one("#re-tab-preview", Button).variant = "primary"

    def on_button_pressed(self, event: Button.Pressed) -> None:
        btn_id = event.button.id
        if btn_id == "re-tab-edit":
            self._show_edit_tab()
        elif btn_id == "re-tab-preview":
            self._show_preview_tab()
        elif btn_id == "re-save":
            self.action_save()
        elif btn_id == "re-close":
            self.action_request_close()
        elif btn_id == "re-external":
            self.action_open_external()

    def on_text_area_changed(self, event: TextArea.Changed) -> None:
        if event.text_area.id == "re-textarea":
            self._dirty = event.text_area.text != self._original_text

    def action_save(self) -> None:
        textarea = self.query_one("#re-textarea", TextArea)
        rubric_dir = nc.resolve_path(self._config, "rubrics")
        target = rubric_dir / self._rubric_filename
        try:
            target.write_text(textarea.text, encoding="utf-8")
            self._original_text = textarea.text
            self._dirty = False
            self.app.notify(
                f"Rubrik gespeichert: {self._rubric_filename}",
                severity="information",
            )
        except Exception as e:
            self.app.notify(f"Fehler beim Speichern: {e}", severity="error")

    def action_open_external(self) -> None:
        rubric_dir = nc.resolve_path(self._config, "rubrics")
        target = rubric_dir / self._rubric_filename
        if not target.exists():
            self.app.notify(f"Datei nicht gefunden: {target}", severity="error")
            return
        nc.open_file(target)
        self.app.notify(f"Geöffnet: {self._rubric_filename}", severity="information")

    def action_request_close(self) -> None:
        if self._dirty:

            def _handle(result: bool) -> None:
                if result:
                    self.dismiss(False)

            self.app.push_screen(
                ConfirmScreen("Ungespeicherte Aenderungen verwerfen?"),
                _handle,
            )
        else:
            self.dismiss(False)


class AddAufgabeScreen(ModalScreen[bool]):
    """Modal zum Anlegen einer neuen Aufgabe (Schularbeit) für eine Klasse."""

    BINDINGS = [("escape", "dismiss", "Abbrechen")]

    _TEXTSORTEN_DEUTSCH = [
        "Erörterung",
        "Kommentar",
        "Leserbrief",
        "Textanalyse",
        "Textinterpretation",
        "Zusammenfassung",
        "Offener Brief",
        "Meinungsrede",
        "Empfehlung/Rezension",
    ]
    _TEXTSORTEN_ENGLISCH = [
        "Article",
        "Report",
        "Essay",
        "Email/Letter",
        "Review",
        "Blog Post",
    ]

    def __init__(self, klasse: str, existing_slugs: list[str], config: dict[str, Any]) -> None:
        super().__init__()
        self._klasse = klasse
        self._existing = existing_slugs
        self._config = config
        self._current_fach = "Deutsch"
        self._current_schulstufe = "Oberstufe"

    def _textsorte_options(self, fach: str) -> list[tuple[str, str]]:
        lst = self._TEXTSORTEN_ENGLISCH if fach == "Englisch" else self._TEXTSORTEN_DEUTSCH
        return [(t, t) for t in lst]

    def _rubric_options(self) -> list[tuple[str, str]]:
        opts = nc.rubric_options_for(self._current_fach, self._current_schulstufe, self._config)
        return [(o, o) for o in opts] if opts else [("(keine Rubrik gefunden)", "")]

    def compose(self) -> ComposeResult:
        ts_opts = self._textsorte_options(self._current_fach)
        rubric_opts = self._rubric_options()
        with Container(classes="settings-container"):
            yield Label(f"Neue Aufgabe für Klasse {self._klasse}", classes="panel-title")
            yield Label("Bezeichnung (z.B. SA1 – Kommentar):")
            yield Input(placeholder="Label", id="auf-label")
            yield Label("Fach:")
            yield Select(
                options=[("Deutsch", "Deutsch"), ("Englisch", "Englisch")],
                value="Deutsch",
                id="auf-fach",
            )
            yield Label("Schulstufe:")
            yield Select(
                options=[("Oberstufe", "Oberstufe"), ("Unterstufe", "Unterstufe")],
                value="Oberstufe",
                id="auf-schulstufe",
            )
            yield Label("Textsorte:")
            yield Select(options=ts_opts, value=ts_opts[0][1], id="auf-textsorte")
            yield Label("Rubrik:")
            yield Select(options=rubric_opts, value=rubric_opts[0][1], id="auf-rubric")
            with Horizontal():
                yield Button("Anlegen", variant="success", id="auf-ok")
                yield Button("Abbrechen", variant="error", id="auf-cancel")

    def on_select_changed(self, event: Select.Changed) -> None:
        if event.select.id == "auf-fach" and event.value is not Select.BLANK:
            self._current_fach = str(event.value)
            ts_opts = self._textsorte_options(self._current_fach)
            self.query_one("#auf-textsorte", Select).set_options(ts_opts)
            self._refresh_rubric()
        elif event.select.id == "auf-schulstufe" and event.value is not Select.BLANK:
            self._current_schulstufe = str(event.value)
            self._refresh_rubric()

    def _refresh_rubric(self) -> None:
        opts = self._rubric_options()
        rubric_sel = self.query_one("#auf-rubric", Select)
        rubric_sel.set_options(opts)
        default = nc.default_rubric_for(self._current_fach, self._current_schulstufe, self._config)
        values = [v for _, v in opts]
        rubric_sel.value = default if default in values else (values[0] if values else "")

    def on_button_pressed(self, event: Button.Pressed) -> None:
        if event.button.id == "auf-cancel":
            self.dismiss(False)
        elif event.button.id == "auf-ok":
            self._do_save()

    def on_input_submitted(self, event: Input.Submitted) -> None:
        self._do_save()

    def _do_save(self) -> None:
        label = self.query_one("#auf-label", Input).value.strip()
        if not label:
            self.app.notify("Bitte eine Bezeichnung eingeben.", severity="warning")
            return
        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", label)
        if slug in self._existing:
            self.app.notify(f"Aufgabe '{slug}' existiert bereits.", severity="warning")
            return
        fach_sel = self.query_one("#auf-fach", Select)
        stufe_sel = self.query_one("#auf-schulstufe", Select)
        ts_sel = self.query_one("#auf-textsorte", Select)
        rubric_sel = self.query_one("#auf-rubric", Select)
        fach = str(fach_sel.value) if fach_sel.value is not Select.BLANK else "Deutsch"
        schulstufe = str(stufe_sel.value) if stufe_sel.value is not Select.BLANK else "Oberstufe"
        textsorte = str(ts_sel.value) if ts_sel.value is not Select.BLANK else ""
        rubric = str(rubric_sel.value) if rubric_sel.value is not Select.BLANK else ""
        try:
            nc.add_aufgabe_to_config(self._klasse, slug, label, fach, schulstufe, textsorte, rubric)
            self.dismiss(True)
        except Exception as e:
            self.app.notify(f"Fehler: {e}", severity="error")


class AddClassScreen(ModalScreen[bool]):
    """Modal zum Anlegen einer neuen Klasse."""

    BINDINGS = [("escape", "dismiss", "Abbrechen")]

    def __init__(self, existing_names: list[str]) -> None:
        super().__init__()
        self._existing = existing_names

    def compose(self) -> ComposeResult:
        with Container(classes="confirm-container"):
            yield Label("Neue Klasse anlegen", classes="panel-title")
            yield Label("Name der Klasse (z.B. 7A, 8B):")
            yield Input(placeholder="Klassenname", id="new-class-name")
            with Horizontal():
                yield Button("Anlegen", variant="success", id="add-class-ok")
                yield Button("Abbrechen", variant="error", id="add-class-cancel")
        yield Static("", id="add-class-error")

    def on_button_pressed(self, event: Button.Pressed) -> None:
        if event.button.id == "add-class-cancel":
            self.dismiss(False)
            return
        if event.button.id == "add-class-ok":
            self._do_save()

    def on_input_submitted(self, event: Input.Submitted) -> None:
        self._do_save()

    def _do_save(self) -> None:
        name_input = self.query_one("#new-class-name", Input)
        name = name_input.value.strip()
        if not name:
            self.app.notify("Bitte einen Klassennamen eingeben.", severity="warning")
            return
        if name in self._existing:
            self.app.notify(f"Klasse '{name}' existiert bereits.", severity="warning")
            return
        slug = re.sub(r"[^a-zA-Z0-9_\-]", "_", name).lower()
        input_rel = f"input/{slug}"
        output_rel = f"output/{slug}"
        try:
            nc.add_class_to_config(name, input_rel, output_rel)
            (nc.PROJECT_ROOT / input_rel).mkdir(parents=True, exist_ok=True)
            (nc.PROJECT_ROOT / output_rel).mkdir(parents=True, exist_ok=True)
            self.dismiss(True)
        except Exception as e:
            self.app.notify(f"Fehler: {e}", severity="error")


class AttachRubricScreen(ModalScreen[bool]):
    """Modal zum Anhängen einer Rubrik-Datei (.md) an eine Aufgabe."""

    BINDINGS = [("escape", "dismiss", "Abbrechen")]

    def __init__(
        self,
        klasse: str,
        aufgabe: str,
        aufgabe_label: str,
        config: dict[str, Any],
    ) -> None:
        super().__init__()
        self._klasse = klasse
        self._aufgabe = aufgabe
        self._aufgabe_label = aufgabe_label
        self._config = config

    def compose(self) -> ComposeResult:
        all_rubrics = nc.list_all_rubrics(self._config)
        rubric_opts: list[tuple[str, str]] = [(r, r) for r in all_rubrics]
        if not rubric_opts:
            rubric_opts = [("(keine Rubrik vorhanden)", "")]
        with Container(classes="settings-container"):
            yield Label(f"Rubrik für: {self._aufgabe_label}", classes="panel-title")
            yield Label("Pfad zur Rubrik-Datei (.md) – tippen oder einfügen:")
            yield Input(
                placeholder="/home/natascha/meine_rubrik.md",
                id="rubric-path-input",
            )
            yield Label("Oder vorhandene Rubrik aus rubrics/ wählen:")
            yield Select(
                options=rubric_opts,
                id="rubric-existing-select",
                prompt="Rubrik wählen…",
            )
            with Horizontal():
                yield Button("Speichern", variant="success", id="rubric-save")
                yield Button("Abbrechen", variant="error", id="rubric-cancel")

    def on_button_pressed(self, event: Button.Pressed) -> None:
        if event.button.id == "rubric-cancel":
            self.dismiss(False)
        elif event.button.id == "rubric-save":
            self._do_save()

    def on_input_submitted(self, _: Input.Submitted) -> None:
        self._do_save()

    def _do_save(self) -> None:
        path_str = self.query_one("#rubric-path-input", Input).value.strip()
        existing_sel = self.query_one("#rubric-existing-select", Select)

        if path_str:
            source = Path(path_str)
            if not source.exists():
                self.app.notify(f"Datei nicht gefunden: {path_str}", severity="error")
                return
            if source.suffix.lower() != ".md":
                self.app.notify("Nur .md-Dateien erlaubt.", severity="warning")
                return
            try:
                dest_name = nc.attach_rubric_to_aufgabe(self._klasse, self._aufgabe, source)
                self.app.notify(f"Rubrik gespeichert: {dest_name}", severity="information")
                self.dismiss(True)
            except Exception as e:
                self.app.notify(f"Fehler: {e}", severity="error")
        elif existing_sel.value is not Select.BLANK and str(existing_sel.value):
            rubric_name = str(existing_sel.value)
            try:
                nc.set_rubric_for_aufgabe(self._klasse, self._aufgabe, rubric_name)
                self.app.notify(f"Rubrik gesetzt: {rubric_name}", severity="information")
                self.dismiss(True)
            except Exception as e:
                self.app.notify(f"Fehler: {e}", severity="error")
        else:
            self.app.notify("Bitte Pfad eingeben oder Rubrik wählen.", severity="warning")


class ErwartungshorizontGeneratorScreen(ModalScreen[bool]):
    """Modal zum Generieren eines Erwartungshorizonts via LLM."""

    BINDINGS = [("escape", "dismiss", "Abbrechen")]

    _PROVIDERS = [
        ("Anthropic / Claude", "anthropic"),
        ("OpenAI", "openai"),
        ("DeepSeek", "deepseek"),
        ("Mistral", "mistral"),
        ("Qwen / DashScope", "qwen"),
        ("Kimi / Moonshot", "kimi"),
        ("Ollama (lokal)", "ollama"),
    ]

    def __init__(self, config: dict[str, Any], klasse: str, aufgabe: str) -> None:
        super().__init__()
        self.config = config
        self.klasse = klasse
        self.aufgabe = aufgabe
        self._result_text: str = ""
        self._cancel_event = threading.Event()

    def compose(self) -> ComposeResult:
        auf_cfg = nc.get_aufgabe_cfg(self.config, self.klasse, self.aufgabe)
        label = auf_cfg.get("label", self.aufgabe)
        ausgangstext = nc.detect_ausgangstext(self.config, self.klasse, self.aufgabe)
        at_status = f"✓ {ausgangstext.name}" if ausgangstext else "✗ Kein Ausgangstext in ausgangstext/"

        eh_cfg = self.config.get("erwartungshorizont", {})
        api_cfg = self.config.get("api", {})
        cur_prov = eh_cfg.get("default_provider", api_cfg.get("provider", "anthropic"))
        cur_model = eh_cfg.get("default_model", api_cfg.get("model", ""))
        safe_prov = _safe_select_value(cur_prov, self._PROVIDERS, "anthropic")

        with Container(classes="settings-container"):
            yield Label(f"Erwartungshorizont generieren: {label}", classes="panel-title")
            yield Static(f"Klasse: {self.klasse}  |  Aufgabe: {self.aufgabe}  |  Textsorte: {auf_cfg.get('textsorte', '—')}")
            yield Static(at_status, id="eh-at-status")
            yield Label("LLM-Provider:")
            yield Select(options=self._PROVIDERS, value=safe_prov, id="eh-provider")
            yield Label("Modell (leer = Config-Default):")
            yield Input(value=cur_model, placeholder="z. B. claude-opus-4-20250514", id="eh-model")
            yield Label("Vorschau:")
            with VerticalScroll(id="eh-preview-scroll"):
                yield Markdown("*(Noch nicht generiert)*", id="eh-preview")
            with Horizontal():
                yield Button("Generieren", variant="success", id="eh-generate")
                yield Button("Übernehmen", variant="primary", id="eh-accept", disabled=True)
                yield Button("Abbrechen", variant="error", id="eh-cancel")

    def on_button_pressed(self, event: Button.Pressed) -> None:
        if event.button.id == "eh-generate":
            self._cancel_event.clear()
            self._start_generation()
        elif event.button.id == "eh-accept" and self._result_text:
            self._save_and_accept()
        elif event.button.id == "eh-cancel":
            self._cancel_event.set()
            self.dismiss(False)

    @work(thread=True)
    def _start_generation(self) -> None:
        prov_sel = self.query_one("#eh-provider", Select)
        model_input = self.query_one("#eh-model", Input)
        provider = str(prov_sel.value) if prov_sel.value is not Select.BLANK else ""
        model = model_input.value.strip()

        self.call_from_thread(
            self.query_one("#eh-preview", Markdown).update, "*(Generierung läuft…)*"
        )
        self.call_from_thread(self.query_one("#eh-generate", Button).__setattr__, "disabled", True)

        try:
            result = nc.generate_erwartungshorizont(
                self.config, self.klasse, self.aufgabe,
                provider=provider, model=model,
                cancel_event=self._cancel_event,
            )
            self._result_text = result
            self.call_from_thread(self.query_one("#eh-preview", Markdown).update, result)
            self.call_from_thread(
                self.query_one("#eh-accept", Button).__setattr__, "disabled", False
            )
        except Exception as e:
            self.call_from_thread(
                self.query_one("#eh-preview", Markdown).update, f"**Fehler:** {e}"
            )
        finally:
            self.call_from_thread(
                self.query_one("#eh-generate", Button).__setattr__, "disabled", False
            )

    def _save_and_accept(self) -> None:
        slug = self.aufgabe.lower().replace(" ", "_")
        eh_filename = f"erwartungshorizont_{slug}.md"
        rubrics_dir = nc.resolve_path(self.config, "rubrics")
        rubrics_dir.mkdir(parents=True, exist_ok=True)
        (rubrics_dir / eh_filename).write_text(self._result_text, encoding="utf-8")
        try:
            nc.save_erwartungshorizont_to_config(self.klasse, self.aufgabe, eh_filename)
            self.app.notify(f"Erwartungshorizont gespeichert: {eh_filename}", severity="information")
        except Exception as e:
            self.app.notify(f"Config-Speichern fehlgeschlagen: {e}", severity="warning")
        self.dismiss(True)


class SchuelerVerwaltungScreen(ModalScreen[None]):
    """Modal-Screen zur Verwaltung der Schüler pro Klasse."""

    BINDINGS = [("escape", "dismiss", "Schließen")]

    def __init__(self, config: dict[str, Any]) -> None:
        super().__init__()
        self._config = config
        self._klasse = nc.active_klasse(config) or ""
        self._schueler_liste: list[dict[str, Any]] = []

    def compose(self) -> ComposeResult:
        with VerticalScroll(classes="schueler-container"):
            yield Label(f"Schüler-Verwaltung – {self._klasse}", classes="panel-title")
            klassen = nc.list_classes(self._config)
            if klassen:
                options = [(k, k) for k in klassen]
                active = self._klasse if self._klasse in klassen else None
                sel = Select(options, id="schueler-klassen-select", prompt="Klasse wählen…")
                if active:
                    sel.value = active
                yield sel
            yield DataTable(id="schueler-table")
            with Horizontal(classes="schueler-input-row"):
                yield Input(placeholder="Vorname", id="schueler-vorname")
                yield Input(placeholder="Nachname (optional)", id="schueler-nachname")
                yield Button("➕ Hinzufügen", id="schueler-add", variant="success")
            with Horizontal(classes="schueler-buttons"):
                yield Button("🗑 Löschen", id="schueler-del", variant="error", classes="btn-destructive")
                yield Button("📈 Verlauf", id="schueler-detail", variant="primary")
                yield Button("📁 CSV-Import", id="schueler-csv", variant="primary")
                yield Button("📁 CSV-Export Noten", id="schueler-export", variant="primary")
                yield Button("Schließen", id="schueler-close", variant="error")

    def on_mount(self) -> None:
        table = self.query_one("#schueler-table", DataTable)
        table.add_columns("ID", "Vorname", "Nachname", "Eintrag")
        self._refresh_table()

    def _refresh_table(self) -> None:
        table = self.query_one("#schueler-table", DataTable)
        table.clear()
        self._schueler_liste = []
        if ndb is not None and self._klasse:
            try:
                db_path = ndb.get_db_path(self._config)
                ndb.init_db(db_path)
                self._schueler_liste = ndb.get_schueler_by_klasse(db_path, self._klasse)
            except Exception as e:
                self.app.notify(f"DB-Fehler: {e}", severity="error")
        for s in self._schueler_liste:
            table.add_row(
                str(s["id"]),
                s["vorname"],
                s.get("nachname") or "",
                s.get("created_at", "")[:10],
            )

    def on_select_changed(self, event: Select.Changed) -> None:
        if event.select.id == "schueler-klassen-select":
            self._klasse = str(event.value) if event.value is not Select.BLANK else ""
            lbl = self.query_one(".panel-title", Label)
            lbl.update(f"Schüler-Verwaltung – {self._klasse or '(keine Klasse)'}")
            self._refresh_table()

    def on_button_pressed(self, event: Button.Pressed) -> None:
        btn_id = event.button.id
        if btn_id == "schueler-close":
            self.dismiss(None)
            return
        if ndb is None:
            self.app.notify("Datenbank-Modul nicht verfügbar.", severity="error")
            return
        db_path = ndb.get_db_path(self._config)
        try:
            ndb.init_db(db_path)
        except Exception as e:
            self.app.notify(f"DB-Fehler: {e}", severity="error")
            return

        if btn_id == "schueler-add":
            vn = self.query_one("#schueler-vorname", Input).value.strip()
            nn = self.query_one("#schueler-nachname", Input).value.strip()
            if not vn:
                self.app.notify("Vorname erforderlich.", severity="warning")
                return
            try:
                ndb.insert_schueler(db_path, self._klasse, vn, nn)
                self.app.notify(f"'{vn}' hinzugefuegt.", severity="information")
                self.query_one("#schueler-vorname", Input).value = ""
                self.query_one("#schueler-nachname", Input).value = ""
                self._refresh_table()
            except Exception as e:
                self.app.notify(f"Fehler: {e}", severity="error")

        elif btn_id == "schueler-del":
            table = self.query_one("#schueler-table", DataTable)
            if table.cursor_row is None:
                self.app.notify("Bitte Zeile auswählen.", severity="warning")
                return
            idx = table.cursor_row
            if 0 <= idx < len(self._schueler_liste):
                sid = self._schueler_liste[idx]["id"]
                name = self._schueler_liste[idx]["vorname"]

                def _do_delete(confirmed: bool) -> None:
                    if not confirmed:
                        return
                    try:
                        _db_path = ndb.get_db_path(self._config)
                        ndb.delete_schueler(_db_path, sid)
                        self.app.notify("Schüler gelöscht.", severity="information")
                        self._refresh_table()
                    except Exception as e:
                        self.app.notify(f"Fehler: {e}", severity="error")

                self.app.push_screen(
                    ConfirmScreen(f"Schüler '{name}' wirklich löschen?"),
                    _do_delete,
                )

        elif btn_id == "schueler-csv":
            self._import_csv(db_path)

        elif btn_id == "schueler-export":
            self._export_csv(db_path)

        elif btn_id == "schueler-detail":
            table = self.query_one("#schueler-table", DataTable)
            idx = table.cursor_row
            if idx is None or not (0 <= idx < len(self._schueler_liste)):
                self.app.notify("Bitte Zeile auswählen.", severity="warning")
                return
            s = self._schueler_liste[idx]
            name = s["vorname"] + ((" " + s["nachname"]) if s.get("nachname") else "")
            self.app.push_screen(SchuelerDetailScreen(self._config, s["id"], name))

    def _import_csv(self, db_path: Path) -> None:
        try:
            from tkinter import Tk, filedialog
            root = Tk()
            root.withdraw()
            path_str = filedialog.askopenfilename(
                title="CSV-Datei auswählen",
                filetypes=[("CSV", "*.csv")],
            )
            root.destroy()
            if not path_str:
                return
            csv_path = Path(path_str)
            count = ndb.import_schueler_csv(db_path, csv_path, self._klasse)
            self.app.notify(f"{count} Schüler importiert.", severity="information")
            self._refresh_table()
        except Exception as e:
            self.app.notify(f"Import-Fehler: {e}", severity="error")

    def _export_csv(self, db_path: Path) -> None:
        try:
            csv_data = ndb.export_noten_csv(db_path, self._klasse)
            if not csv_data or len(csv_data.splitlines()) <= 1:
                self.app.notify("Keine Noten zum Exportieren vorhanden.", severity="warning")
                return
            from tkinter import Tk, filedialog
            root = Tk()
            root.withdraw()
            path_str = filedialog.asksaveasfilename(
                title="Noten-CSV speichern",
                defaultextension=".csv",
                filetypes=[("CSV", "*.csv")],
                initialfile=f"noten_{self._klasse}.csv",
            )
            root.destroy()
            if path_str:
                Path(path_str).write_text(csv_data, encoding="utf-8-sig")
                self.app.notify(f"Noten exportiert: {Path(path_str).name}", severity="information")
        except Exception as e:
            self.app.notify(f"Export-Fehler: {e}", severity="error")


class SchuelerDetailScreen(ModalScreen[None]):
    """Lernverlauf (Längsschnitt) EINES Schülers – reine Anzeige, keine LLM-Calls.

    Ruft ndb.get_schueler_laengsschnitt() auf und stellt Notenverlauf (App vs. Lehrer),
    Kriterien-/K1-K3-Trend und Fehlerschwerpunkte textbasiert dar. Trendpfeile sind an die
    pädagogische Verbesserung gekoppelt (↑ = besser), nicht an die reine Zahlrichtung.
    """

    BINDINGS = [("escape", "dismiss", "Schließen")]

    def __init__(self, config: dict[str, Any], schueler_id: int, schueler_name: str = "") -> None:
        super().__init__()
        self._config = config
        self._schueler_id = schueler_id
        self._schueler_name = schueler_name

    def compose(self) -> ComposeResult:
        with Container(classes="stats-container"):
            yield Label(f"Lernverlauf – {self._schueler_name}", classes="panel-title")
            with VerticalScroll(classes="stats-scroll"):
                yield Static("", id="detail-body")
                yield Static(
                    "Sendet anonymisierte Kennzahlen an den aktiven LLM-Provider.",
                    id="profil-hinweis",
                    classes="dim",
                )
                yield Static("", id="profil-meta")
                yield Static("", id="profil-ergebnis")
            with Horizontal(classes="stats-buttons"):
                yield Button("KI-Profil erstellen", id="btn-profil", variant="primary")
                yield Button("Letztes Profil laden", id="btn-profil-load", variant="primary")
                yield Button("Schließen", id="detail-close", variant="error")

    def on_mount(self) -> None:
        body = self.query_one("#detail-body", Static)
        if ndb is None:
            body.update("Datenbank-Modul nicht verfügbar.")
            return
        try:
            db_path = ndb.get_db_path(self._config)
            ndb.init_db(db_path)
            self._laengsschnitt = ndb.get_schueler_laengsschnitt(
                db_path, self._schueler_id
            )
        except Exception as e:
            body.update(f"DB-Fehler: {e}")
            return
        body.update(self._format_laengsschnitt(self._laengsschnitt))

    def on_button_pressed(self, event: Button.Pressed) -> None:
        if event.button.id == "detail-close":
            self.dismiss(None)
            event.stop()
        elif event.button.id == "btn-profil":
            self._on_profil_click()
            event.stop()
        elif event.button.id == "btn-profil-load":
            self._on_profil_load()
            event.stop()

    def _on_profil_click(self) -> None:
        """Startet LLM-Profil-Erstellung für den aktuellen Schüler."""
        if not hasattr(self, "_laengsschnitt") or not self._laengsschnitt:
            self.app.notify(
                "Keine Längsschnittdaten vorhanden.", severity="warning"
            )
            return
        if not nc.api_key_available(self._config):
            provider = self._config.get("api", {}).get("provider", "anthropic")
            self.app.notify(
                f"API-Key für '{provider}' nicht gesetzt.", severity="error"
            )
            return
        self.app.notify("Profil wird erstellt …")
        self._start_profil_worker()

    @work(thread=True)
    def _start_profil_worker(self) -> None:
        """Worker-Thread: Prompt bauen → LLM-Call → JSON extrahieren → Anzeige."""
        try:
            prompt = nc.build_schueler_profil_prompt(self._laengsschnitt)
        except Exception as e:
            self.app.call_from_thread(
                self.app.notify,
                f"Prompt-Fehler: {e}",
                severity="error",
            )
            return

        try:
            raw = nc.run_llm_api(prompt, self._config)
        except Exception as e:
            self.app.call_from_thread(
                self.app.notify,
                f"API-Fehler: {e}",
                severity="error",
            )
            return

        if raw.startswith("FEHLER"):
            self.app.call_from_thread(
                self.app.notify,
                f"LLM-Fehler: {raw}",
                severity="error",
            )
            return

        try:
            data = nc.extract_json_from_llm(raw)
        except Exception as e:
            self.app.call_from_thread(
                self.app.notify,
                f"JSON-Extraktion fehlgeschlagen: {e}",
                severity="error",
            )
            return

        # Ergebnis formatieren und anzeigen
        self.app.call_from_thread(self._show_profil_ergebnis, data)

        # Automatisch speichern (historisch, kein Überschreiben)
        self._save_profil(data)

    def _show_profil_ergebnis(
        self,
        data: dict[str, Any],
        is_loaded: bool = False,
        basis_count: int | None = None,
        current_count: int | None = None,
        modell: str = "",
        erstellt_am: str = "",
    ) -> None:
        """Zeigt das LLM-Profil im Static-Widget an."""
        try:
            ergebnis = self.query_one("#profil-ergebnis", Static)
        except Exception:
            return

        lines: list[str] = ["", "[b]KI-Profil[/b]", ""]

        # Veraltungserkennung
        if is_loaded and basis_count is not None and current_count is not None:
            if current_count > basis_count:
                lines.append(
                    f"[yellow]⚠ Veraltet: Basis {basis_count} Abgaben, "
                    f"aktuell {current_count} Abgaben.[/yellow]"
                )
            else:
                lines.append("[green]✓ Aktuell[/green]")
            if erstellt_am:
                lines.append(f"[dim]Erstellt: {erstellt_am}[/dim]")
            if modell:
                lines.append(f"[dim]Modell: {modell}[/dim]")
            lines.append("")

        kurzbild = data.get("kurzbild", "")
        if kurzbild:
            lines.append(f"[i]{kurzbild}[/i]")
            lines.append("")

        staerken = data.get("staerken", [])
        if staerken:
            lines.append("[b]Stärken[/b]")
            for s in staerken:
                lines.append(f"  • {s}")
            lines.append("")

        foerder = data.get("foerderbereiche", [])
        if foerder:
            lines.append("[b]Förderbereiche[/b]")
            for f in foerder:
                kat = f.get("kategorie", "")
                befund = f.get("befund", "")
                uebung = f.get("uebung", "")
                lines.append(f"  • [b]{kat}[/b]")
                if befund:
                    lines.append(f"    Befund: {befund}")
                if uebung:
                    lines.append(f"    Übung: {uebung}")
            lines.append("")

        matura = data.get("maturabezug", "")
        if matura:
            lines.append(f"[b]Matura-Bezug[/b]\n  {matura}")

        ergebnis.update("\n".join(lines))
        if is_loaded:
            self.app.notify("Gespeichertes Profil geladen.")
        else:
            self.app.notify("Profil erstellt.")

    def _save_profil(self, data: dict[str, Any]) -> None:
        """Speichert das Profil automatisch nach erfolgreicher Erstellung."""
        if ndb is None or not hasattr(self, "_laengsschnitt"):
            return
        try:
            db_path = ndb.get_db_path(self._config)
            ndb.save_schueler_profil(
                db_path=db_path,
                schueler_id=self._schueler_id,
                profil=data,
                basis_anzahl_abgaben=self._laengsschnitt.get("anzahl_abgaben", 0),
                modell=self._config.get("api", {}).get("model", ""),
            )
            self.app.call_from_thread(
                self.app.notify, "Profil gespeichert.", severity="information"
            )
        except Exception:
            self.app.call_from_thread(
                self.app.notify, "Profil konnte nicht gespeichert werden.", severity="warning"
            )

    def _on_profil_load(self) -> None:
        """Lädt das jüngste gespeicherte Profil für den Schüler."""
        if ndb is None or not hasattr(self, "_laengsschnitt"):
            return
        db_path = ndb.get_db_path(self._config)
        profil_row = ndb.get_latest_schueler_profil(db_path, self._schueler_id)
        if not profil_row:
            self.app.notify("Noch kein gespeichertes Profil vorhanden.", severity="information")
            return
        data = profil_row["profil"]
        basis = profil_row["basis_anzahl_abgaben"]
        current = self._laengsschnitt.get("anzahl_abgaben", 0)
        self._show_profil_ergebnis(
            data,
            is_loaded=True,
            basis_count=basis,
            current_count=current,
            modell=profil_row.get("modell", ""),
            erstellt_am=profil_row.get("erstellt_am", ""),
        )

    @staticmethod
    def _pfeil(richtung: str) -> str:
        return {"steigt": "↑", "faellt": "↓"}.get(richtung, "→")

    @staticmethod
    def _fmt(x: Any) -> str:
        if x is None:
            return "–"
        return f"{x:g}" if isinstance(x, float) else str(x)

    def _format_laengsschnitt(self, ls: dict[str, Any]) -> str:
        sch = ls.get("schueler", {})
        anzahl = ls.get("anzahl_abgaben", 0)
        verlauf = ls.get("verlauf", [])
        trend = ls.get("trend", {})

        kopf_klasse = f"  |  Klasse: {sch.get('klasse')}" if sch.get("klasse") else ""
        lines: list[str] = [f"[b]{anzahl}[/b] Abgabe(n){kopf_klasse}", ""]

        if anzahl == 0:
            lines.append("[dim]Noch keine analysierten Abgaben für diesen Schüler erfasst.[/dim]")
            return "\n".join(lines)

        # Notenverlauf App vs. Lehrer
        app_cells = "  ".join(f"{v['aufgabe']}: {self._fmt(v['note_app'])}" for v in verlauf)
        lehrer_cells = "  ".join(f"{v['aufgabe']}: {self._fmt(v['note_lehrer'])}" for v in verlauf)
        t_app = trend.get("note_app", {})
        t_leh = trend.get("note_lehrer", {})
        lines.append("[b]Notenverlauf[/b]")
        lines.append(
            f"  Note (App):    {app_cells}   "
            f"(Trend: {t_app.get('richtung', 'stabil')} {self._pfeil(t_app.get('richtung', ''))})"
        )
        lines.append(
            f"  Note (Lehrer): {lehrer_cells}   "
            f"(Trend: {t_leh.get('richtung', 'stabil')} {self._pfeil(t_leh.get('richtung', ''))})"
        )
        lines.append("")

        # Kriterien-Trend (↑ = Verbesserung)
        hat_kriterien = any(
            any(v.get("kriterien", {}).get(k) is not None for k in
                ("inhalt", "textstruktur", "ausdruck", "sprachrichtigkeit"))
            for v in verlauf
        )
        if hat_kriterien:
            lines.append("[b]Kriterien-Trend[/b] (Stufe 1–5, ↑ = Verbesserung)")
            for key, titel in (
                ("inhalt", "Inhalt"),
                ("textstruktur", "Textstruktur"),
                ("ausdruck", "Ausdruck"),
                ("sprachrichtigkeit", "Sprachrichtigkeit"),
            ):
                t = trend.get(key, {})
                lines.append(
                    f"  {titel:<18} {self._fmt(t.get('start'))} → {self._fmt(t.get('ende'))}  "
                    f"{self._pfeil(t.get('richtung', ''))}"
                )
            lines.append("")
            for key, titel in (
                ("k1", "K1 (Inhalt + Textstruktur)"),
                ("k3", "K3 (Stil + Sprachnormen)"),
            ):
                t = trend.get(key, {})
                lines.append(
                    f"  [b]{titel}[/b]: Stufe {self._fmt(t.get('start'))} → "
                    f"{self._fmt(t.get('ende'))}  {self._pfeil(t.get('richtung', ''))}"
                )
            lines.append("")
        else:
            lines.append("[dim]Keine Kriteriendaten vorhanden (nur Noten erfasst).[/dim]")
            lines.append("")

        # Fehlerschwerpunkte
        schwerpunkte = ls.get("fehlerschwerpunkte", [])
        if schwerpunkte:
            lines.append("[b]Fehlerschwerpunkte[/b] (über alle Abgaben)")
            for f in schwerpunkte[:3]:
                beispiel = ""
                if f.get("beispiele"):
                    b = f["beispiele"][0]
                    beispiel = f"   z. B. „{b['zitat']}“ → „{b['korrektur']}“"
                lines.append(f"  • {f['label']}: {f['anzahl']}×{beispiel}")
            lines.append("")

        # Kalibrierung (nur wenn Vergleichspaare vorhanden)
        kal = ls.get("kalibrierung", {})
        if kal.get("paare", 0) > 0:
            tendenz_txt = {
                "app strenger": "App tendenziell strenger",
                "app milder": "App tendenziell milder",
                "deckungsgleich": "App und Lehrernote weitgehend deckungsgleich",
            }.get(kal.get("tendenz", ""), kal.get("tendenz", ""))
            lines.append(
                f"[dim]App vs. Lehrer: Ø Abweichung {self._fmt(kal.get('mittlere_abweichung'))} "
                f"Notenstufen, {tendenz_txt} ({kal['paare']} Vergleichspaare)[/dim]"
            )

        return "\n".join(lines)


class KlassenFeedbackScreen(ModalScreen[None]):
    """Modal-Screen: Klassen-Ansicht mit drei Reitern – Feedback, Heatmap, Statistik."""

    BINDINGS = [("escape", "dismiss", "Schließen")]

    def __init__(self, config: dict[str, Any], aufgabe: str = "") -> None:
        super().__init__()
        self._config = config
        self._klasse = nc.active_klasse(config) or ""
        self._aufgabe = aufgabe
        self._feedback_data: dict[str, Any] | None = None
        self._aggregat: dict[str, Any] | None = None
        self._briefing_data: dict[str, Any] | None = None
        self._heatmap_initialized = False
        self._statistik_initialized = False
        self._stats_view_mode = "stats"
        self._klassen_statistik: dict[str, Any] | None = None
        self._aufgaben: list[str] = []
        if self._klasse:
            try:
                self._aufgaben = list(
                    (config.get("classes", {}).get(self._klasse, {}).get("aufgaben", {})).keys()
                )
            except Exception:
                logging.debug(
                    "Aufgabenliste für Klasse %s nicht ladbar", self._klasse, exc_info=True
                )

    def compose(self) -> ComposeResult:
        aufgaben_opt = [("Alle Aufgaben", "")] + [(a, a) for a in self._aufgaben]
        with Container(classes="feedback-container"):
            yield Label(f"👥 Klasse – {self._klasse}", classes="panel-title")
            yield Button("Schließen", id="feedback-close", variant="error")
            with TabbedContent(initial="tab-feedback"):
                with TabPane("📊 Feedback", id="tab-feedback"):
                    with Horizontal(classes="feedback-controls"):
                        yield Button("🔄 Aktualisieren", id="feedback-refresh", variant="primary")
                        yield Button(
                            "🤖 Klassen-Briefing", id="feedback-briefing", variant="success"
                        )
                        yield Button(
                            "📂 Briefing laden", id="feedback-briefing-load", variant="default"
                        )
                        yield Button("📄 Als DOCX", id="feedback-docx", variant="default")
                    yield Static("", id="feedback-body")
                    yield Static(
                        "Sendet anonymisierte Kennzahlen an den aktiven LLM-Provider.",
                        id="feedback-briefing-hinweis",
                        classes="dim",
                    )
                    yield Static("", id="feedback-briefing-meta")
                    yield Static("", id="feedback-briefing-ergebnis")
                with TabPane("🔥 Heatmap", id="tab-heatmap"):
                    with Horizontal(classes="heatmap-controls"):
                        yield Select(aufgaben_opt, id="heatmap-aufgabe", value="")
                        yield Button("🔄 Aktualisieren", id="heatmap-refresh", variant="primary")
                        yield Button("📄 Als DOCX", id="heatmap-docx", variant="default")
                        yield Button("🎯 Für Übungs-Tool", id="heatmap-bridge", variant="default")
                    yield DataTable(id="heatmap-table")
                    yield Static("", id="heatmap-detail")
                with TabPane("📈 Statistik", id="tab-statistik"):
                    yield Static("", id="stats-body")
                    with Horizontal(classes="stats-buttons"):
                        yield Button("📊 Als DOCX", variant="default", id="stats-docx")
                        yield Button("📈 Fortschritt", variant="default", id="stats-progress")

    def on_mount(self) -> None:
        self._refresh_feedback()

    def on_tabbed_content_tab_activated(self, event: TabbedContent.TabActivated) -> None:
        """Lädt Tab-Inhalte beim ersten Öffnen (lazy)."""
        pane_id = event.pane.id if event.pane else ""
        if pane_id == "tab-heatmap" and not self._heatmap_initialized:
            table = self.query_one("#heatmap-table", DataTable)
            table.add_columns("Fehlertyp", "Anzahl", "% der Klasse")
            self._refresh_heatmap()
            self._heatmap_initialized = True
        elif pane_id == "tab-statistik" and not self._statistik_initialized:
            self._refresh_statistik()
            self._statistik_initialized = True

    def _refresh_feedback(self) -> None:
        body = self.query_one("#feedback-body", Static)
        if ndb is None or not self._klasse:
            body.update("Datenbank nicht verfügbar.")
            return
        db_path = ndb.get_db_path(self._config)
        try:
            ndb.init_db(db_path)
        except Exception as e:
            body.update(f"DB-Fehler: {e}")
            return
        aufgabe = self._aufgabe if self._aufgabe else None
        try:
            self._feedback_data = ndb.get_klassen_feedback(db_path, self._klasse, aufgabe)
            # Neue Aggregate laden
            k1k3 = ndb.get_klassen_k1_k3(db_path, self._klasse, aufgabe)
            noten = ndb.get_notenverteilung(db_path, self._klasse, aufgabe)
            kalib = ndb.get_klassen_kalibrierung(db_path, self._klasse, aufgabe)
            trend = ndb.get_klassen_trend(db_path, self._klasse)
            self._aggregat = {
                "k1k3": k1k3,
                "notenverteilung": noten,
                "kalibrierung": kalib,
                "trend": trend,
                "anzahl_abgaben": self._feedback_data.get("anzahl_abgaben", 0),
                "gesamt_fehler": self._feedback_data.get("gesamt_fehler", 0),
                "heatmap": self._feedback_data.get("heatmap", []),
                "beispiele": self._feedback_data.get("beispiele", []),
                "kriterien": self._feedback_data.get("kriterien", {}),
                "aufgabe": self._feedback_data.get("aufgabe", ""),
            }
            text = self._build_text()
            body.update(text)
        except Exception as e:
            body.update(f"Fehler: {e}")

    def _build_text(self) -> Text:
        from rich.text import Text as RichText
        fb = self._feedback_data
        if not fb:
            return RichText("Keine Daten.")
        t = RichText()
        t.append(f"Klasse: {fb['klasse']}  |  Aufgabe: {fb['aufgabe']}\n", style="bold")
        t.append(f"Abgaben analysiert: {fb['anzahl_abgaben']}  |  Fehler gesamt: {fb['gesamt_fehler']}\n\n", style="dim")

        # K1/K3
        agg = self._aggregat or {}
        k1k3 = agg.get("k1k3", {})
        if k1k3:
            t.append("K1 / Inhalt & Struktur  |  K3 / Sprache & Ausdruck\n", style="bold underline")
            k1 = k1k3.get("k1", {})
            k3 = k1k3.get("k3", {})
            k1_avg = k1.get("avg")
            k3_avg = k3.get("avg")
            k1_color = "#ff6b6b" if (k1_avg is not None and k1_avg < 2.5) else "#feca57" if (k1_avg is not None and k1_avg < 3.5) else "#1dd1a1"
            k3_color = "#ff6b6b" if (k3_avg is not None and k3_avg < 2.5) else "#feca57" if (k3_avg is not None and k3_avg < 3.5) else "#1dd1a1"
            t.append("  K1 Ø: ", style="default")
            t.append(f"{k1_avg:.2f}" if k1_avg is not None else "—", style=f"bold {k1_color}")
            t.append(f"  (n={k1.get('count', 0)})  |  ", style="dim")
            t.append("K3 Ø: ", style="default")
            t.append(f"{k3_avg:.2f}" if k3_avg is not None else "—", style=f"bold {k3_color}")
            t.append(f"  (n={k3.get('count', 0)})\n\n", style="dim")

        # Notenverteilung
        noten = agg.get("notenverteilung", {})
        if noten and any(v > 0 for v in noten.values()):
            t.append("Notenverteilung\n", style="bold underline")
            total_noten = sum(noten.values())
            for note in range(1, 6):
                count = noten.get(note, 0)
                pct = (count / total_noten * 100) if total_noten > 0 else 0
                color = "#1dd1a1" if note <= 2 else "#feca57" if note == 3 else "#ff6b6b"
                t.append(f"  Note {note}: ", style="default")
                t.append(f"{count:2d}", style=f"bold {color}")
                t.append(f"  ({pct:4.1f}%)\n", style="dim")
            t.append("\n")

        # Kalibrierung
        kalib = agg.get("kalibrierung", {})
        if kalib and kalib.get("n_mit_feedback", 0) > 0:
            t.append("Kalibrierung (App vs. Lehrer)\n", style="bold underline")
            t.append(f"  App-Ø: {kalib.get('app_avg', '—')}  |  ", style="default")
            t.append(f"Lehrer-Ø: {kalib.get('lehrer_avg', '—')}  |  ", style="default")
            diff = kalib.get("diff")
            diff_color = "#ff6b6b" if diff is not None and abs(diff) > 0.5 else "#feca57" if diff is not None and abs(diff) > 0.3 else "#1dd1a1"
            t.append(f"Diff: {diff:+.2f}" if diff is not None else "Diff: —", style=f"bold {diff_color}")
            t.append(f"  ({kalib.get('tendenz', 'n/a')})\n", style="dim")
            t.append(f"  (basierend auf {kalib.get('n_mit_feedback', 0)} von {kalib.get('n_gesamt', 0)} Abgaben)\n\n", style="dim")

        # Trend
        trend = agg.get("trend", [])
        if trend:
            t.append("Trend über Aufgaben\n", style="bold underline")
            for tr in trend:
                lehrer_str = f", L-Ø {tr.get('avg_note_lehrer', '—')}" if tr.get("avg_note_lehrer") is not None else ""
                t.append(
                    f"  {tr.get('aufgabe', '?')}: App-Ø {tr.get('avg_note_app', '—')}{lehrer_str} "
                    f"(n={tr.get('n', 0)})\n",
                    style="default",
                )
            t.append("\n")

        # Empfehlungen
        if fb["empfehlungen"]:
            t.append("Empfehlungen\n", style="bold underline")
            for e in sorted(fb["empfehlungen"], key=lambda x: 0 if x["prio"] == "hoch" else 1):
                color = "#ff6b6b" if e["prio"] == "hoch" else "#feca57"
                icon = "🔴" if e["prio"] == "hoch" else "🟡"
                t.append(f"  {icon} ", style=color)
                t.append(f"{e['text']}\n", style="default")
            t.append("\n")
        else:
            t.append("Keine auffälligen Defizite erkannt.\n\n", style="green")

        # Kriterien
        if fb["kriterien"]:
            t.append("Kriterien-Durchschnitte\n", style="bold underline")
            for k_name, k_data in fb["kriterien"].items():
                avg = k_data.get("avg")
                if avg is not None:
                    color = "#ff6b6b" if avg < 2.5 else "#feca57" if avg < 3.5 else "#1dd1a1"
                    t.append(f"  • {k_name.replace('_', ' ').title()}: ", style="default")
                    t.append(f"Stufe {avg:.2f}", style=f"bold {color}")
                    t.append(f"  (n={k_data['count']})\n", style="dim")
            t.append("\n")

        # Beispiel-Zitate
        if fb["beispiele"]:
            t.append("Häufige Fehler-Beispiele\n", style="bold underline")
            for b in fb["beispiele"]:
                t.append(f"  • \"{b['zitat']}\" → \"{b['korrektur']}\"\n", style="default")
                if b.get("erklaerung"):
                    t.append(f"    ({b['erklaerung']})\n", style="dim")
                t.append(f"    → {b['haeufigkeit']}x aufgetreten\n\n", style="dim")

        return t

    def on_button_pressed(self, event: Button.Pressed) -> None:
        btn_id = event.button.id
        if btn_id == "feedback-close":
            self.dismiss(None)
        elif btn_id == "feedback-refresh":
            self._refresh_feedback()
        elif btn_id == "feedback-briefing":
            self._on_briefing_click()
        elif btn_id == "feedback-briefing-load":
            self._on_briefing_load()
        elif btn_id == "feedback-docx":
            self._export_docx()
        elif btn_id == "heatmap-refresh":
            self._refresh_heatmap()
        elif btn_id == "heatmap-docx":
            self._export_docx_heatmap()
        elif btn_id == "heatmap-bridge":
            self._export_bridge()
        elif btn_id == "stats-docx":
            self._save_stats_docx()
        elif btn_id == "stats-progress":
            self._toggle_stats_view()

    # ── Heatmap-Tab ───────────────────────────────────────────────────────────

    def _refresh_heatmap(self) -> None:
        """Lädt Heatmap-Daten aus der DB und aktualisiert die Tabelle."""
        table = self.query_one("#heatmap-table", DataTable)
        table.clear()
        detail = self.query_one("#heatmap-detail", Static)
        detail.update("")
        if ndb is None or not self._klasse:
            detail.update("Datenbank nicht verfügbar.")
            return
        db_path = ndb.get_db_path(self._config)
        try:
            ndb.init_db(db_path)
        except Exception as e:
            detail.update(f"DB-Fehler: {e}")
            return
        aufgabe_widget = self.query_one("#heatmap-aufgabe", Select)
        aufgabe_val = aufgabe_widget.value
        aufgabe = str(aufgabe_val) if aufgabe_val else None
        try:
            heatmap = ndb.get_fehler_heatmap(db_path, self._klasse, aufgabe)
            if not heatmap:
                detail.update("Keine Fehlerdaten für diese Auswahl vorhanden.")
                return
            typ_labels = {"R": "Rechtschreibung", "G": "Grammatik",
                          "Z": "Zeichensetzung", "A": "Ausdruck/Stil"}
            for row in heatmap:
                label = typ_labels.get(row["typ"], row["typ"])
                table.add_row(label, str(row["anzahl"]), f"{row['prozent']:.1f}%")
            total = sum(r["anzahl"] for r in heatmap)
            detail.update(f"Gesamt: {total} Fehler erfasst.")
        except Exception as e:
            detail.update(f"Fehler: {e}")

    def _export_docx_heatmap(self) -> None:
        """Exportiert Heatmap als DOCX."""
        if ndb is None:
            return
        try:
            db_path = ndb.get_db_path(self._config)
            ndb.init_db(db_path)
            aufgabe_widget = self.query_one("#heatmap-aufgabe", Select)
            aufgabe_val = aufgabe_widget.value
            aufgabe = str(aufgabe_val) if aufgabe_val else None
            heatmap = ndb.get_fehler_heatmap(db_path, self._klasse, aufgabe)
            if not heatmap:
                self.app.notify("Keine Daten zum Exportieren.", severity="warning")
                return
            from docx import Document
            doc = Document()
            doc.add_heading(f"Fehler-Heatmap – {self._klasse}", level=1)
            if aufgabe:
                doc.add_paragraph(f"Aufgabe: {aufgabe}")
            dtable = doc.add_table(rows=1, cols=3)
            dtable.style = "Table Grid"
            hdr = dtable.rows[0].cells
            hdr[0].text = "Fehlertyp"
            hdr[1].text = "Anzahl"
            hdr[2].text = "% der Klasse"
            typ_labels = {"R": "Rechtschreibung", "G": "Grammatik",
                          "Z": "Zeichensetzung", "A": "Ausdruck/Stil"}
            for row in heatmap:
                r = dtable.add_row().cells
                r[0].text = typ_labels.get(row["typ"], row["typ"])
                r[1].text = str(row["anzahl"])
                r[2].text = f"{row['prozent']:.1f}%"
            out_dir = Path(self._config.get("paths", {}).get("output", "output"))
            out_dir.mkdir(parents=True, exist_ok=True)
            today = time.strftime("%Y-%m-%d")
            fname = f"heatmap_{self._klasse}_{today}.docx"
            out_path = out_dir / fname
            doc.save(str(out_path))
            self.app.notify(f"Heatmap gespeichert: {fname}", severity="information")
            nc.open_file(out_path)
        except Exception as e:
            self.app.notify(f"Export-Fehler: {e}", severity="error")

    def _export_bridge(self) -> None:
        """Exportiert die Korrekturdaten als JSON-Brücke für das Lehrunterlagen-Tool."""
        if ndb is None or natascha_bridge is None:
            self.app.notify("Brücken-Modul nicht verfügbar.", severity="warning")
            return
        try:
            db_path = ndb.get_db_path(self._config)
            ndb.init_db(db_path)
            aufgabe_widget = self.query_one("#heatmap-aufgabe", Select)
            aufgabe_val = aufgabe_widget.value
            aufgabe = str(aufgabe_val) if aufgabe_val else None
            if not aufgabe:
                self.app.notify(
                    "Bitte oben eine Aufgabe wählen (der Export ist pro Aufgabe).",
                    severity="warning",
                )
                return
            heatmap = ndb.get_fehler_heatmap(db_path, self._klasse, aufgabe)
            if not heatmap:
                self.app.notify("Keine Daten zum Exportieren.", severity="warning")
                return
            target = natascha_bridge.export_klassen_bridge(
                db_path, self._klasse, aufgabe, config=self._config
            )
            self.app.notify(
                f"Für Übungs-Tool exportiert: {target.name}", severity="information"
            )
        except Exception as e:
            self.app.notify(f"Brücken-Export-Fehler: {e}", severity="error")

    # ── Statistik-Tab ─────────────────────────────────────────────────────────

    def _refresh_statistik(self) -> None:
        """Lädt Klassenstatistik aus der DB und befüllt den Statistik-Tab."""
        body = self.query_one("#stats-body", Static)
        if ndb is None or not self._klasse:
            body.update("Datenbank nicht verfügbar.")
            return
        try:
            db_path = ndb.get_db_path(self._config)
            ndb.init_db(db_path)
            aufgabe = self._aufgabe if self._aufgabe else None
            self._klassen_statistik = ndb.get_klassen_statistik(db_path, self._klasse, aufgabe)
            self._stats_view_mode = "stats"
            body.update(self._build_stats_content())
        except Exception as e:
            body.update(f"Fehler beim Laden der Statistik: {e}")

    def _build_stats_content(self) -> Text:
        """Baut den Statistik-Text (Notenverteilung + Kriterien) aus gecachten Daten."""
        text = Text()
        if not self._klassen_statistik:
            text.append("Keine Statistikdaten vorhanden.", style="dim")
            return text
        s = self._klassen_statistik["stats"]
        total = s["total"]
        text.append(f"Ausgewertet: {total} Abgaben\n", style="bold")
        if total == 0:
            text.append("\nNoch keine Analysen in der Datenbank.", style="dim")
            return text
        avg = s["grade_average"]
        avg_note = round(avg)
        avg_color = _NOTE_COLORS.get(avg_note, "#ffffff")
        text.append("Gesamtdurchschnitt: ", style="bold")
        text.append(f"{avg:.2f}\n\n", style=f"bold {avg_color}")
        text.append("Notenverteilung\n", style="bold underline")
        dist = s["grade_distribution"]
        max_count = max(dist.values()) if dist else 1
        for note in range(1, 6):
            count = dist.get(note, 0)
            pct = (count / total * 100) if total > 0 else 0
            filled = round(count / max_count * _BAR_WIDTH) if max_count > 0 else 0
            bar = "█" * filled + "░" * (_BAR_WIDTH - filled)
            color = _NOTE_COLORS.get(note, "#ffffff")
            label = _NOTE_LABELS[note]
            text.append(f"  {note} {label:<12s} ", style="default")
            text.append(bar, style=color)
            text.append(f"  {count:2d}  ({pct:4.1f}%)\n", style="default")
        crit_avgs = s["criteria_averages"]
        if crit_avgs:
            text.append("\nKriterien\n", style="bold underline")
            weakest = s["weakest_criterion"]
            strongest = s["strongest_criterion"]
            for key, vals in sorted(crit_avgs.items(), key=lambda x: x[1]["avg"] or 0):
                lbl = key.replace("_", " ").title()
                avg_c = vals["avg"] or 0.0
                bar_filled = round(avg_c / 5 * _BAR_WIDTH)
                bar_c = "█" * bar_filled + "░" * (_BAR_WIDTH - bar_filled)
                if key == weakest:
                    style = f"bold {_NOTE_COLORS[5]}"
                    marker = " ◀ schwächstes"
                elif key == strongest:
                    style = f"bold {_NOTE_COLORS[1]}"
                    marker = " ◀ stärkstes"
                else:
                    style = "default"
                    marker = ""
                text.append(f"  {lbl:<22s} ", style=style)
                text.append(bar_c, style=style)
                text.append(f"  Ø {avg_c:.2f}  (n={vals['count']}){marker}\n", style=style)
        return text

    def _build_stats_progress_content(self) -> Text:
        """Baut den Fortschritts-Text (Verlauf über Aufgaben) aus gecachten Daten."""
        text = Text()
        if not self._klassen_statistik:
            text.append("Keine Daten vorhanden.", style="dim")
            return text
        progress = self._klassen_statistik["progress"]
        if not progress:
            text.append("Lernfortschritt\n\n", style="bold underline")
            text.append(
                "Keine Aufgaben mit Analysen in der Datenbank vorhanden.", style="dim"
            )
            return text
        text.append("Lernfortschritt\n\n", style="bold underline")
        short_labels = []
        for p in progress:
            lbl = p["label"]
            parts = lbl.split("–")
            short = parts[-1].strip() if len(parts) > 1 else lbl
            short_labels.append(short[:10])
        notes = [p["avg_note"] for p in progress]
        for note_val in range(5, 0, -1):
            text.append(f"  {note_val} │", style="default")
            for avg in notes:
                if avg is None:
                    text.append("   ", style="default")
                    continue
                rounded = round(avg)
                if rounded == note_val:
                    text.append("  ●", style=f"bold {_NOTE_COLORS.get(rounded, '#ffffff')}")
                else:
                    text.append("   ", style="default")
            text.append("\n", style="default")
        text.append("    └", style="default")
        for _ in short_labels:
            text.append("───", style="default")
        text.append("─\n", style="default")
        text.append("     ", style="default")
        for sl in short_labels:
            text.append(f"{sl:^10s}", style="dim")
        text.append("\n\n", style="default")
        for p in progress:
            avg_note = p["avg_note"]
            if avg_note is not None:
                n_color = _NOTE_COLORS.get(round(avg_note), "#ffffff")
                text.append(f"  {p['label']:<30s}", style="default")
                text.append(f"Ø {avg_note:.2f}", style=f"bold {n_color}")
                text.append(f"  (n={p['n']})\n", style="dim")
        if len(progress) >= 2:
            text.append("\nKriterien-Vergleich\n", style="bold underline")
            all_keys: set[str] = set()
            for p in progress:
                all_keys.update(p["avg_criteria"].keys())
            for key in sorted(all_keys):
                lbl = key.replace("_", " ").title()
                text.append(f"  {lbl:<22s}", style="default")
                for i, p in enumerate(progress):
                    val = p["avg_criteria"].get(key, 0.0)
                    trend = ""
                    if i > 0:
                        prev = progress[i - 1]["avg_criteria"].get(key)
                        if prev is not None:
                            trend = "↑" if val > prev else "↓" if val < prev else "="
                    text.append(f" {val:.1f}{trend}", style="default")
                text.append("\n", style="default")
        return text

    def _toggle_stats_view(self) -> None:
        """Wechselt zwischen Statistik- und Fortschritts-Ansicht."""
        self._stats_view_mode = "progress" if self._stats_view_mode == "stats" else "stats"
        body = self.query_one("#stats-body", Static)
        if self._stats_view_mode == "progress":
            body.update(self._build_stats_progress_content())
        else:
            body.update(self._build_stats_content())

    def _save_stats_docx(self) -> None:
        """Exportiert die Klassenstatistik als DOCX."""
        if not self._klassen_statistik:
            self.app.notify("Keine Statistikdaten vorhanden.", severity="warning")
            return
        try:
            paths = nc.build_project_paths(self._config)
            paths.output_dir.mkdir(parents=True, exist_ok=True)
            doc = gf.build_statistics_document(
                self._klassen_statistik["stats"], config=self._config, klasse_name=self._klasse
            )
            today = time.strftime("%Y-%m-%d")
            out_name = f"klassenreport_{self._klasse}_{today}.docx" if self._klasse else f"klassenreport_{today}.docx"
            out_path = paths.output_dir / out_name
            doc.save(str(out_path))
            self.app.notify(f"Klassenreport gespeichert: {out_name}", severity="information")
            nc.open_file(out_path)
        except Exception as e:
            self.app.notify(f"Fehler: {e}", severity="error")

    def _on_briefing_click(self) -> None:
        if not self._aggregat:
            self.app.notify("Keine aggregierten Daten vorhanden.", severity="warning")
            return
        if not nc.api_key_available(self._config):
            provider = self._config.get("api", {}).get("provider", "anthropic")
            self.app.notify(
                f"API-Key für '{provider}' nicht gesetzt.", severity="error"
            )
            return
        self.app.notify("Klassen-Briefing wird erstellt …")
        self._start_briefing_worker()

    @work(thread=True)
    def _start_briefing_worker(self) -> None:
        if not self._aggregat:
            self.app.call_from_thread(
                self.app.notify, "Keine Daten für Briefing.", severity="warning"
            )
            return
        try:
            prompt = nc.build_klassen_briefing_prompt(self._aggregat)
        except Exception as e:
            self.app.call_from_thread(
                self.app.notify, f"Prompt-Fehler: {e}", severity="error"
            )
            return

        try:
            raw = nc.run_llm_api(prompt, self._config)
        except Exception as e:
            self.app.call_from_thread(
                self.app.notify, f"API-Fehler: {e}", severity="error"
            )
            return

        if raw.startswith("FEHLER"):
            self.app.call_from_thread(
                self.app.notify, f"LLM-Fehler: {raw}", severity="error"
            )
            return

        try:
            data = nc.extract_json_from_llm(raw)
        except Exception as e:
            self.app.call_from_thread(
                self.app.notify,
                f"JSON-Extraktion fehlgeschlagen: {e}",
                severity="error",
            )
            return

        self.app.call_from_thread(self._show_briefing_ergebnis, data)
        self._save_briefing(data)

    def _show_briefing_ergebnis(
        self,
        data: dict[str, Any],
        is_loaded: bool = False,
        basis_abgaben: int | None = None,
        current_abgaben: int | None = None,
        basis_fehler: int | None = None,
        current_fehler: int | None = None,
        modell: str = "",
        erstellt_am: str = "",
    ) -> None:
        try:
            ergebnis = self.query_one("#feedback-briefing-ergebnis", Static)
        except Exception:
            return

        lines: list[str] = ["", "[b]KI-Klassen-Briefing[/b]", ""]

        if is_loaded:
            if (
                basis_abgaben is not None
                and current_abgaben is not None
                and basis_abgaben < current_abgaben
            ):
                lines.append(
                    f"[yellow]⚠ Veraltet: Basis {basis_abgaben} Abgaben, "
                    f"aktuell {current_abgaben} Abgaben.[/yellow]"
                )
            else:
                lines.append("[green]✓ Aktuell[/green]")
            if erstellt_am:
                lines.append(f"[dim]Erstellt: {erstellt_am}[/dim]")
            if modell:
                lines.append(f"[dim]Modell: {modell}[/dim]")
            lines.append("")

        kurzbild = data.get("kurzbild", "")
        if kurzbild:
            lines.append(f"[i]{kurzbild}[/i]")
            lines.append("")

        schwerpunkte = data.get("schwerpunkte", [])
        if schwerpunkte:
            lines.append("[b]Schwerpunkte[/b]")
            for s in schwerpunkte:
                bereich = s.get("bereich", "")
                befund = s.get("befund", "")
                empfehlung = s.get("empfehlung", "")
                lines.append(f"  • [b]{bereich}[/b]")
                if befund:
                    lines.append(f"    Befund: {befund}")
                if empfehlung:
                    lines.append(f"    Empfehlung: {empfehlung}")
            lines.append("")

        empfehlungen = data.get("unterrichtsempfehlungen", [])
        if empfehlungen:
            lines.append("[b]Unterrichtsempfehlungen[/b]")
            for e in empfehlungen:
                fokus = e.get("fokus", "")
                stundenidee = e.get("stundenidee", "")
                material = e.get("material", "")
                zielgruppe = e.get("zielgruppe", "")
                lines.append(f"  • [b]{fokus}[/b]")
                if stundenidee:
                    lines.append(f"    Stundenidee: {stundenidee}")
                if material:
                    lines.append(f"    Material: {material}")
                if zielgruppe:
                    lines.append(f"    Zielgruppe: {zielgruppe}")
            lines.append("")

        matura = data.get("matura_fokus", "")
        if matura:
            lines.append(f"[b]Matura-Fokus[/b]\n  {matura}")

        self._briefing_data = data
        ergebnis.update("\n".join(lines))
        if is_loaded:
            self.app.notify("Gespeichertes Briefing geladen.")
        else:
            self.app.notify("Klassen-Briefing erstellt.")

    def _save_briefing(self, data: dict[str, Any]) -> None:
        if ndb is None or not self._aggregat:
            return
        try:
            db_path = ndb.get_db_path(self._config)
            ndb.save_klassen_briefing(
                db_path=db_path,
                klasse=self._klasse,
                aufgabe=self._aufgabe or None,
                briefing=data,
                basis_anzahl_abgaben=self._aggregat.get("anzahl_abgaben", 0),
                basis_anzahl_fehler=self._aggregat.get("gesamt_fehler", 0),
                modell=self._config.get("api", {}).get("model", ""),
            )
            self.app.call_from_thread(
                self.app.notify, "Briefing gespeichert.", severity="information"
            )
        except Exception:
            self.app.call_from_thread(
                self.app.notify,
                "Briefing konnte nicht gespeichert werden.",
                severity="warning",
            )

    def _on_briefing_load(self) -> None:
        if ndb is None or not self._aggregat:
            return
        db_path = ndb.get_db_path(self._config)
        row = ndb.get_latest_klassen_briefing(
            db_path, self._klasse, self._aufgabe or None
        )
        if not row:
            self.app.notify(
                "Noch kein gespeichertes Briefing vorhanden.",
                severity="information",
            )
            return
        data = row["briefing"]
        basis_abg = row["basis_anzahl_abgaben"]
        basis_feh = row["basis_anzahl_fehler"]
        current_abg = self._aggregat.get("anzahl_abgaben", 0)
        current_feh = self._aggregat.get("gesamt_fehler", 0)
        self._show_briefing_ergebnis(
            data,
            is_loaded=True,
            basis_abgaben=basis_abg,
            current_abgaben=current_abg,
            basis_fehler=basis_feh,
            current_fehler=current_feh,
            modell=row.get("modell", ""),
            erstellt_am=row.get("erstellt_am", ""),
        )


    def _export_docx(self) -> None:
        if not self._feedback_data:
            self.app.notify("Keine Daten zum Exportieren.", severity="warning")
            return
        try:
            from docx import Document
            from docx.shared import RGBColor

            fb = self._feedback_data
            agg = self._aggregat or {}
            doc = Document()
            doc.add_heading(f"Klassen-Feedback – {fb['klasse']}", level=1)
            doc.add_paragraph(f"Aufgabe: {fb['aufgabe']}")
            doc.add_paragraph(f"Abgaben analysiert: {fb['anzahl_abgaben']}")
            doc.add_paragraph(f"Fehler gesamt: {fb['gesamt_fehler']}")

            # ── K1/K3, Notenverteilung, Kalibrierung ──
            k1k3 = agg.get("k1k3", {})
            if k1k3:
                doc.add_heading("K1 / Inhalt & Struktur  |  K3 / Sprache & Ausdruck", level=2)
                k1 = k1k3.get("k1", {})
                k3 = k1k3.get("k3", {})
                p = doc.add_paragraph()
                p.add_run(f"K1 Ø: {k1.get('avg', '–'):.2f} (n={k1.get('count', 0)})  |  ")
                p.add_run(f"K3 Ø: {k3.get('avg', '–'):.2f} (n={k3.get('count', 0)})")

            noten = agg.get("notenverteilung", {})
            if noten and any(v > 0 for v in noten.values()):
                doc.add_heading("Notenverteilung", level=2)
                for note in range(1, 6):
                    count = noten.get(note, 0)
                    doc.add_paragraph(f"Note {note}: {count} Abgaben", style="List Bullet")

            kalib = agg.get("kalibrierung", {})
            if kalib and kalib.get("n_mit_feedback", 0) > 0:
                doc.add_heading("Kalibrierung (App vs. Lehrer)", level=2)
                p = doc.add_paragraph()
                p.add_run(f"App-Ø: {kalib.get('app_avg', '–')}  |  ")
                p.add_run(f"Lehrer-Ø: {kalib.get('lehrer_avg', '–')}  |  ")
                p.add_run(f"Diff: {kalib.get('diff', '–'):+.2f}  ({kalib.get('tendenz', 'n/a')})")
                doc.add_paragraph(
                    f"(basierend auf {kalib.get('n_mit_feedback', 0)} von {kalib.get('n_gesamt', 0)} Abgaben)",
                    style="Intense Quote",
                )

            trend = agg.get("trend", [])
            if trend:
                doc.add_heading("Trend über Aufgaben", level=2)
                for tr in trend:
                    lehrer_str = f", Lehrer-Ø {tr.get('avg_note_lehrer', '–')}" if tr.get("avg_note_lehrer") is not None else ""
                    doc.add_paragraph(
                        f"{tr.get('aufgabe', '?')}: App-Ø {tr.get('avg_note_app', '–')}{lehrer_str} (n={tr.get('n', 0)})",
                        style="List Bullet",
                    )

            # ── Regelbasierte Empfehlungen ──
            if fb["empfehlungen"]:
                doc.add_heading("Empfehlungen", level=2)
                for e in sorted(fb["empfehlungen"], key=lambda x: 0 if x["prio"] == "hoch" else 1):
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"[{e['prio'].upper()}] ").bold = True
                    p.add_run(e["text"])

            if fb["kriterien"]:
                doc.add_heading("Kriterien-Durchschnitte", level=2)
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Kriterium"
                hdr[1].text = "Ø Stufe"
                hdr[2].text = "Anzahl"
                for k_name, k_data in fb["kriterien"].items():
                    r = table.add_row().cells
                    r[0].text = k_name.replace("_", " ").title()
                    r[1].text = f"{k_data['avg']:.2f}" if k_data.get("avg") is not None else "–"
                    r[2].text = str(k_data.get("count", "–"))

            if fb["beispiele"]:
                doc.add_heading("Häufige Fehler-Beispiele", level=2)
                for b in fb["beispiele"]:
                    p = doc.add_paragraph()
                    p.add_run(f"\"{b['zitat']}\" → \"{b['korrektur']}\"").italic = True
                    if b.get("erklaerung"):
                        p.add_run(f" ({b['erklaerung']})")
                    p.add_run(f" – {b['haeufigkeit']}x")

            # ── KI-Klassen-Briefing ──
            if self._briefing_data:
                briefing = self._briefing_data
                doc.add_page_break()
                doc.add_heading("KI-Klassen-Briefing", level=1)

                kurzbild = briefing.get("kurzbild", "")
                if kurzbild:
                    p = doc.add_paragraph()
                    p.add_run(kurzbild).italic = True

                schwerpunkte = briefing.get("schwerpunkte", [])
                if schwerpunkte:
                    doc.add_heading("Schwerpunkte", level=2)
                    for s in schwerpunkte:
                        p = doc.add_paragraph()
                        p.add_run(s.get("bereich", "")).bold = True
                        if s.get("befund"):
                            doc.add_paragraph(f"Befund: {s['befund']}", style="List Bullet")
                        if s.get("empfehlung"):
                            doc.add_paragraph(f"Empfehlung: {s['empfehlung']}", style="List Bullet")

                empfehlungen = briefing.get("unterrichtsempfehlungen", [])
                if empfehlungen:
                    doc.add_heading("Unterrichtsempfehlungen", level=2)
                    for e in empfehlungen:
                        p = doc.add_paragraph()
                        p.add_run(e.get("fokus", "")).bold = True
                        if e.get("stundenidee"):
                            doc.add_paragraph(f"Stundenidee: {e['stundenidee']}", style="List Bullet")
                        if e.get("material"):
                            doc.add_paragraph(f"Material: {e['material']}", style="List Bullet")
                        if e.get("zielgruppe"):
                            doc.add_paragraph(f"Zielgruppe: {e['zielgruppe']}", style="List Bullet")

                matura = briefing.get("matura_fokus", "")
                if matura:
                    doc.add_heading("Matura-Fokus", level=2)
                    doc.add_paragraph(matura)

            out_dir = Path(self._config.get("paths", {}).get("output", "output"))
            out_dir.mkdir(parents=True, exist_ok=True)
            today = time.strftime("%Y-%m-%d")
            fname = f"klassenfeedback_{fb['klasse']}_{today}.docx"
            out_path = out_dir / fname
            doc.save(str(out_path))
            self.app.notify(f"Feedback gespeichert: {fname}", severity="information")
            nc.open_file(out_path)
        except Exception as e:
            self.app.notify(f"Export-Fehler: {e}", severity="error")


class RetroImportScreen(ModalScreen[None]):
    """Assistent zum retroaktiven Import bestehender JSON-Analysen."""

    BINDINGS = [("escape", "dismiss", "Schließen")]

    def __init__(self, config: dict[str, Any]) -> None:
        super().__init__()
        self._config = config
        self._klasse = nc.active_klasse(config) or ""
        self._aufgabe = ""
        if self._klasse:
            kl_cfg = config.get("classes", {}).get(self._klasse, {})
            self._aufgabe = kl_cfg.get("active_aufgabe", "")

    def compose(self) -> ComposeResult:
        with VerticalScroll(classes="retro-container"):
            yield Label("Retro-Import bestehender Analysen", classes="panel-title")
            yield Static(
                f"Klasse: {self._klasse} | Aufgabe: {self._aufgabe}\n"
                "Importiert alle JSON-Dateien aus output/feedback_data/ in die Datenbank.",
                id="retro-info",
            )
            yield ProgressBar(id="retro-progress", total=100)
            yield Static("", id="retro-log")
            with Horizontal(classes="retro-buttons"):
                yield Button("▶ Starten", id="retro-start", variant="success")
                yield Button("Schließen", id="retro-close", variant="error")

    def on_button_pressed(self, event: Button.Pressed) -> None:
        if event.button.id == "retro-close":
            self.dismiss(None)
        elif event.button.id == "retro-start":
            self._start_import()

    def _start_import(self) -> None:
        if ndb is None:
            self.query_one("#retro-log", Static).update("Datenbank nicht verfügbar.")
            return
        if not self._klasse or not self._aufgabe:
            self.query_one("#retro-log", Static).update("Keine Klasse/Aufgabe aktiv.")
            return
        log = self.query_one("#retro-log", Static)
        progress = self.query_one("#retro-progress", ProgressBar)
        log.update("Suche JSON-Dateien...")
        try:
            paths = nc.build_project_paths(self._config)
            feedback_dir = paths.feedback_data_dir
            if not feedback_dir.exists():
                log.update(f"Ordner nicht gefunden: {feedback_dir}")
                return
            json_files = list(feedback_dir.glob("*.json"))
            if not json_files:
                log.update("Keine JSON-Dateien gefunden.")
                return
            db_path = ndb.get_db_path(self._config)
            ndb.init_db(db_path)
            # Input-Dir fuer Originaldateien (Hash-Berechnung)
            input_dir = paths.input_dir
            erfolgreich = 0
            uebersprungen = 0
            for i, json_path in enumerate(json_files):
                progress.progress = (i + 1) / len(json_files) * 100
                # Versuche Originaldatei zu finden
                orig = None
                if input_dir:
                    candid = input_dir / json_path.name.replace("_analysis.json", ".docx")
                    if candid.exists():
                        orig = candid
                abg_id = ndb.import_json_to_db(
                    db_path, json_path, self._klasse, self._aufgabe, orig
                )
                if abg_id >= 0:
                    erfolgreich += 1
                else:
                    uebersprungen += 1
            log.update(
                f"Import abgeschlossen:\n"
                f"  Erfolgreich: {erfolgreich}\n"
                f"  Übersprungen (Duplikate): {uebersprungen}"
            )
            progress.progress = 100
            self.app.notify(f"Retro-Import: {erfolgreich} importiert.", severity="information")
        except Exception as e:
            log.update(f"Fehler: {e}")
            self.app.notify(f"Import-Fehler: {e}", severity="error")


class NataschaFooter(Widget):
    """Zweizeiliger Footer: 9 Aktions-Buttons oben, reaktive Statuszeile unten."""

    def __init__(self, app_instance: "NataschaApp") -> None:
        super().__init__()
        self._app = app_instance

    def compose(self) -> ComposeResult:
        # Reihenfolge nach UX-Best-Practice: primäre Aktionen links, Navigation,
        # destruktive Aktion (Beenden) ganz rechts.
        with Horizontal(id="footer-action-bar"):
            yield Button("✓ Analysieren", id="footer-analyze", classes="footer-btn")
            yield Button("📄 DOCX", id="footer-docx", classes="footer-btn")
            yield Button("⚙  Zuordnung", id="footer-edit", classes="footer-btn")
            yield Button("⚙  Einstellungen", id="footer-settings", classes="footer-btn")
            yield Button("? Hilfe", id="footer-help", classes="footer-btn")
            yield Button("👤 Schüler", id="footer-schueler", classes="footer-btn")
            yield Button("👥 Klasse", id="footer-klasse", classes="footer-btn")
            yield Button("📂 Ordner", id="footer-open-folder", classes="footer-btn")
            yield Button("✕ Beenden", id="footer-quit", classes="footer-btn footer-btn-quit")

    def on_button_pressed(self, event: Button.Pressed) -> None:
        action_map: dict[str, Any] = {
            "footer-open-folder": self._app.action_open_folder,
            "footer-analyze": self._app.action_analyze_file,
            "footer-edit": self._app.action_edit_assignment,
            "footer-settings": self._app.action_show_settings,
            "footer-docx": self._app.action_generate_docx,
            "footer-help": self._app.action_show_help,
            "footer-quit": self._app.action_quit,
            "footer-schueler": self._app.action_show_schueler,
            "footer-klasse": self._app.action_show_klassen_feedback,
        }
        handler = action_map.get(event.button.id)
        if handler:
            handler()
            event.stop()


class NataschaApp(App):
    TITLE = "NATASCHA"
    CSS_PATH = "natascha.tcss"

    BINDINGS = [
        Binding("q", "quit", "Beenden", show=False),
        Binding("question_mark", "show_help", "Hilfe", key_display="?", show=False),
        Binding("f1", "show_help", "Hilfe", show=False),
        Binding("a", "analyze_file", "Analysieren", show=False),
        Binding("shift+a", "analyze_marked", "Alle analysieren", show=False),
        Binding("d", "generate_docx", "DOCX", show=False),
        Binding("shift+d", "generate_docx_marked", "Alle DOCX", show=False),
        Binding("e", "edit_assignment", "Zuordnung", show=False),
        Binding("r", "review_file", "Review", show=False),
        Binding("slash", "toggle_search", "Suche", key_display="/", show=False),
        Binding("1", "preview_tab('text')", "Text", show=False),
        Binding("2", "preview_tab('rating')", "Bewertung", show=False),
        Binding("3", "preview_tab('rubrik')", "Rubrik", show=False),
        Binding("4", "preview_tab('output')", "Output", show=False),
        Binding("s", "show_settings", "Einstellungen", show=False),
        Binding("tab", "next_panel", "Naechstes Panel", show=False),
        Binding("shift+e", "apply_aufgabe_to_all", "Aufgabe→Alle", show=False),
        Binding("o", "toggle_sort", "Sortierung", show=False),
        Binding("ctrl+a", "select_all", "Alle markieren", show=False),
        Binding("ctrl+d", "deselect_all", "Alle demarkieren", show=False),
        Binding("g", "generate_erwartungshorizont", "EH generieren", show=False),
        Binding("shift+s", "show_schueler", "Schüler", show=False),
        Binding("shift+f", "show_klassen_feedback", "Klasse", show=False),
        Binding("f2", "retro_import", "Import", show=False),
    ]

    selected_index: reactive[int] = reactive(0)
    preview_mode: reactive[str] = reactive("text")
    sort_mode: reactive[str] = reactive("name")
    search_active: reactive[bool] = reactive(False)

    def __init__(self) -> None:
        super().__init__()
        self.config = nc.load_config()
        self.files: list[FileInfo] = []
        self._id_to_index: dict[str, int] = {}
        self._search_filter: str = ""
        self._filtered_indices: list[int] = []
        self._focus_panel: int = 0
        self._analysis_cancelled = False
        self._rebuilding_list: bool = False
        self._cancel_event = threading.Event()
        self._switching_context = False
        self._known_files: set[Path] = set()
        self._last_click_id: str = ""
        self._last_click_time: float = 0.0
        self._output_paths: dict[str, Path] = {}  # output-list item id → Path

    def compose(self) -> ComposeResult:
        yield NataschaHeader(id="app-header")
        with Horizontal(id="main-container"):
            with Vertical(id="files-panel"):
                yield Label("DATEIEN", classes="panel-title")
                yield Select([], id="class-select", prompt="Klasse wählen…")
                yield Button("＋ Klasse", id="add-class-btn", classes="add-class-btn")
                yield Select([], id="aufgabe-select", prompt="Aufgabe wählen…")
                with Vertical(id="action-row-aufgabe"):
                    with Horizontal(classes="action-sub-row"):
                        yield Button("＋ Aufgabe", id="add-aufgabe-btn", classes="action-row-btn")
                        yield Button("📎 Rubrik", id="attach-rubric-btn", classes="action-row-btn")
                        yield Button(
                            "⬇ →Alle", id="apply-aufgabe-all-btn", classes="action-row-btn"
                        )
                    with Horizontal(classes="action-sub-row"):
                        yield Button(
                            "📂 Ordner öffnen", id="btn-open-folder", classes="action-row-btn"
                        )
                        yield Button(
                            "📁 Ausgangstext", id="btn-open-ausgangstext", classes="action-row-btn"
                        )
                yield Input(placeholder="Suche...", id="search-input")
                yield ListView(id="file-list")
                yield Static("", id="file-counter")
            with Vertical(id="middle-panel"):
                yield Label("ZUORDNUNG", classes="panel-title")
                yield Collapsible(
                    VerticalScroll(
                        Static("App-Note: –", id="lf-app-note"),
                        RadioSet(
                            RadioButton("1 – Sehr gut", id="lf-note-1"),
                            RadioButton("2 – Gut", id="lf-note-2"),
                            RadioButton("3 – Befriedigend", id="lf-note-3"),
                            RadioButton("4 – Genügend", id="lf-note-4"),
                            RadioButton("5 – Nicht genügend", id="lf-note-5"),
                            id="lf-note-radioset",
                        ),
                        TextArea(
                            placeholder="Optionaler Kommentar zur Note …",
                            id="lf-kommentar",
                        ),
                        Button("Speichern", id="lf-speichern", variant="success"),
                        Static("", id="lf-status"),
                        id="lf-content",
                    ),
                    title="Echte Note eintragen",
                    id="lehrer-feedback-collapsible",
                )
                yield VerticalScroll(Static("", id="middle-content"), id="middle-scroll")
                with Vertical(id="action-bar"):
                    yield Button(
                        "▶  Analysieren", id="btn-analyze", variant="success", classes="action-btn"
                    )
                    yield Button(
                        "👁  Review öffnen", id="btn-review", variant="primary", classes="action-btn"
                    )
                    yield Button(
                        "📄  DOCX erstellen", id="btn-docx", variant="primary", classes="action-btn"
                    )
                    yield Button(
                        "✏  Zuordnung bearbeiten",
                        id="btn-edit",
                        variant="default",
                        classes="action-btn",
                    )
                    yield Button(
                        "🔄  Neu analysieren",
                        id="btn-reanalyze",
                        variant="warning",
                        classes="action-btn",
                    )
            with VerticalScroll(id="preview-panel"):
                yield Label("VORSCHAU", classes="panel-title")
                with Horizontal(id="preview-tabs"):
                    yield Button("1 Text", id="tab-text", classes="tab-btn tab-active")
                    yield Button("2 Bewertung", id="tab-rating", classes="tab-btn")
                    yield Button("3 Rubrik", id="tab-rubrik", classes="tab-btn")
                    yield Button("4 Output", id="tab-output", classes="tab-btn")
                yield Static("", id="preview-content")
                yield Button(
                    "✏️  Rubrik bearbeiten",
                    id="btn-edit-rubric",
                    variant="default",
                    classes="action-btn",
                )
                yield ListView(id="output-list")
        yield NataschaFooter(self)

    def on_mount(self) -> None:
        self._populate_class_select()
        self._populate_aufgabe_select()
        self._load_files()
        self._update_header()
        self._apply_defaults()
        self._update_all_panels()

        self._lf_status_cache: dict[str, bool] = {}

        search_input = self.query_one("#search-input", Input)
        search_input.display = False

        output_list = self.query_one("#output-list", ListView)
        output_list.display = False

        lf_collapsible = self.query_one("#lehrer-feedback-collapsible", Collapsible)
        lf_collapsible.display = False

        if not nc.api_key_available(self.config):
            self._check_first_run()

        self._watch_input_dir()

    def _populate_class_select(self) -> None:
        """Befüllt das Klassen-Select-Widget aus der Config."""
        class_sel = self.query_one("#class-select", Select)
        names = nc.list_classes(self.config)
        if not names:
            class_sel.display = False
            self.query_one("#add-class-btn", Button).display = False
            return
        options = [(n, n) for n in names]
        class_sel.set_options(options)
        active = nc.active_klasse(self.config)
        if active and active in names:
            class_sel.value = active

    def _populate_aufgabe_select(self) -> None:
        """Befüllt das Aufgaben-Select-Widget für die aktive Klasse."""
        aufgabe_sel = self.query_one("#aufgabe-select", Select)
        action_row = self.query_one("#action-row-aufgabe")
        rubric_btn = self.query_one("#attach-rubric-btn", Button)
        klasse = nc.active_klasse(self.config)
        if not klasse:
            aufgabe_sel.display = False
            action_row.display = False
            return
        slugs = nc.list_aufgaben(self.config, klasse)
        aufgabe_sel.display = True
        action_row.display = True
        rubric_btn.display = bool(slugs)
        if not slugs:
            aufgabe_sel.set_options([])
            return
        cls_cfg = self.config.get("classes", {}).get(klasse, {})
        options = [(cls_cfg.get("aufgaben", {}).get(s, {}).get("label", s), s) for s in slugs]
        aufgabe_sel.set_options(options)
        active = nc.active_aufgabe(self.config, klasse)
        if active and active in slugs:
            aufgabe_sel.value = active

    _WATCHED_PATTERNS: list[str] = ["*.docx", "*.odt", "*.pdf", "*.jpg", "*.jpeg", "*.png"]

    def _load_files(self) -> None:
        input_dir = nc.build_project_paths(self.config).input_dir
        if not input_dir.exists():
            return

        all_files: list[Path] = []
        for pattern in self._WATCHED_PATTERNS:
            all_files.extend(input_dir.glob(pattern))
        all_files = sorted(all_files)

        self.files = []
        for p in all_files:
            ft = _detect_file_type(p)
            wc = nc.count_words(p) if ft == "docx" else nc.get_file_size_kb(p)
            fi = FileInfo(path=p, word_count=wc, file_type=ft)
            self._load_existing_analysis(fi)
            self.files.append(fi)

        self._known_files = {fi.path for fi in self.files}
        self._rebuild_id_map()
        self._apply_filter()

    def _load_existing_analysis(self, fi: FileInfo) -> None:
        paths = nc.build_project_paths(self.config)
        analysis_name = fi.path.stem + "_analysis.json"
        analysis_path = paths.feedback_data_dir / analysis_name
        if analysis_path.exists():
            try:
                data = json.loads(analysis_path.read_text(encoding="utf-8"))
                fi.analysis = data
                fi.fach = data.get("fach", fi.fach)
                fi.schulstufe = data.get("schulstufe", fi.schulstufe)
                fi.textsorte = data.get("textsorte", fi.textsorte)
                fi.rubric = data.get("rubrik", fi.rubric)
                fi.schueler = data.get("schueler", fi.schueler)
                fi.bewertungsmodus = data.get("bewertungsmodus", "benotet")
                fi.erwartungshorizont = data.get("erwartungshorizont", fi.erwartungshorizont)
                docx_name = gf.output_filename(data.get("datei", fi.path.name))
                if (paths.output_dir / docx_name).exists():
                    fi.status = FileStatus.DONE
                else:
                    fi.status = FileStatus.ANALYZED
            except (json.JSONDecodeError, Exception):
                fi.status = FileStatus.ERROR

    def _apply_defaults(self) -> None:
        klasse = nc.active_klasse(self.config)
        aufgabe = nc.active_aufgabe(self.config, klasse) if klasse else None
        auf_defs = nc.aufgabe_defaults(self.config, klasse, aufgabe)
        global_defs = self.config.get("defaults", {})
        for fi in self.files:
            if not fi.fach:
                fi.fach = auf_defs.get("fach") or global_defs.get("fach", "Deutsch")
            if not fi.schulstufe:
                fi.schulstufe = auf_defs.get("schulstufe") or global_defs.get(
                    "schulstufe", "Oberstufe"
                )
            if not fi.textsorte:
                fi.textsorte = auf_defs.get("textsorte") or "Kommentar"
            if not fi.rubric:
                fi.rubric = auf_defs.get("rubric") or nc.default_rubric_for(
                    fi.fach, fi.schulstufe, self.config
                )
            if fi.bewertungsmodus == "benotet":
                fi.bewertungsmodus = auf_defs.get("bewertungsmodus", "benotet")
            if not fi.erwartungshorizont:
                fi.erwartungshorizont = auf_defs.get("erwartungshorizont", "")

    @work(thread=True)
    def _watch_input_dir(self) -> None:
        """Pollt alle 10 Sekunden den input/-Ordner auf neue Dateien (DOCX, PDF, Bilder)."""
        while True:
            time.sleep(10)
            input_dir = nc.build_project_paths(self.config).input_dir
            if not input_dir.exists():
                continue
            try:
                current_paths: set[Path] = set()
                for pattern in self._WATCHED_PATTERNS:
                    current_paths.update(input_dir.glob(pattern))
                known_paths = {fi.path for fi in self.files}
                new_paths = sorted(current_paths - known_paths)
                if new_paths:
                    self.call_from_thread(self._on_new_files_detected, new_paths)
            except Exception:
                logging.debug("File-Watcher-Poll fehlgeschlagen", exc_info=True)

    def _on_new_files_detected(self, new_paths: list[Path]) -> None:
        """Fügt neu erkannte Dateien zur Dateiliste hinzu (läuft im Main-Thread)."""
        klasse = nc.active_klasse(self.config)
        aufgabe = nc.active_aufgabe(self.config, klasse) if klasse else None
        auf_defs = nc.aufgabe_defaults(self.config, klasse, aufgabe)
        global_defs = self.config.get("defaults", {})
        added = 0
        for p in new_paths:
            if any(fi.path == p for fi in self.files):
                continue
            ft = _detect_file_type(p)
            wc = nc.count_words(p) if ft == "docx" else nc.get_file_size_kb(p)
            fi = FileInfo(path=p, word_count=wc, file_type=ft)
            self._load_existing_analysis(fi)
            if not fi.fach:
                fi.fach = auf_defs.get("fach") or global_defs.get("fach", "Deutsch")
            if not fi.schulstufe:
                fi.schulstufe = auf_defs.get("schulstufe") or global_defs.get(
                    "schulstufe", "Oberstufe"
                )
            if not fi.textsorte:
                fi.textsorte = auf_defs.get("textsorte") or "Kommentar"
            if not fi.rubric:
                fi.rubric = auf_defs.get("rubric") or nc.default_rubric_for(
                    fi.fach, fi.schulstufe, self.config
                )
            if fi.bewertungsmodus == "benotet":
                fi.bewertungsmodus = auf_defs.get("bewertungsmodus", "benotet")  # type: ignore[assignment]
            self.files.append(fi)
            added += 1

        if added > 0:
            self._known_files = {fi.path for fi in self.files}
            self._rebuild_id_map()
            self._apply_filter()
            self._update_all_panels()
            self._update_header()
            self.notify(
                f"{added} neue Datei{'en' if added > 1 else ''} erkannt",
                severity="information",
            )

    def _rebuild_id_map(self) -> None:
        self._id_to_index = {}
        for i, fi in enumerate(self.files):
            sid = safe_id("fi", fi.path.name)
            self._id_to_index[sid] = i

    def _apply_filter(self) -> None:
        if self._search_filter:
            self._filtered_indices = [
                i
                for i, fi in enumerate(self.files)
                if self._search_filter.lower() in fi.path.name.lower()
            ]
        else:
            self._filtered_indices = list(range(len(self.files)))

        if self.sort_mode == "status":
            self._filtered_indices.sort(
                key=lambda i: (
                    0
                    if self.files[i].status == FileStatus.ERROR
                    else 1
                    if self.files[i].status == FileStatus.PENDING
                    else 2
                    if self.files[i].status == FileStatus.ANALYZED
                    else 3
                    if self.files[i].status == FileStatus.PROGRESS
                    else 4
                )
            )
        elif self.sort_mode == "words":
            self._filtered_indices.sort(key=lambda i: self.files[i].word_count)
        else:
            self._filtered_indices.sort(key=lambda i: self.files[i].path.name.lower())

        if self.selected_index >= len(self._filtered_indices):
            self.selected_index = max(0, len(self._filtered_indices) - 1)

    def _update_header(self) -> None:
        api_status = "✓" if nc.api_key_available(self.config) else "✗"
        self.query_one("#app-header", NataschaHeader).update_status(
            api_status, len(self.files), nc.VERSION
        )

    def _update_all_panels(self) -> None:
        for fn in (
            self._update_file_list,
            self._update_middle_panel,
            self._update_preview_panel,
            self._update_counter,
        ):
            try:
                fn()
            except Exception:
                logging.exception("Panel update failed: %s", fn.__name__)

    def _has_lehrer_feedback_cached(self, fi: FileInfo) -> bool:
        """Prüft (mit Cache), ob für eine Datei bereits Lehrer-Feedback existiert."""
        path_str = str(fi.path)
        if path_str in self._lf_status_cache:
            return self._lf_status_cache[path_str]
        if ndb is None:
            return False
        try:
            db_path = ndb.get_db_path(self.config)
            if not db_path.exists():
                return False
            file_hash = ndb._file_hash(fi.path)
            result = ndb.has_lehrer_feedback_for_file(db_path, file_hash)
            self._lf_status_cache[path_str] = result
            return result
        except Exception:
            return False

    def _update_file_list(self) -> None:
        list_view = self.query_one("#file-list", ListView)
        self._rebuilding_list = True

        wanted_ids = [
            safe_id("fi", self.files[i].path.name)
            for i in self._filtered_indices
            if i < len(self.files)
        ]

        # Vorhandene IDs ermitteln
        existing_ids = {item.id for item in list_view.query(ListItem)}

        # Überflüssige Items entfernen
        for item in list(list_view.query(ListItem)):
            if item.id not in wanted_ids:
                item.remove()

        # Items erstellen oder Label aktualisieren
        for pos, real_idx in enumerate(self._filtered_indices):
            if real_idx >= len(self.files):
                continue
            fi = self.files[real_idx]
            sym, cls = STATUS_SYMBOLS[fi.status]
            mark = "☑ " if fi.marked else "☐ "
            size_str = f"{fi.word_count} W" if fi.file_type == "docx" else f"{fi.word_count} KB"
            lf_marker = "✓" if self._has_lehrer_feedback_cached(fi) else ""
            label_text = f" {mark}{lf_marker}{sym} {fi.path.name}  ({size_str})"
            item_id = safe_id("fi", fi.path.name)
            if item_id in existing_ids:
                # Nur Label aktualisieren — kein neues Widget mounten
                try:
                    existing = list_view.query_one(f"#{item_id}", ListItem)
                    existing.query_one(Static).update(Text(label_text))
                except Exception:
                    logging.debug("ListItem-Update fehlgeschlagen (id=%s)", item_id, exc_info=True)
            else:
                item = ListItem(Static(Text(label_text)), id=item_id)
                list_view.append(item)

        if self._filtered_indices:
            list_view.index = min(self.selected_index, len(self._filtered_indices) - 1)
        self._rebuilding_list = False

    def _update_counter(self) -> None:
        """Aktualisiert den Datei-Zähler im Datei-Panel."""
        total = len(self.files)
        analyzed = sum(
            1 for fi in self.files if fi.status in (FileStatus.ANALYZED, FileStatus.DONE)
        )
        marked = sum(1 for fi in self.files if fi.marked)

        pct = int((analyzed / total) * 100) if total > 0 else 0
        filled = int(pct / 5)
        bar_str = "█" * filled + "░" * (20 - filled)

        try:
            counter = self.query_one("#file-counter", Static)
            parts = [f"{analyzed}/{total}  {bar_str} {pct}%"]
            if marked > 0:
                parts.append(f"☑ {marked}")
            counter.update("  ".join(parts))
        except Exception:
            logging.debug("Datei-Counter-Update fehlgeschlagen", exc_info=True)

    def _update_middle_panel(self) -> None:
        content = self.query_one("#middle-content", Static)
        if not self._filtered_indices:
            content.update("Keine Dateien gefunden.")
            return

        idx = (
            self._filtered_indices[self.selected_index]
            if self.selected_index < len(self._filtered_indices)
            else None
        )
        if idx is None or idx >= len(self.files):
            content.update("")
            return

        fi = self.files[idx]
        if fi.file_type == "docx":
            groesse_zeile = f"Woerter:   {fi.word_count}"
        else:
            typ_label = "PDF" if fi.file_type == "pdf" else "Bild"
            groesse_zeile = f"Groesse:   {fi.word_count} KB  [{typ_label}]"
        lines = [
            f"Datei:     {fi.path.name}",
            groesse_zeile,
            f"Fach:      {fi.fach}",
            f"Schulstufe:{fi.schulstufe}",
            f"Textsorte: {fi.textsorte}",
            f"Rubrik:    {fi.rubric}",
            f"Erw.-Hor.: {fi.erwartungshorizont or '(keiner)'}",
        ]
        if fi.schueler:
            lines.append(f"Schüler/in:{fi.schueler}")
        modus_label = "Hausaufgabe (unbenotet)" if fi.bewertungsmodus == "unbenotet" else "Schularbeit (benotet)"
        lines.append(f"Modus:     {modus_label}")
        _klasse = nc.active_klasse(self.config)
        _aufgabe = nc.active_aufgabe(self.config, _klasse) if _klasse else None
        ausgangstext_line: tuple[str, str] | None = None
        if _klasse and _aufgabe:
            _at = nc.detect_ausgangstext(self.config, _klasse, _aufgabe)
            if _at:
                ausgangstext_line = (f"✓ Ausgangstext: {_at.name}", "bold green")
            else:
                ausgangstext_line = ("⚠ Ausgangstext: keiner  [📁 Ausgangstext → Datei ablegen]", "yellow")

        text = Text("\n".join(lines))
        if ausgangstext_line:
            text.append("\n")
            text.append(ausgangstext_line[0], style=ausgangstext_line[1])

        if fi.analysis:
            modell = fi.analysis.get("modell", "")
            provider_name = fi.analysis.get("provider", "")
            if modell or provider_name:
                llm_info = modell
                if provider_name:
                    llm_info += f" ({provider_name})" if modell else provider_name
                text.append(f"\nAnalysiert mit: {llm_info}", style="italic #808080")
            paths = nc.build_project_paths(self.config)
            history_dir = paths.feedback_data_dir / ".history"
            if history_dir.exists():
                version_count = len(list(history_dir.glob(f"{fi.path.stem}_analysis_*.json")))
                if version_count:
                    text.append(f"\nVersionen: {version_count}", style="italic #808080")
            note_data = fi.analysis.get("notenempfehlung")
            notendetail = fi.analysis.get("notendetail")
            text.append("\n")
            if note_data and fi.bewertungsmodus != "unbenotet":
                note = note_data.get("note", "?")
                bez = note_data.get("bezeichnung", "?")
                schnitt = note_data.get("durchschnitt", "?")
                note_color = _NOTE_COLORS.get(int(note) if str(note).isdigit() else 0, "#ffffff")
                text.append(
                    f"\nNote: {note} – {bez}",
                    style=f"bold {note_color}",
                )
                text.append(f"\nDurchschnitt: {schnitt}")

                if notendetail:
                    k1_note = notendetail.get("k1_note")
                    k1_schnitt = notendetail.get("k1_schnitt")
                    k3_note = notendetail.get("k3_note")
                    k3_schnitt = notendetail.get("k3_schnitt")
                    if k1_note is not None:
                        k1_s = f" [Stufe {k1_schnitt:.1f}]" if k1_schnitt is not None else ""
                        text.append(f"\n  K1 (Inhalt + Textstruktur): Note {k1_note}{k1_s}")
                    if k3_note is not None:
                        k3_s = f" [Stufe {k3_schnitt:.1f}]" if k3_schnitt is not None else ""
                        text.append(f"\n  K3/1 (Stil + Sprachnormen): Note {k3_note}{k3_s}")
                    sonderregel = notendetail.get("sonderregel")
                    if sonderregel:
                        text.append(
                            f"\n  ⚠ Sonderregel: {sonderregel}",
                            style="bold #C00000",
                        )

                llm_nd = fi.analysis.get("notenempfehlung_llm")
                if llm_nd and str(llm_nd.get("note")) != str(note):
                    text.append(
                        f"\n⚠ LLM empfiehlt: {llm_nd.get('note')} – {llm_nd.get('bezeichnung', '?')} (Abweichung)",
                        style="italic #808080",
                    )
            else:
                text.append("\nHausaufgabe – keine Benotung", style="italic #808080")
            text.append("\n\nKriterien:")

            bewertung = fi.analysis.get("bewertung", {})
            for key, crit in bewertung.items():
                if isinstance(crit, dict):
                    punkte = crit.get("punkte", 0)
                    filled = int(punkte) if isinstance(punkte, (int, float)) else 0
                    bar = "●" * filled + "○" * (5 - filled)
                    text.append(f"\n  {key.replace('_', ' ').title():20s} {punkte}  {bar}")

        content.update(text)
        self._update_action_bar(fi)
        self._update_lehrer_feedback_panel(fi)

    def _update_lehrer_feedback_panel(self, fi: FileInfo) -> None:
        """Aktualisiert den Lehrer-Feedback-Block im mittleren Panel."""
        try:
            collapsible = self.query_one("#lehrer-feedback-collapsible", Collapsible)
        except Exception:
            return

        if not fi.analysis or fi.bewertungsmodus == "unbenotet":
            collapsible.display = False
            return

        note_data = fi.analysis.get("notenempfehlung")
        if not note_data:
            collapsible.display = False
            return

        collapsible.display = True

        app_note = note_data.get("note", "?")
        try:
            app_note_static = self.query_one("#lf-app-note", Static)
            app_note_static.update(f"App-Note: {app_note}")
        except Exception:
            pass

        abgabe_id = fi.analysis.get("_abgabe_id")
        # Wenn keine _abgabe_id (z.B. aus JSON geladen), versuche über Hash zu finden
        if not abgabe_id and ndb is not None:
            try:
                db_path = ndb.get_db_path(self.config)
                ndb.init_db(db_path)
                file_hash = ndb._file_hash(fi.path)
                existing = ndb.get_abgabe_by_hash(db_path, file_hash)
                if existing and existing.get("id"):
                    abgabe_id = existing["id"]
                    fi.analysis["_abgabe_id"] = abgabe_id
            except Exception:
                pass

        note_final = None
        kommentar = ""
        if abgabe_id and ndb is not None:
            try:
                db_path = ndb.get_db_path(self.config)
                ndb.init_db(db_path)
                lf = ndb.get_lehrer_feedback(db_path, abgabe_id)
                if lf:
                    note_final = lf.get("note_final")
                    kommentar = lf.get("lehrer_kommentar") or ""
            except Exception:
                pass

        try:
            radio_set = self.query_one("#lf-note-radioset", RadioSet)
            selected_note = None
            if note_final is not None:
                selected_note = int(note_final)
            elif str(app_note).isdigit():
                selected_note = int(app_note)
            if selected_note is not None and 1 <= selected_note <= 5:
                radio_set.pressed_index = selected_note - 1
        except Exception:
            pass

        try:
            text_area = self.query_one("#lf-kommentar", TextArea)
            text_area.text = kommentar
        except Exception:
            pass

        try:
            status_static = self.query_one("#lf-status", Static)
            status_static.update("")
        except Exception:
            pass

    def _update_action_bar(self, fi: FileInfo) -> None:
        """Blendet Buttons je nach Dateistatus ein oder aus."""
        status = fi.status
        analyzed = status in (FileStatus.ANALYZED, FileStatus.DONE)
        pending = status == FileStatus.PENDING
        in_progress = status == FileStatus.PROGRESS
        error = status == FileStatus.ERROR

        self.query_one("#btn-analyze", Button).display = pending or error
        self.query_one("#btn-review", Button).display = analyzed
        self.query_one("#btn-docx", Button).display = analyzed
        self.query_one("#btn-edit", Button).display = True
        self.query_one("#btn-reanalyze", Button).display = analyzed
        # Während Analyse läuft: alle deaktivieren
        for btn_id in ("#btn-analyze", "#btn-review", "#btn-docx", "#btn-edit", "#btn-reanalyze"):
            self.query_one(btn_id, Button).disabled = in_progress

    def _action_save_lehrer_feedback(self) -> None:
        """Speichert die vom Lehrer eingetragene Note und den Kommentar in die DB."""
        fi = self._get_selected_file()
        if not fi or not fi.analysis:
            self.notify("Keine Analyse vorhanden.", severity="warning")
            return

        abgabe_id = fi.analysis.get("_abgabe_id")
        if not abgabe_id:
            self.notify(
                "Keine Abgabe-ID vorhanden (DB-Speicherung fehlgeschlagen"
                " oder Duplikat ohne Cache).",
                severity="warning",
            )
            return

        try:
            radio_set = self.query_one("#lf-note-radioset", RadioSet)
            pressed = radio_set.pressed_index
            if pressed is None:
                self.notify("Bitte eine Note auswählen.", severity="warning")
                return
            note_final = float(pressed + 1)
        except Exception:
            self.notify("Fehler beim Lesen der Note.", severity="error")
            return

        try:
            text_area = self.query_one("#lf-kommentar", TextArea)
            kommentar = text_area.text
        except Exception:
            kommentar = ""

        _klasse = nc.active_klasse(self.config)
        _aufgabe = nc.active_aufgabe(self.config, _klasse) if _klasse else None

        note_app_snapshot = None
        note_data = fi.analysis.get("notenempfehlung")
        if note_data:
            raw_note = note_data.get("note")
            note_app_snapshot = float(raw_note) if raw_note is not None else None

        if ndb is not None:
            try:
                db_path = ndb.get_db_path(self.config)
                ndb.init_db(db_path)
                lf_id = ndb.upsert_lehrer_feedback(
                    db_path=db_path,
                    abgabe_id=abgabe_id,
                    klasse=_klasse or "",
                    aufgabe=_aufgabe or "",
                    note_final=note_final,
                    note_app_snapshot=note_app_snapshot,
                    lehrer_kommentar=kommentar,
                )
                # Cache invalidieren
                self._lf_status_cache.pop(str(fi.path), None)
                self._update_file_list()

                status_static = self.query_one("#lf-status", Static)
                status_static.update(f"✓ Eingetragen: Note {int(note_final)}")
                if kommentar:
                    first_line = kommentar.split("\n")[0][:40]
                    status_static.update(f"✓ Eingetragen: Note {int(note_final)} – {first_line}…")

                self.notify(f"Lehrer-Feedback gespeichert (ID: {lf_id}).")
            except Exception as e:
                self.notify(f"Speichern fehlgeschlagen: {e}", severity="error")
        else:
            self.notify("Datenbank-Modul nicht verfügbar.", severity="error")

    def on_button_pressed(self, event: Button.Pressed) -> None:
        """Aktionsbuttons im Middle Panel → gleiche Aktionen wie Tastenkürzel."""
        btn_id = event.button.id
        if btn_id == "lf-speichern":
            self._action_save_lehrer_feedback()
            return
        if btn_id == "btn-analyze":
            self.action_analyze_file()
        elif btn_id == "btn-review":
            self.action_review_file()
        elif btn_id == "btn-docx":
            self.action_generate_docx()
        elif btn_id == "btn-edit":
            self.action_edit_assignment()
        elif btn_id == "btn-reanalyze":
            fi = self._get_selected_file()
            if fi:
                fi.status = FileStatus.PENDING
                fi.analysis = None
                self._update_all_panels()
        elif btn_id == "add-class-btn":
            self._action_add_class()
        elif btn_id == "add-aufgabe-btn":
            self._action_add_aufgabe()
        elif btn_id == "attach-rubric-btn":
            self._action_attach_rubric()
        elif btn_id == "btn-open-folder":
            klasse = nc.active_klasse(self.config)
            aufgabe = nc.active_aufgabe(self.config, klasse) if klasse else None
            paths = nc.build_project_paths(self.config, klasse, aufgabe)
            paths.input_dir.mkdir(parents=True, exist_ok=True)
            nc.open_file(paths.input_dir)
            self.notify(f"Ordner: {paths.input_dir}")
        elif btn_id == "btn-open-ausgangstext":
            klasse = nc.active_klasse(self.config)
            aufgabe = nc.active_aufgabe(self.config, klasse) if klasse else None
            if not klasse or not aufgabe:
                self.notify("Keine Aufgabe ausgewählt.", severity="warning")
            else:
                paths = nc.build_project_paths(self.config, klasse, aufgabe)
                at_dir = paths.input_dir / "ausgangstext"
                at_dir.mkdir(parents=True, exist_ok=True)
                nc.open_file(at_dir)
                at_file = nc.detect_ausgangstext(self.config, klasse, aufgabe)
                if at_file:
                    self.notify(f"Ausgangstext: {at_file.name}")
                else:
                    self.notify("Ausgangstext-Ordner geöffnet — Datei hier ablegen (DOCX/PDF/Bild).")
        elif btn_id == "apply-aufgabe-all-btn":
            self.action_apply_aufgabe_to_all()
        elif btn_id in ("tab-text", "tab-rating", "tab-rubrik", "tab-output"):
            mode = btn_id.removeprefix("tab-")
            self.action_preview_tab(mode)
        elif btn_id == "btn-edit-rubric":
            self._action_edit_rubric()

    def _action_add_class(self) -> None:
        existing = nc.list_classes(self.config)

        def _after_add(result: bool) -> None:
            if result:
                self.config = nc.load_config()
                self._populate_class_select()
                self._populate_aufgabe_select()
                self._load_files()
                self._update_all_panels()
                new_active = nc.active_klasse(self.config)
                if new_active:
                    self.notify(f"Klasse '{new_active}' angelegt.", severity="information")

        self.push_screen(AddClassScreen(existing), _after_add)

    def _action_add_aufgabe(self) -> None:
        klasse = nc.active_klasse(self.config)
        if not klasse:
            self.notify("Erst eine Klasse wählen.", severity="warning")
            return
        existing_slugs = nc.list_aufgaben(self.config, klasse)

        def _after_add(result: bool) -> None:
            if result:
                self.config = nc.load_config()
                self._populate_aufgabe_select()
                self.files = []
                self._load_files()
                self._apply_defaults()
                self._update_all_panels()
                new_auf = nc.active_aufgabe(self.config, klasse)
                auf_cfg = nc.get_aufgabe_cfg(self.config, klasse, new_auf or "")
                label = auf_cfg.get("label", new_auf or "")
                if label:
                    self.notify(f"Aufgabe '{label}' angelegt.", severity="information")

        self.push_screen(AddAufgabeScreen(klasse, existing_slugs, self.config), _after_add)

    def _action_attach_rubric(self) -> None:
        klasse = nc.active_klasse(self.config)
        if not klasse:
            self.notify("Erst eine Klasse wählen.", severity="warning")
            return
        aufgabe = nc.active_aufgabe(self.config, klasse)
        if not aufgabe:
            self.notify("Erst eine Aufgabe wählen.", severity="warning")
            return
        auf_cfg = nc.get_aufgabe_cfg(self.config, klasse, aufgabe)
        label = auf_cfg.get("label", aufgabe)

        def _after(result: bool) -> None:
            if result:
                self.config = nc.load_config()
                self._update_preview_panel()

        self.push_screen(AttachRubricScreen(klasse, aufgabe, label, self.config), _after)

    def on_select_changed(self, event: Select.Changed) -> None:
        """Klassen- oder Aufgaben-Wechsel: Dateiliste neu laden."""
        if self._switching_context:
            return
        self._switching_context = True
        try:
            if event.select.id == "class-select" and event.value is not Select.BLANK:
                new_klasse = str(event.value)
                self.config.setdefault("classes", {})["active"] = new_klasse
                try:
                    nc.save_active_klasse(new_klasse)
                except Exception:
                    logging.warning("Aktive Klasse konnte nicht gespeichert werden", exc_info=True)
                self._cancel_event.set()
                self.selected_index = 0
                self._populate_aufgabe_select()
                self.files = []
                try:
                    self._load_files()
                    self._apply_defaults()
                    self._update_all_panels()
                except Exception as exc:
                    self.notify(f"Fehler beim Laden: {exc}", severity="error")
                self.notify(f"Klasse: {new_klasse}")

            elif event.select.id == "aufgabe-select" and event.value is not Select.BLANK:
                new_aufgabe = str(event.value)
                klasse = nc.active_klasse(self.config)
                if klasse:
                    self.config.setdefault("classes", {}).setdefault(klasse, {})[
                        "active_aufgabe"
                    ] = new_aufgabe
                    try:
                        nc.save_active_aufgabe(klasse, new_aufgabe)
                    except Exception:
                        logging.warning(
                            "Aktive Aufgabe konnte nicht gespeichert werden", exc_info=True
                        )
                self._cancel_event.set()
                self.selected_index = 0
                self.files = []
                try:
                    self._load_files()
                    self._apply_defaults()
                    self._update_all_panels()
                except Exception as exc:
                    self.notify(f"Fehler beim Laden: {exc}", severity="error")
                auf_cfg = nc.get_aufgabe_cfg(self.config, klasse or "", new_aufgabe)
                label = auf_cfg.get("label", new_aufgabe)
                self.notify(f"Aufgabe: {label}")
        finally:
            self._switching_context = False

    def _update_preview_panel(self) -> None:
        content = self.query_one("#preview-content", Static)
        output_list = self.query_one("#output-list", ListView)

        # Toggle content vs output-list visibility
        is_output_tab = self.preview_mode == "output"
        is_rubrik_tab = self.preview_mode == "rubrik"
        content.display = not is_output_tab
        output_list.display = is_output_tab

        edit_rubric_btn = self.query_one("#btn-edit-rubric", Button)
        edit_rubric_btn.display = is_rubrik_tab

        # Aktiven Tab-Button hervorheben
        tab_btn_ids = {
            "text": "tab-text",
            "rating": "tab-rating",
            "rubrik": "tab-rubrik",
            "output": "tab-output",
        }
        for mode, btn_id in tab_btn_ids.items():
            btn = self.query_one(f"#{btn_id}", Button)
            if mode == self.preview_mode:
                btn.add_class("tab-active")
            else:
                btn.remove_class("tab-active")

        if is_output_tab:
            self._populate_output_list(output_list)
            return

        if not self._filtered_indices:
            content.update("")
            return

        idx = (
            self._filtered_indices[self.selected_index]
            if self.selected_index < len(self._filtered_indices)
            else None
        )
        if idx is None or idx >= len(self.files):
            content.update("")
            return

        fi = self.files[idx]

        if self.preview_mode == "text":
            if fi.file_type == "docx":
                try:
                    content.update(nc.read_docx_rich(fi.path))
                except Exception:
                    content.update("(Datei konnte nicht gelesen werden)")
            else:
                typ_label = "PDF" if fi.file_type == "pdf" else "Bilddatei"
                content.update(
                    f"[{typ_label}: {fi.path.name}]\n"
                    f"Größe: {fi.word_count} KB\n\n"
                    f"Handgeschriebene Schülerarbeit — Inhalt wird durch\n"
                    f"Vision-Modell bei der Analyse gelesen.\n"
                    f"Text-Vorschau im TUI nicht verfügbar."
                )
        elif self.preview_mode == "rubrik":
            try:
                klasse = nc.active_klasse(self.config)
                aufgabe = nc.active_aufgabe(self.config, klasse) if klasse else None
                if fi.rubric:
                    rubric_content = nc.load_rubric(fi.rubric, self.config)
                else:
                    rubric_content = nc.load_rubric_for_aufgabe(self.config, klasse, aufgabe)
                content.update(
                    rubric_content[:1200] + ("..." if len(rubric_content) > 1200 else "")
                )
            except Exception:
                content.update("(Rubrik konnte nicht geladen werden)")
        else:
            if fi.analysis:
                lines = []
                bewertung = fi.analysis.get("bewertung", {})
                for key, crit in bewertung.items():
                    if not isinstance(crit, dict):
                        continue
                    lines.append(f"── {key.replace('_', ' ').title()} ──")
                    for s in crit.get("staerken", []):
                        lines.append(f"  + {s}")
                    for s in crit.get("schwaechen", []):
                        lines.append(f"  - {s}")
                    lines.append("")
                content.update("\n".join(lines))
            else:
                content.update("(Keine Analyse vorhanden)")

    def _get_selected_file(self) -> FileInfo | None:
        if not self._filtered_indices or self.selected_index >= len(self._filtered_indices):
            return None
        idx = self._filtered_indices[self.selected_index]
        if idx >= len(self.files):
            return None
        return self.files[idx]

    def _populate_output_list(self, output_list: ListView) -> None:
        """Befüllt die Output-ListView mit DOCX-Dateien aus dem output/-Ordner.

        Verwendet In-Place-Update (kein clear()) um DuplicateIds zu vermeiden —
        dasselbe Muster wie _update_file_list. IDs sind Hash-basiert damit
        dieselbe Datei immer dieselbe ID bekommt und nie doppelt eingefügt wird.
        """
        output_dir = nc.build_project_paths(self.config).output_dir

        # Neue Items aufbauen: id → (path, anzeigetext)
        new_items: dict[str, tuple[Path, str]] = {}
        if output_dir.exists():
            for p in sorted(
                output_dir.glob("*.docx"),
                key=lambda f: f.stat().st_mtime,
                reverse=True,
            ):
                try:
                    st = p.stat()
                    mtime = time.strftime("%d.%m.%y %H:%M", time.localtime(st.st_mtime))
                    size_kb = st.st_size // 1024
                    label_text = f"{mtime}  {p.name}  ({size_kb} KB)"
                    item_id = "out-" + hashlib.md5(p.name.encode()).hexdigest()[:8]
                    new_items[item_id] = (p, label_text)
                except OSError:
                    continue

        # _output_paths aktualisieren
        self._output_paths = {k: v[0] for k, v in new_items.items()}

        # Existierende IDs im ListView ermitteln
        existing_ids = {item.id for item in output_list.query(ListItem) if item.id}

        # Stale Items entfernen (ohne clear())
        for item in list(output_list.query(ListItem)):
            if item.id and item.id not in new_items and item.id != "out-empty":
                item.remove()

        if not new_items:
            if "out-empty" not in existing_ids:
                output_list.append(ListItem(Label("Noch keine DOCX erstellt."), id="out-empty"))
            return

        # Platzhalter entfernen falls vorhanden
        if "out-empty" in existing_ids:
            try:
                output_list.query_one("#out-empty", ListItem).remove()
            except Exception:
                logging.debug("Platzhalter 'out-empty' konnte nicht entfernt werden", exc_info=True)

        # Nur wirklich neue Items anhängen (keine DuplicateIds möglich)
        for item_id, (_path, label_text) in new_items.items():
            if item_id not in existing_ids:
                output_list.append(ListItem(Label(label_text), id=item_id))

    def on_list_view_selected(self, event: ListView.Selected) -> None:
        item = event.item
        # Handle output-list selections
        if item.id and item.id in self._output_paths:
            p = self._output_paths[item.id]
            if p.exists():
                nc.open_file(p)
                self.notify(f"Öffne: {p.name}")
            return
        if not (item.id and item.id in self._id_to_index):
            return
        real_idx = self._id_to_index[item.id]
        if real_idx not in self._filtered_indices:
            return

        now = time.time()
        is_double_click = item.id == self._last_click_id and now - self._last_click_time < 0.5
        self._last_click_id = item.id
        self._last_click_time = now

        self.selected_index = self._filtered_indices.index(real_idx)
        self._update_middle_panel()
        self._update_preview_panel()

        if is_double_click:
            if real_idx >= len(self.files):
                return
            fi = self.files[real_idx]
            if fi.status in (FileStatus.ANALYZED, FileStatus.DONE):
                self.action_review_file()
            elif fi.status == FileStatus.PENDING:
                self.action_analyze_file()

    def on_list_view_highlighted(self, event: ListView.Highlighted) -> None:
        if self._rebuilding_list or event.item is None:
            return
        if event.item.id and event.item.id in self._id_to_index:
            real_idx = self._id_to_index[event.item.id]
            if real_idx in self._filtered_indices:
                self.selected_index = self._filtered_indices.index(real_idx)
                self._update_middle_panel()
                self._update_preview_panel()

    def on_input_changed(self, event: Input.Changed) -> None:
        if event.input.id == "search-input":
            self._search_filter = event.value
            self._apply_filter()
            self.selected_index = 0
            self._update_file_list()
            self._update_counter()

    def action_toggle_search(self) -> None:
        search_input = self.query_one("#search-input", Input)
        self.search_active = not self.search_active
        search_input.display = self.search_active
        if self.search_active:
            search_input.focus()
        else:
            self._search_filter = ""
            search_input.value = ""
            self._apply_filter()
            self._update_file_list()
            self._update_counter()

    def action_next_panel(self) -> None:
        panels = ["#files-panel", "#middle-panel", "#preview-panel"]
        for p in panels:
            try:
                self.query_one(p).remove_class("panel-focused")
            except Exception:
                logging.debug(
                    "Panel-Fokus-Klasse konnte nicht entfernt werden (%s)", p, exc_info=True
                )
        self._focus_panel = (self._focus_panel + 1) % 3
        try:
            panel = self.query_one(panels[self._focus_panel])
            panel.add_class("panel-focused")
            panel.focus()
        except Exception:
            logging.debug("Panel-Fokus konnte nicht gesetzt werden", exc_info=True)

    def action_show_help(self) -> None:
        self.push_screen(HelpScreen())

    def action_show_settings(self) -> None:
        self.push_screen(SettingsScreen(self.config))

    def action_show_schueler(self) -> None:
        self.push_screen(SchuelerVerwaltungScreen(self.config))

    def action_show_klassen_feedback(self) -> None:
        klasse = nc.active_klasse(self.config)
        aufgabe = nc.active_aufgabe(self.config, klasse) if klasse else None
        self.push_screen(KlassenFeedbackScreen(self.config, aufgabe=aufgabe or ""))

    def action_retro_import(self) -> None:
        self.push_screen(RetroImportScreen(self.config))

    def action_generate_erwartungshorizont(self) -> None:
        klasse = nc.active_klasse(self.config)
        aufgabe = nc.active_aufgabe(self.config, klasse) if klasse else None
        if not klasse or not aufgabe:
            self.notify("Keine aktive Klasse/Aufgabe — zuerst Aufgabe anlegen.", severity="warning")
            return

        def _after(result: bool) -> None:
            if result:
                self._update_all_panels()

        self.push_screen(
            ErwartungshorizontGeneratorScreen(self.config, klasse, aufgabe), _after
        )

    def action_toggle_sort(self) -> None:
        modes = ["name", "status", "words"]
        current = modes.index(self.sort_mode) if self.sort_mode in modes else 0
        self.sort_mode = modes[(current + 1) % len(modes)]
        self._apply_filter()
        self._update_file_list()
        self._update_counter()
        self.notify(f"Sortierung: {self.sort_mode}")

    def action_edit_assignment(self) -> None:
        marked = [fi for fi in self.files if fi.marked]
        if marked:
            self._edit_assignment_sequence(marked, 0)
        else:
            fi = self._get_selected_file()
            if fi:
                self._open_edit_dialog(fi, 1, 1)

    def _open_edit_dialog(self, fi: FileInfo, idx: int, total: int) -> None:
        """Füllt Vornamen vor (wenn leer) und öffnet den Edit-Dialog."""
        if not fi.schueler:
            suggested = (
                nc.extract_schueler_name(fi.path)
                if fi.file_type == "docx"
                else nc.extract_vorname_from_filename(fi.path.name)
            )
            if suggested:
                fi.schueler = suggested

        def _after(result: bool) -> None:
            if result:
                self._update_all_panels()
                if total == 1:
                    self.notify(f"Zuordnung gespeichert: {fi.path.name}")

        self.push_screen(EditAssignmentScreen(fi, self.config, idx, total), _after)

    def _edit_assignment_sequence(self, targets: list[FileInfo], idx: int) -> None:
        """Öffnet den Edit-Dialog für targets[idx], danach rekursiv für idx+1."""
        if idx >= len(targets):
            self.notify(f"Zuordnung abgeschlossen: {len(targets)} Dateien bearbeitet.")
            self._update_all_panels()
            return

        fi = targets[idx]
        total = len(targets)

        if not fi.schueler:
            suggested = nc.extract_vorname_from_filename(fi.path.name)
            if suggested:
                fi.schueler = suggested

        def _after(result: bool) -> None:
            if result:
                self._update_all_panels()
            self._edit_assignment_sequence(targets, idx + 1)

        self.push_screen(
            EditAssignmentScreen(fi, self.config, idx + 1, total), _after
        )

    def action_review_file(self) -> None:
        fi = self._get_selected_file()
        if not fi or not fi.analysis:
            self.notify("Keine Analyse vorhanden.", severity="warning")
            return
        self.push_screen(ReviewScreen(fi, self.config))
        self._update_all_panels()

    def action_analyze_file(self) -> None:
        marked = [fi for fi in self.files if fi.marked]
        if marked:
            self._run_analysis(marked)
        else:
            fi = self._get_selected_file()
            if fi:
                self._run_analysis([fi])

    def action_analyze_marked(self) -> None:
        marked = [fi for fi in self.files if fi.marked]
        if not marked:
            fi = self._get_selected_file()
            if fi:
                marked = [fi]
        if marked:
            self._run_analysis(marked)

    @work(thread=True)
    def _run_analysis(self, targets: list[FileInfo]) -> None:
        """Fuehrt die LLM-Analyse mit robuster Retry-Logik aus."""
        if not nc.api_key_available(self.config):
            provider = self.config.get("api", {}).get("provider", "anthropic")
            self.call_from_thread(
                self.notify,
                f"API-Key fuer Provider '{provider}' nicht gesetzt. .env pruefen.",
                severity="error",
            )
            return

        # Vision-Capability prüfen: jede Vision-Datei gegen Provider+Modell testen
        provider = self.config.get("api", {}).get("provider", "anthropic")
        model = self.config.get("api", {}).get("model", "")
        for fi in targets:
            if nc.is_vision_file(fi.path) and not nc.is_vision_capable(provider, model, fi.path):
                suffix = fi.path.suffix.upper().lstrip(".")
                pdf_hint = " (OpenAI unterstützt keine PDFs — bitte Anthropic wählen)" if suffix == "PDF" else ""
                self.call_from_thread(
                    self.notify,
                    f"'{model}' ({provider}) kann {suffix}-Dateien nicht analysieren.{pdf_hint}",
                    severity="error",
                )
                return

        self._cancel_event.clear()
        progress_screen = ProgressScreen()
        self.call_from_thread(self.push_screen, progress_screen)

        total = len(targets)
        for i, fi in enumerate(targets):
            if self._cancel_event.is_set():
                break

            fi.status = FileStatus.PROGRESS
            self.call_from_thread(self._update_file_list)

            queue = [t.path.name for t in targets[i + 1 :]]
            if progress_screen in self.screen_stack:
                self.call_from_thread(
                    progress_screen.update_progress,
                    fi.path.name,
                    "Analyse läuft...",
                    queue,
                    i,
                    total,
                )

            try:
                if fi.file_type == "docx":
                    if fi.path.suffix.lower() == ".odt":
                        docx_text = nc.read_odt_text(fi.path)
                    else:
                        docx_text = nc.read_docx_text(fi.path)
                    file_path_for_vision = None
                else:
                    docx_text = ""
                    file_path_for_vision = fi.path
                rubric_content = nc.load_rubric(fi.rubric, self.config)

                # Ausgangstext erkennen (subfolder ausgangstext/ im Aufgabenordner)
                _klasse = nc.active_klasse(self.config)
                _aufgabe = nc.active_aufgabe(self.config, _klasse) if _klasse else None
                ausgangstext_path = (
                    nc.detect_ausgangstext(self.config, _klasse, _aufgabe)
                    if _klasse and _aufgabe else None
                )
                # Frühe Warnung wenn Ausgangstext-PDF mit OpenAI nicht unterstützt (non-fatal)
                if (ausgangstext_path and nc.is_vision_file(ausgangstext_path)
                        and not nc.is_vision_capable(provider, model, ausgangstext_path)):
                    self.call_from_thread(
                        self.notify,
                        f"Ausgangstext PDF wird mit '{provider}' übersprungen — kein PDF-Support.",
                        severity="warning",
                        timeout=6,
                    )

                data, errors = nc.run_llm_analysis(
                    docx_text,
                    rubric_content,
                    fi.fach,
                    fi.schulstufe,
                    fi.textsorte,
                    self.config,
                    schueler=fi.schueler,
                    cancel_event=self._cancel_event,
                    file_path=file_path_for_vision,
                    bewertungsmodus=fi.bewertungsmodus,
                    ausgangstext_path=ausgangstext_path,
                    klasse=_klasse or "",
                    aufgabe=_aufgabe or "",
                    erwartungshorizont_name=fi.erwartungshorizont,
                    rubric_name=fi.rubric,
                )

                if data is not None:
                    data["datei"] = fi.path.name  # LLM kennt den Dateinamen nicht zuverlässig
                    fi.analysis = data
                    fi.status = FileStatus.ANALYZED
                    self._save_analysis(fi)
                    if errors:  # non-fatale Warnungen loggen
                        paths = nc.build_project_paths(self.config)
                        for warn in errors:
                            nc.log_tui_error(paths, f"WARNUNG {fi.path.name}: {warn}")
                elif errors:
                    last_error = errors[-1]
                    if "abgebrochen" in last_error.lower():
                        fi.status = FileStatus.PENDING
                    else:
                        fi.status = FileStatus.ERROR
                        paths = nc.build_project_paths(self.config)
                        for err in errors:
                            nc.log_tui_error(paths, f"{fi.path.name}: {err}")
                else:
                    fi.status = FileStatus.ERROR

            except Exception as e:
                if self._cancel_event.is_set():
                    fi.status = FileStatus.PENDING
                else:
                    fi.status = FileStatus.ERROR
                    paths = nc.build_project_paths(self.config)
                    nc.log_tui_error(paths, f"{fi.path.name}: {e}")

            status_text = (
                "✓ Fertig"
                if fi.status == FileStatus.ANALYZED
                else "⏹ Abgebrochen"
                if fi.status == FileStatus.PENDING
                else "✗ Fehler"
            )
            if progress_screen in self.screen_stack:
                self.call_from_thread(
                    progress_screen.update_progress,
                    fi.path.name,
                    status_text,
                    [],
                    i + 1,
                    total,
                )

        if progress_screen in self.screen_stack:
            self.call_from_thread(progress_screen.dismiss)

        self.call_from_thread(self._update_all_panels)
        done_count = sum(1 for f in targets if f.status == FileStatus.ANALYZED)
        self.call_from_thread(
            self.notify,
            f"Analyse abgeschlossen: {done_count}/{total} erfolgreich",
        )

        # Pruefen, ob alle Dateien der aktuellen Aufgabe jetzt analysiert sind
        if done_count == total and done_count > 0:
            pending_in_aufgabe = sum(
                1 for f in self.files if f.status == FileStatus.PENDING
            )
            if pending_in_aufgabe == 0:
                self.call_from_thread(
                    self.notify,
                    "Alle Schüler analysiert! Klassen-Feedback verfügbar (Button: 📊 Feedback).",
                    severity="information",
                    timeout=8,
                )

    def _save_analysis(self, fi: FileInfo) -> None:
        if not fi.analysis:
            return
        paths = nc.build_project_paths(self.config)
        paths.feedback_data_dir.mkdir(parents=True, exist_ok=True)
        if fi.schueler:
            fi.analysis["schueler"] = fi.schueler
        fi.analysis["fach"] = fi.fach
        fi.analysis["schulstufe"] = fi.schulstufe
        fi.analysis["textsorte"] = fi.textsorte
        fi.analysis["rubrik"] = fi.rubric
        fi.analysis["erwartungshorizont"] = fi.erwartungshorizont
        fi.analysis["bewertungsmodus"] = fi.bewertungsmodus
        api_cfg = self.config.get("api", {})
        fi.analysis["modell"] = api_cfg.get("model", "")
        fi.analysis["provider"] = api_cfg.get("provider", "anthropic")
        out_name = fi.path.stem + "_analysis.json"
        analysis_path = paths.feedback_data_dir / out_name
        nc.archive_existing_analysis(analysis_path)
        analysis_path.write_text(
            json.dumps(fi.analysis, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )

    def action_generate_docx(self) -> None:
        marked = [fi for fi in self.files if fi.marked]
        if marked:
            self._generate_docx_files(marked)
        else:
            fi = self._get_selected_file()
            if not fi or not fi.analysis:
                self.notify("Keine Analyse vorhanden.", severity="warning")
                return
            self._generate_docx_files([fi])

    def action_generate_docx_marked(self) -> None:
        marked = [fi for fi in self.files if fi.marked]
        if not marked:
            marked = [fi for fi in self.files if fi.analysis]
        if not marked:
            fi = self._get_selected_file()
            if fi:
                marked = [fi]
        if marked:
            self._generate_docx_files(marked)

    def _generate_docx_files(self, targets: list[FileInfo]) -> None:
        paths = nc.build_project_paths(self.config)
        paths.output_dir.mkdir(parents=True, exist_ok=True)
        _klasse = nc.active_klasse(self.config)
        _aufgabe = nc.active_aufgabe(self.config, _klasse) if _klasse else None
        success = 0
        last_out_path = None
        for fi in targets:
            if not fi.analysis:
                continue
            try:
                feedback = gf.parse_feedback_data(fi.analysis)
                out_name = gf.output_filename(feedback.datei)
                out_path = paths.output_dir / out_name
                if fi.file_type == "docx" and fi.path.exists():
                    if fi.path.suffix.lower() == ".odt":
                        original_text = nc.read_odt_text(fi.path)
                    else:
                        original_text = nc.read_docx_text(fi.path)
                    source_note = None
                else:
                    transkription = _load_transkription(fi, paths)
                    original_text = transkription if transkription else None
                    source_note = None if transkription else f"Analyse basiert auf eingescannter Datei: {fi.path.name}"
                _auf_label_b = (
                    nc.get_aufgabe_cfg(self.config, _klasse or "", _aufgabe or "").get("label", "")
                    if _aufgabe
                    else ""
                )
                doc = gf.build_feedback_document(
                    feedback, config=self.config, original_text=original_text,
                    source_note=source_note, aufgabe_label=_auf_label_b,
                    bewertungsmodus=fi.bewertungsmodus,
                )
                doc.save(str(out_path))
                fi.status = FileStatus.DONE
                success += 1
                last_out_path = out_path
            except Exception as e:
                fi.status = FileStatus.ERROR
                nc.log_tui_error(paths, f"{fi.path.name}: {e}")

        self._update_all_panels()
        self.notify(f"DOCX: {success}/{len(targets)} generiert")
        if success == 1 and last_out_path:
            nc.open_file(last_out_path)

    def action_apply_aufgabe_to_all(self) -> None:
        if not self.files:
            self.notify("Keine Dateien geladen.", severity="warning")
            return
        klasse = nc.active_klasse(self.config)
        aufgabe = nc.active_aufgabe(self.config, klasse) if klasse else None
        auf_defs = nc.aufgabe_defaults(self.config, klasse, aufgabe) if aufgabe else {}
        global_defs = self.config.get("defaults", {})
        fach = auf_defs.get("fach") or global_defs.get("fach", "Deutsch")
        schulstufe = auf_defs.get("schulstufe") or global_defs.get("schulstufe", "Oberstufe")
        textsorte = auf_defs.get("textsorte", "")
        rubric = auf_defs.get("rubric") or nc.default_rubric_for(fach, schulstufe, self.config)
        bewertungsmodus = auf_defs.get("bewertungsmodus", "benotet")
        eh = auf_defs.get("erwartungshorizont", "")
        for fi in self.files:
            fi.fach = fach
            fi.schulstufe = schulstufe
            if textsorte:
                fi.textsorte = textsorte
            if rubric:
                fi.rubric = rubric
            fi.bewertungsmodus = bewertungsmodus  # type: ignore[assignment]
            if eh:
                fi.erwartungshorizont = eh
        self._update_all_panels()
        self.notify(f"Aufgabe auf {len(self.files)} Dateien angewendet: {fach} / {schulstufe}")

    def action_open_folder(self) -> None:
        """Öffnet den aktuellen Input-Ordner im Betriebssystem-Dateimanager."""
        input_dir = nc.build_project_paths(self.config).input_dir
        input_dir.mkdir(parents=True, exist_ok=True)
        nc.open_file(input_dir)
        self.notify(f"Ordner geöffnet: {input_dir.name}")

    def action_quit(self) -> None:
        def _handle_confirm(result: bool) -> None:
            if result:
                self.exit()

        self.push_screen(ConfirmScreen("NATASCHA beenden?"), _handle_confirm)

    def _check_first_run(self) -> None:
        marker = nc.PROJECT_ROOT / ".natascha_first_run_done"
        if marker.exists():
            return
        provider = self.config.get("api", {}).get("provider", "anthropic")
        self.notify(
            f"API-Key für Provider '{provider}' nicht gefunden. .env prüfen.",
            severity="warning",
        )
        try:
            marker.touch()
        except OSError:
            pass

    def key_space(self) -> None:
        if self.search_active:
            return
        fi = self._get_selected_file()
        if fi:
            fi.marked = not fi.marked
            self._update_file_list()
            self._update_counter()
            lv = self.query_one("#file-list", ListView)
            if lv.index is not None and lv.index < len(self._filtered_indices) - 1:
                lv.index = lv.index + 1

    def action_select_all(self) -> None:
        for fi in self.files:
            fi.marked = True
        self._update_file_list()
        self._update_counter()

    def action_deselect_all(self) -> None:
        for fi in self.files:
            fi.marked = False
        self._update_file_list()
        self._update_counter()

    def key_enter(self) -> None:
        modes = ["text", "rating", "rubrik", "output"]
        cur = self.preview_mode if self.preview_mode in modes else "text"
        self.preview_mode = modes[(modes.index(cur) + 1) % len(modes)]
        self._update_preview_panel()

    def action_preview_tab(self, mode: str) -> None:
        self.preview_mode = mode
        self._update_preview_panel()
        names = {"text": "Text", "rating": "Bewertung", "rubrik": "Rubrik", "output": "Output"}
        self.notify(f"Vorschau: {names.get(mode, mode)}")

    def _action_edit_rubric(self) -> None:
        fi = self._get_selected_file()
        rubric_name = ""
        if fi and fi.rubric:
            rubric_name = fi.rubric
        else:
            klasse = nc.active_klasse(self.config)
            aufgabe = nc.active_aufgabe(self.config, klasse) if klasse else None
            auf_cfg = nc.get_aufgabe_cfg(self.config, klasse or "", aufgabe or "")
            rubric_name = auf_cfg.get("rubric", "")
        if not rubric_name:
            self.notify("Keine Rubrik zugeordnet.", severity="warning")
            return

        def _after_edit(result: bool) -> None:
            self._update_preview_panel()

        self.push_screen(RubrikEditorScreen(rubric_name, self.config), _after_edit)

    def key_delete(self) -> None:
        fi = self._get_selected_file()
        if not fi:
            return

        def _handle(result: bool) -> None:
            if result:
                self.files.remove(fi)
                self._rebuild_id_map()
                self._apply_filter()
                self._update_all_panels()

        self.push_screen(ConfirmScreen(f"{fi.path.name} entfernen?"), _handle)


def main() -> None:
    app = NataschaApp()
    app.run()


if __name__ == "__main__":
    main()
