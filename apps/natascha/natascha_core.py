#!/usr/bin/env python3
"""
NATASCHA Core – Gemeinsame Logik-Funktionen fuer Wizard und Dashboard.
"""

from __future__ import annotations

import base64
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import threading
import time
import tomllib
import urllib.error
import urllib.request
from datetime import datetime
from pathlib import Path
from typing import Any

import tomlkit

if sys.version_info < (3, 11):
    print("Python 3.11+ wird benoetigt.")
    raise SystemExit(1)

try:
    from docx import Document as DocxDocument
except ImportError:
    print("python-docx fehlt: pip install python-docx")
    raise SystemExit(1)

try:
    import jsonschema as _jsmod
except ImportError:
    _jsmod = None

sys.path.insert(0, str(Path(__file__).resolve().parent))
import generate_feedback as gf

# Schueler-Tracking (optional – falls natascha_db.py fehlt, wird es uebersprungen)
try:
    import natascha_db as ndb
except Exception:
    ndb = None  # type: ignore[assignment]

# Pseudonymisierung vor LLM-Versand (optional wie ndb; ohne DB kein Roster)
try:
    import pseudonymisierung as pseu
except Exception:
    pseu = None  # type: ignore[assignment]

# Dateien, die beim ersten Start eines gebündelten Sidecars ins persistente
# Datenverzeichnis kopiert werden. Einzeldateien nur, wenn sie fehlen —
# Nutzeränderungen (eigene Rubriken, angepasste Config) werden NIE überschrieben;
# neue mitgelieferte Dateien kommen bei App-Updates trotzdem an.
_SEED_DATEIEN = ("natascha_config.toml", "feedback_schema.json")
_SEED_ORDNER = ("rubrics", "prompts")


def _seed_data_dir(bundle_root: Path, data_root: Path) -> None:
    """Kopiert mitgelieferte Ressourcen einmalig ins Datenverzeichnis (idempotent)."""
    import shutil as _shutil

    data_root.mkdir(parents=True, exist_ok=True)
    for name in _SEED_DATEIEN:
        quelle = bundle_root / name
        ziel = data_root / name
        if quelle.is_file() and not ziel.exists():
            _shutil.copy2(quelle, ziel)
    for ordner in _SEED_ORDNER:
        quell_dir = bundle_root / ordner
        if not quell_dir.is_dir():
            continue
        ziel_dir = data_root / ordner
        ziel_dir.mkdir(parents=True, exist_ok=True)
        for quelle in quell_dir.rglob("*"):
            rel = quelle.relative_to(quell_dir)
            ziel = ziel_dir / rel
            if quelle.is_dir():
                ziel.mkdir(parents=True, exist_ok=True)
            elif not ziel.exists():
                _shutil.copy2(quelle, ziel)


def _resolve_project_root() -> Path:
    """Projekt-Wurzel für Config, Rubriken und alle Schreibpfade.

    Dev-Modus: das echte Repo-Verzeichnis (unverändertes Verhalten).
    PyInstaller-Onefile: __file__ zeigt ins flüchtige _MEIPASS-Temp-Verzeichnis,
    das bei jedem Start neu entpackt wird — jeder Schreibvorgang dorthin wäre
    nach Prozessende verloren (Config-Änderungen, gespeicherte Rubriken,
    Analyse-JSONs). Deshalb dient im Bundle-Modus ein persistentes
    Datenverzeichnis neben der gemeinsamen DB als Wurzel; die mitgelieferten
    Ressourcen werden beim Start dorthin geseedet.
    """
    source_root = Path(__file__).resolve().parent
    if not getattr(sys, "frozen", False):
        return source_root
    bundle_root = Path(getattr(sys, "_MEIPASS", source_root))
    data_root = Path.home() / "lehr-suite-bridge" / "natascha"
    try:
        _seed_data_dir(bundle_root, data_root)
    except OSError:
        # Datenverzeichnis nicht beschreibbar → wenigstens lesend lauffähig bleiben.
        return bundle_root
    return data_root


PROJECT_ROOT = _resolve_project_root()

VERSION = "0.6.0"


def _load_dotenv() -> None:
    """Lädt .env aus dem Projektverzeichnis — überschreibt keine bereits gesetzten Variablen."""
    env_path = PROJECT_ROOT / ".env"
    if not env_path.exists():
        return
    try:
        from dotenv import load_dotenv

        load_dotenv(env_path, override=False)
        return
    except ImportError:
        pass
    # Fallback: manuelles Parsen ohne externe Abhängigkeit
    with env_path.open(encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, _, value = line.partition("=")
            key = key.strip()
            value = value.strip().strip('"').strip("'")
            if key and key not in os.environ:
                os.environ[key] = value


def load_config() -> dict[str, Any]:
    _load_dotenv()
    config_path = PROJECT_ROOT / "natascha_config.toml"
    if not config_path.exists():
        raise FileNotFoundError(f"Konfigurationsdatei nicht gefunden: {config_path}")
    with config_path.open("rb") as f:
        return tomllib.load(f)


def _load_toml_doc() -> tomlkit.TOMLDocument:
    """Lädt natascha_config.toml als bearbeitbares tomlkit-Dokument (Kommentare bleiben erhalten)."""
    return tomlkit.parse((PROJECT_ROOT / "natascha_config.toml").read_text(encoding="utf-8"))


def _save_toml_doc(doc: tomlkit.TOMLDocument) -> None:
    """Schreibt ein tomlkit-Dokument zurück nach natascha_config.toml."""
    (PROJECT_ROOT / "natascha_config.toml").write_text(tomlkit.dumps(doc), encoding="utf-8")


def resolve_path(config: dict[str, Any], key: str) -> Path:
    return PROJECT_ROOT / config["paths"][key]


def count_words(path: Path) -> int:
    try:
        if path.suffix.lower() == ".odt":
            text = read_odt_text(path)
            return len(text.split())
        doc = DocxDocument(str(path))
        return sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
    except Exception:
        return 0


def read_docx_text(docx_path: Path, preserve_italic: bool = True) -> str:
    doc = DocxDocument(str(docx_path))
    parts = []
    for para in doc.paragraphs:
        runs_text = []
        for run in para.runs:
            text = run.text
            if preserve_italic and run.italic and text.strip():
                text = f"*{text}*"
            runs_text.append(text)
        parts.append("".join(runs_text))
    return "\n".join(parts)


def read_odt_text(path: Path) -> str:
    """Extrahiert den Text aus einer ODT-Datei (OpenDocument Text)."""
    import zipfile
    from xml.etree import ElementTree as ET
    with zipfile.ZipFile(path) as z:
        with z.open("content.xml") as f:
            tree = ET.parse(f)
    ns = {"text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0"}
    paragraphs = tree.findall(".//text:p", ns)
    return "\n".join("".join(el.itertext()) for el in paragraphs).strip()


def read_docx_rich(docx_path: Path) -> str:
    """Return the full document text with basic Rich markup for terminal display.

    Paragraph separation is the top priority; headings get bold+underline,
    inline bold/italic are also rendered.  Empty paragraphs are collapsed so
    that a single blank line always separates paragraphs.
    """
    doc = DocxDocument(str(docx_path))
    lines: list[str] = []
    prev_empty = False

    for para in doc.paragraphs:
        raw = para.text

        # Completely empty paragraph → use as separator (max one blank line)
        if not raw.strip():
            if not prev_empty and lines:
                lines.append("")
            prev_empty = True
            continue
        prev_empty = False

        # Detect heading style
        style_name = (para.style.name or "").lower()
        is_heading = style_name.startswith("heading") or style_name.startswith("überschrift")

        # Build the line with per-run inline markup
        parts: list[str] = []
        for run in para.runs:
            if not run.text:
                continue
            # Escape Rich markup characters in the raw text
            text = run.text.replace("[", r"\[")
            if run.bold and run.italic:
                text = f"[bold italic]{text}[/bold italic]"
            elif run.bold:
                text = f"[bold]{text}[/bold]"
            elif run.italic:
                text = f"[italic]{text}[/italic]"
            parts.append(text)

        line = "".join(parts) if parts else raw.replace("[", r"\[")

        if is_heading:
            line = f"[bold underline]{line}[/bold underline]"

        lines.append(line)

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Vision-Datei-Unterstützung (PDF, JPEG, PNG)
# ---------------------------------------------------------------------------

_VISION_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".jpg", ".jpeg", ".png"})
_OPENAI_VISION_MODELS: frozenset[str] = frozenset({
    "gpt-4o",
    "gpt-4o-mini",
    "gpt-4.1",
    "gpt-4.1-mini",
    "gpt-4-turbo",
    "gpt-5",
    "gpt-5.1",
    "gpt-5.2",
})


def is_vision_file(path: Path) -> bool:
    """True wenn der Dateityp Vision-Analyse erfordert (PDF, JPEG, PNG)."""
    return path.suffix.lower() in _VISION_EXTENSIONS


def get_file_size_kb(path: Path) -> int:
    try:
        return path.stat().st_size // 1024
    except OSError:
        return 0


def is_vision_capable(provider: str, model: str, file_path: Path | None = None) -> bool:
    """True wenn Provider+Modell die gegebene Datei per Vision analysieren kann.

    file_path: Wenn angegeben, wird auch der Dateityp geprüft.
    OpenAI unterstützt Bilder, aber keine nativen PDFs.
    """
    p = provider.lower()
    if p == "anthropic":
        return True  # alle aktuellen Claude-Modelle unterstützen Vision + PDF
    if p == "openai":
        if file_path is not None and file_path.suffix.lower() == ".pdf":
            return False  # OpenAI hat keinen nativen PDF-Support
        return model in _OPENAI_VISION_MODELS
    if p == "qwen":
        if file_path is not None and file_path.suffix.lower() == ".pdf":
            return False
        m = model.lower()
        return "vl" in m or "omni" in m
    if p == "mistral":
        return False  # Mistral wird in dieser Integration nur fuer Text genutzt
    return False  # GLM, Kimi, Mistral, Ollama: kein Vision-Support


def encode_file_for_vision(path: Path) -> tuple[dict, bool]:
    """Enkodiert eine PDF- oder Bilddatei als Anthropic-Content-Block.

    Returns:
        (content_block, needs_pdf_beta) – needs_pdf_beta ist True für PDFs.

    Raises:
        ValueError: Dateityp nicht unterstützt oder Datei zu groß.
        OSError: Datei nicht lesbar.
    """
    suffix = path.suffix.lower()
    size_kb = get_file_size_kb(path)
    limit_kb = 32 * 1024 if suffix == ".pdf" else 20 * 1024
    if size_kb > limit_kb:
        raise ValueError(f"Datei zu groß: {size_kb} KB (Limit: {limit_kb} KB für {suffix})")
    media_type_map = {
        ".pdf": "application/pdf",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
    }
    media_type = media_type_map.get(suffix)
    if not media_type:
        raise ValueError(f"Nicht unterstützter Dateityp: {suffix}")
    data = base64.standard_b64encode(path.read_bytes()).decode("ascii")
    if suffix == ".pdf":
        block = {
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf", "data": data},
        }
        return block, True
    block = {
        "type": "image",
        "source": {"type": "base64", "media_type": media_type, "data": data},
    }
    return block, False


_NAME_BLOCKLIST: frozenset[str] = frozenset(
    {
        "Deutsch", "Deutschtext", "Kommentar", "Weber", "Die", "Der", "Das",
        "Textinterpretation", "Interpretation", "Analyse", "Jugendsprache",
        "Text", "Krass", "Digga", "Und", "Oder", "Eine", "Einen", "Einem",
    }
)


def extract_vorname_from_filename(filename: str) -> str:
    """Schlägt einen Vornamen aus dem Dateinamen vor (DSGVO: nur Vorname).

    Sucht das erste Wort mit Großbuchstabe + mind. 3 Kleinbuchstaben,
    das nicht in der Blockliste steht. Gibt "" zurück wenn nichts gefunden.
    """
    stem = Path(filename).stem
    parts = re.split(r"[—\-_\s,]+", stem)
    for part in parts:
        part = part.strip()
        if re.match(r"^[A-ZÄÖÜ][a-zäöüß]{2,}$", part) and part not in _NAME_BLOCKLIST:
            return part
    return ""


def extract_schueler_name(docx_path: Path) -> str:
    """Extrahiert den Schülernamen aus Dateiname, DOCX-Header oder ersten Absätzen."""
    name = extract_vorname_from_filename(docx_path.name)
    if name:
        return name
    if docx_path.suffix.lower() != ".docx":
        return ""
    try:
        from docx import Document as _DocxDoc
        doc = _DocxDoc(str(docx_path))
        for section in doc.sections:
            for p in section.header.paragraphs:
                text = p.text.strip()
                if text and len(text) < 40:
                    for part in text.split():
                        if (
                            re.match(r"^[A-ZÄÖÜ][a-zäöüß]{2,}$", part)
                            and part not in _NAME_BLOCKLIST
                        ):
                            return part
        for p in doc.paragraphs[:3]:
            text = p.text.strip()
            if text and len(text) < 40:
                for part in text.split():
                    if re.match(r"^[A-ZÄÖÜ][a-zäöüß]{2,}$", part) and part not in _NAME_BLOCKLIST:
                        return part
    except Exception:
        logging.debug("Schülername-Extraktion fehlgeschlagen", exc_info=True)
    return ""


def extract_criteria_keys(rubric_content: str) -> list[str]:
    """Liest die verbindlichen JSON-Schlüssel aus dem '## JSON-Kriterien'-Abschnitt der Rubrik.

    Gibt eine leere Liste zurück, wenn kein solcher Abschnitt vorhanden ist.
    """
    match = re.search(
        r"##\s*JSON-Kriterien.*?\n(.*?)(?=\n##|\Z)", rubric_content, re.DOTALL
    )
    if not match:
        return []
    return re.findall(r"`([a-z][a-z_]*)`", match.group(1))


_RUBRIK_HEADER_RE = re.compile(
    r"^\ufeff?<!--[ \t]*luka-rubrik[ \t]*\r?\n(?P<body>.*?)[ \t]*\r?\n-->[ \t]*(?:\r?\n[ \t]*)*",
    re.DOTALL,
)
_RUBRIK_HEADER_FIELDS = ("titel", "fach", "schulstufe", "textsorte")


def parse_rubrik_header(text: str) -> dict[str, str]:
    """Liest den optionalen luka-rubrik-Kommentar am Anfang einer Rubrik.

    Alte oder selbst angelegte Rubriken ohne Header bleiben vollständig nutzbar;
    ihre Metadaten sind dann leere Strings.
    """
    header = {field: "" for field in _RUBRIK_HEADER_FIELDS}
    match = _RUBRIK_HEADER_RE.match(text)
    if not match:
        return header
    for line in match.group("body").splitlines():
        key, separator, value = line.partition(":")
        key = key.strip().lower()
        if separator and key in header:
            header[key] = value.strip()
    return header


def strip_rubrik_header(text: str) -> str:
    """Entfernt ausschließlich den optionalen Metadaten-Header einer Rubrik."""
    return _RUBRIK_HEADER_RE.sub("", text, count=1)


def load_rubric(rubric_filename: str, config: dict[str, Any]) -> str:
    rubric_dir = resolve_path(config, "rubrics")
    rubric_path = rubric_dir / rubric_filename
    if not rubric_path.exists():
        raise FileNotFoundError(f"Rubrik nicht gefunden: {rubric_path}")
    return strip_rubrik_header(rubric_path.read_text(encoding="utf-8"))


def list_all_rubrics(config: dict[str, Any]) -> list[str]:
    """Gibt alle .md-Dateien aus rubrics/ zurück (ohne README-Dateien)."""
    rubric_dir = resolve_path(config, "rubrics")
    if not rubric_dir.exists():
        return []
    return sorted(
        f.name for f in rubric_dir.glob("*.md") if not f.name.upper().startswith("README")
    )


def load_rubric_for_aufgabe(config: dict[str, Any], klasse: str | None, aufgabe: str | None) -> str:
    """Lädt die Rubrik für eine Aufgabe — mit Fallback-Kette.

    1. aufgabe_cfg["rubric"] (Dateiname in rubrics/)
    2. default_rubric_for(fach, schulstufe)
    3. Erste verfügbare .md-Datei in rubrics/
    """
    auf_cfg = get_aufgabe_cfg(config, klasse or "", aufgabe or "") if klasse and aufgabe else {}
    rubric_name = auf_cfg.get("rubric", "")
    if rubric_name:
        try:
            return load_rubric(rubric_name, config)
        except FileNotFoundError:
            pass
    fach = auf_cfg.get("fach") or config.get("defaults", {}).get("fach", "Deutsch")
    schulstufe = auf_cfg.get("schulstufe") or config.get("defaults", {}).get(
        "schulstufe", "Oberstufe"
    )
    default = default_rubric_for(fach, schulstufe, config)
    if default:
        return load_rubric(default, config)
    rubric_dir = resolve_path(config, "rubrics")
    for f in sorted(rubric_dir.glob("*.md")):
        if not f.name.upper().startswith("README"):
            return load_rubric(f.name, config)
    raise FileNotFoundError("Keine Rubrik gefunden")


def set_rubric_for_aufgabe(klasse: str, aufgabe: str, rubric_name: str) -> None:
    """Setzt die Rubrik einer Aufgabe in natascha_config.toml (kein Dateikopieren)."""
    doc = _load_toml_doc()
    doc["classes"][klasse]["aufgaben"][aufgabe]["rubric"] = rubric_name
    _save_toml_doc(doc)


def attach_rubric_to_aufgabe(klasse: str, aufgabe: str, source_path: Path) -> str:
    """Kopiert source_path nach rubrics/{aufgabe}_{dateiname} und aktualisiert die Config.

    Gibt den neuen Dateinamen (relativ zu rubrics/) zurück.
    """
    config = load_config()
    rubric_dir = resolve_path(config, "rubrics")
    rubric_dir.mkdir(parents=True, exist_ok=True)
    dest_name = f"{aufgabe}_{source_path.name}"
    dest = rubric_dir / dest_name
    shutil.copy2(str(source_path), str(dest))
    set_rubric_for_aufgabe(klasse, aufgabe, dest_name)
    return dest_name


def load_schema(config: dict[str, Any]) -> dict[str, Any]:
    schema_path = resolve_path(config, "schema")
    if not schema_path.exists():
        return {}
    return json.loads(schema_path.read_text(encoding="utf-8"))


def load_example_fixture() -> str:
    fixtures_dir = PROJECT_ROOT / "tests" / "fixtures"
    # Bewusst gepinnt: das Beispiel-JSON landet in JEDEM Live-Korrektur-Prompt.
    # Ohne Pin würde eine neu abgelegte, alphabetisch frühere Fixture-Datei das
    # Modellverhalten still verändern (Audit N1, docs/AUDIT-prompts-didaktik.md).
    pinned = fixtures_dir / "beispiel_deutsch_kommentar.json"
    if pinned.exists():
        return pinned.read_text(encoding="utf-8")
    candidates = sorted(fixtures_dir.glob("*.json"))
    if not candidates:
        return "{}"
    return candidates[0].read_text(encoding="utf-8")


def _fehler_anweisungen(fach: str, wortanzahl: int = 0) -> str:
    """Gemeinsame 'fehler'-Regeln für Text- und Vision-Prompt (Audit N3/N4/N5).

    Fach-konditioniert (die Nachsuch-Checkliste ist sprachspezifisch) und mit
    an die Textlänge gekoppelter Fehler-Erwartung statt Pauschalspanne.
    wortanzahl=0 (Vision: Text unbekannt) lässt den Zahlen-Anker weg.
    """
    ist_englisch = fach.strip().lower().startswith("engl")
    kopf = (
        "- 'fehler': VOLLSTÄNDIGE Liste ALLER Sprachfehler im Schülertext. KEIN LIMIT.\n"
        "  Jeder Eintrag: 'zitat' (1–6 Wörter, wortgetreu), 'korrektur' (korrekte Fassung),\n"
        "  'typ' (R=Rechtschreibung, G=Grammatik, Z=Zeichensetzung, A=Ausdruck/Stil — Register, Wortwahl),\n"
        "  'erklaerung' (optional, 1 Satz Regelhinweis).\n"
    )
    if ist_englisch:
        beispiele = (
            "  Beispiele:\n"
            "  {\"zitat\":\"he go\",\"korrektur\":\"he goes\",\"typ\":\"G\",\"erklaerung\":\"3rd person singular -s\"}\n"
            "  {\"zitat\":\"informations\",\"korrektur\":\"information\",\"typ\":\"G\",\"erklaerung\":\"uncountable noun\"}\n"
            "  {\"zitat\":\"make a photo\",\"korrektur\":\"take a photo\",\"typ\":\"A\",\"erklaerung\":\"Kollokation\"}\n"
        )
        checkliste = (
            "  Zeitformen und Aspekt (past simple vs. present perfect), Subjekt-Verb-Kongruenz\n"
            "  (3rd person -s), Wortstellung (SVO, Adverbien), Präpositionen, Artikel,\n"
            "  False Friends und Kollokationen.\n"
        )
        austria = ""
    else:
        beispiele = (
            "  Beispiele:\n"
            "  {\"zitat\":\"dem lesen\",\"korrektur\":\"dem Lesen\",\"typ\":\"R\",\"erklaerung\":\"Substantivierter Infinitiv wird großgeschrieben\"}\n"
            "  {\"zitat\":\"Regale die sich\",\"korrektur\":\"Regale, die sich\",\"typ\":\"Z\",\"erklaerung\":\"Komma vor Relativsatz\"}\n"
            "  {\"zitat\":\"bei Seite legen\",\"korrektur\":\"beiseitelegen\",\"typ\":\"R\",\"erklaerung\":\"Zusammenschreibung bei übertragener Bedeutung\"}\n"
        )
        checkliste = (
            "  Kommas vor Nebensätzen (dass, weil, obwohl, wenn, da, ob),\n"
            "  Kommas vor Relativsätzen (der, die, das, welcher),\n"
            "  Groß-/Kleinschreibung bei Substantivierungen,\n"
            "  das/dass-Unterscheidung, Subjekt-Verb-Kongruenz.\n"
        )
        austria = (
            "  ÖSTERREICHISCHES STANDARDDEUTSCH ist KEIN Fehler: Wörter wie Jänner, heuer,\n"
            "  Marille, Paradeiser sowie das Perfekt mit 'sein' bei Positionsverben\n"
            "  ('bin gesessen', 'bin gestanden') sind korrekt. Nimm solche Formen ÜBERHAUPT\n"
            "  NICHT in die 'fehler'-Liste auf — auch nicht als Ausdruck/Stil-Eintrag oder\n"
            "  als Eintrag ohne Änderung.\n"
        )
    if wortanzahl > 0:
        erwartung = (
            f"  Der Schülertext hat etwa {wortanzahl} Wörter. Faustregel bei Schularbeiten:\n"
            "  ungefähr ein Sprachfehler je 25–40 Wörter. Findest du deutlich weniger,\n"
            "  lies den Text nochmals und suche gezielt nach:\n"
        )
    else:
        erwartung = (
            "  Findest du auffällig wenige Fehler, lies den Text nochmals und suche gezielt nach:\n"
        )
    return (
        kopf
        + beispiele
        + "  WICHTIG: ALLE Fehler erfassen — bei 25 Kommafehlern 25 Einträge mit typ=Z.\n"
        "  Die Lehrkraft verlässt sich auf Vollständigkeit. Kein Limit.\n"
        "  JEDER Eintrag muss eine SICHTBARE Korrektur enthalten: 'korrektur' unterscheidet\n"
        "  sich vom 'zitat' (fehlende Satzzeichen gehören INS Zitat-Fenster: zitat=\"Auswahl weil\",\n"
        "  korrektur=\"Auswahl, weil\"). Einträge mit identischem zitat und korrektur sind\n"
        "  UNGÜLTIG — reine Beobachtungen ohne Korrektur gehören in 'hinweise'.\n"
        + erwartung
        + checkliste
        + austria
    )


def load_schema_for_mode(config: dict[str, Any], bewertungsmodus: str) -> dict[str, Any]:
    """Gibt das JSON-Schema zurück — für 'unbenotet' ohne notenempfehlung im required-Array."""
    schema = load_schema(config)
    if bewertungsmodus == "unbenotet":
        schema = dict(schema)
        schema["properties"] = {
            k: v for k, v in schema.get("properties", {}).items() if k != "notenempfehlung"
        }
        schema["required"] = [r for r in schema.get("required", []) if r != "notenempfehlung"]
    else:
        req = list(schema.get("required", []))
        if "notenempfehlung" not in req:
            req.append("notenempfehlung")
        schema = dict(schema, required=req)
    return schema


def build_analysis_prompt(
    docx_text: str,
    rubric_content: str,
    fach: str,
    schulstufe: str,
    textsorte: str,
    config: dict[str, Any],
    schueler: str = "",
    bewertungsmodus: str = "benotet",
    ausgangstext_text: str | None = None,
    klasse: str = "",
    aufgabe: str = "",
    erwartungshorizont_name: str = "",
) -> str:
    schema = load_schema_for_mode(config, bewertungsmodus)
    example = load_example_fixture()
    schema_str = json.dumps(schema, indent=2, ensure_ascii=False)
    example_str = (
        json.dumps(json.loads(example), indent=2, ensure_ascii=False) if example != "{}" else ""
    )

    criteria_keys = extract_criteria_keys(rubric_content)
    criteria_instruction = (
        "KRITERIEN: Verwende im 'bewertung'-Objekt GENAU diese JSON-Schlüssel (keine anderen): "
        + ", ".join(f'"{k}"' for k in criteria_keys)
        + ".\n"
        if criteria_keys
        else ""
    )

    if bewertungsmodus == "unbenotet":
        noten_instruction = (
            "WICHTIG: Dies ist eine Hausaufgabe — keine Note vergeben.\n"
            "Das Feld 'notenempfehlung' WEGLASSEN (nicht im JSON ausgeben).\n"
        )
    else:
        noten_instruction = "Berechne eine Notenempfehlung.\n"

    return (
        "Du bist ein Korrekturassistent für österreichische Gymnasium-Schularbeiten.\n"
        f"Fach: {fach}\nSchulstufe: {schulstufe}\nTextsorte: {textsorte}\n"
        + (f"Schüler/in: {schueler}\n" if schueler else "")
        + "\n"
        "BEWERTUNGSRASTER:\n---\n"
        f"{rubric_content}\n---\n\n"
        + _erwartungshorizont_block(config, klasse, aufgabe, override_name=erwartungshorizont_name)
        + (
            "AUSGANGSTEXT / ARBEITSAUFTRAG:\n---\n"
            f"{ausgangstext_text}\n---\n\n"
            "HINWEIS: Beziehe den Ausgangstext in die Inhaltsbewertung ein. "
            "Prüfe, ob der Schülertext inhaltlich korrekt auf den Ausgangstext eingeht "
            "und die Aufgabenstellung erfüllt.\n\n"
            if ausgangstext_text is not None else ""
        )
        + "SCHÜLERTEXT:\n---\n"
        f"{docx_text}\n---\n\n"
        "JSON-SCHEMA (dein Output MUSS konform sein):\n---\n"
        f"{schema_str}\n---\n\n"
        "BEISPIEL-JSON:\n---\n"
        f"{example_str}\n---\n\n"
        "AUFGABE:\n"
        "Analysiere den Schülertext anhand des Bewertungsrasters.\n"
        + criteria_instruction
        + "PUNKTE-SKALA: Vergib für jedes Kriterium GENAU eine ganzzahlige Punktzahl von 1 bis 5 "
        "(1 = nicht erfüllt, 2 = schwach/unzureichend, 3 = ausreichend/befriedigend, "
        "4 = gut, 5 = sehr gut). Orientiere dich strikt an den Stufenbeschreibungen im Raster. "
        "Verwende KEINE Dezimalzahlen, keine Werte außerhalb 1–5.\n"
        + "Erstelle eine Bewertung für jedes Kriterium mit Stufe, Punkten, Stärken, Schwächen und Vorschlägen.\n"
        + noten_instruction
        + "Erstelle außerdem:\n"
        "- 'zusammenfassung': 2-3 Sätze holistische Gesamteinschätzung der Arbeit.\n"
        "- 'staerken_global': Liste von genau 3 globalen Stärken der Arbeit (kurze Strings).\n"
        "- 'verbesserungsbereiche': Liste von genau 3 konkreten Verbesserungsbereichen (kurze Strings).\n"
        "- 'hinweise': Liste von 2–4 pädagogischen Randnotizen auf Wortebene (Stilbeobachtungen,\n"
        "  Stärken, Registerfragen) — NICHT für Rechtschreib-/Grammatikfehler.\n"
        "  Jeder Eintrag ist ein Objekt mit:\n"
        "  'zitat': GENAU 1–5 Wörter aus dem Schülertext (wortgetreu, kein ganzer Satz!)\n"
        "  'kommentar': kurze direkte Anmerkung dazu, wie eine Randnotiz einer Lehrerin\n"
        "  Beispiel: {\"zitat\": \"mega geil\", \"kommentar\": \"Register überdenken — zu umgangssprachlich\"}\n"
        "  WICHTIG: 'zitat' darf NIE länger als 5 Wörter sein. Keine ganzen Sätze als Zitat!\n"
        + _fehler_anweisungen(fach, len(docx_text.split()))
        + "ACHTUNG: Trage Sprachfehler AUSSCHLIESSLICH im Top-Level-Array 'fehler' ein. "
        "Verwende NICHT 'fehler_detail' oder 'fehlerschwerpunkte' innerhalb der Kriterien. "
        "Diese Felder sind veraltet und werden vom System ignoriert.\n"
        "KEINE KI-Floskeln ('Es wäre ratsam', 'Man könnte', 'Es lässt sich feststellen'). "
        "KEINE Formulierungen die auf KI hinweisen ('Als KI', 'Basierend auf meiner Analyse').\n\n"
        "WICHTIG: Antworte NUR mit validem JSON. Kein Markdown, kein Erklärtext, keine ```json-Blöcke.\n"
        "Das JSON muss dem obigen Schema entsprechen.\n"
        "SPRACHE: Schreibe ausschließlich auf Deutsch mit korrekten Sonderzeichen (ä, ö, ü, Ä, Ö, Ü, ß).\n"
        "Verwende niemals ae/oe/ue als Ersatz für Umlaute.\n"
    )


def build_vision_prompt(
    rubric_content: str,
    fach: str,
    schulstufe: str,
    textsorte: str,
    config: dict[str, Any],
    schueler: str = "",
    bewertungsmodus: str = "benotet",
    ausgangstext_text: str | None = None,
    has_vision_ausgangstext: bool = False,
    klasse: str = "",
    aufgabe: str = "",
    erwartungshorizont_name: str = "",
) -> str:
    """Baut den Text-Teil des Prompts für Vision-Analyse (PDF/Bild-Input).

    Der Schülertext ist als beigefügtes Bild/PDF vorhanden — kein docx_text nötig.
    """
    schema = load_schema_for_mode(config, bewertungsmodus)
    example = load_example_fixture()
    schema_str = json.dumps(schema, indent=2, ensure_ascii=False)
    example_str = (
        json.dumps(json.loads(example), indent=2, ensure_ascii=False) if example != "{}" else ""
    )

    criteria_keys = extract_criteria_keys(rubric_content)
    criteria_instruction = (
        "KRITERIEN: Verwende im 'bewertung'-Objekt GENAU diese JSON-Schlüssel (keine anderen): "
        + ", ".join(f'"{k}"' for k in criteria_keys)
        + ".\n"
        if criteria_keys
        else ""
    )

    if bewertungsmodus == "unbenotet":
        noten_instruction = (
            "WICHTIG: Dies ist eine Hausaufgabe — keine Note vergeben.\n"
            "Das Feld 'notenempfehlung' WEGLASSEN (nicht im JSON ausgeben).\n"
        )
    else:
        noten_instruction = "Berechne eine Notenempfehlung.\n"

    return (
        "Du bist ein Korrekturassistent für österreichische Gymnasium-Schularbeiten.\n"
        f"Fach: {fach}\nSchulstufe: {schulstufe}\nTextsorte: {textsorte}\n"
        + (f"Schüler/in: {schueler}\n" if schueler else "")
        + "\n"
        "BEWERTUNGSRASTER:\n---\n"
        f"{rubric_content}\n---\n\n"
        + _erwartungshorizont_block(config, klasse, aufgabe, override_name=erwartungshorizont_name)
        + (
            "AUSGANGSTEXT / ARBEITSAUFTRAG:\n---\n"
            f"{ausgangstext_text}\n---\n\n"
            "HINWEIS: Beziehe den Ausgangstext in die Inhaltsbewertung ein. "
            "Prüfe, ob der Schülertext inhaltlich korrekt auf den Ausgangstext eingeht "
            "und die Aufgabenstellung erfüllt.\n\n"
            if ausgangstext_text is not None else ""
        )
        + (
            "AUSGANGSTEXT / ARBEITSAUFTRAG: Das erste beigefügte Dokument/Bild ist der "
            "Ausgangstext (Vorlagetext + Aufgabenstellung). Das zweite Dokument/Bild ist "
            "die Schülerarbeit. Beziehe den Ausgangstext in die Inhaltsbewertung ein. "
            "Prüfe, ob der Schülertext inhaltlich korrekt auf den Ausgangstext eingeht "
            "und die Aufgabenstellung erfüllt.\n\n"
            if has_vision_ausgangstext else ""
        )
        + "SCHÜLERARBEIT: Das beigefügte Bild/PDF ENTHÄLT die handgeschriebene Schülerarbeit "
        "samt handschriftlichen Lehrermarkierungen und -korrekturen. "
        "Lies den handgeschriebenen Text vollständig, beachte alle eingetragenen Korrekturen "
        "und Anmerkungen der Lehrkraft, und analysiere die Arbeit anhand des Bewertungsrasters.\n\n"
        "JSON-SCHEMA (dein Output MUSS konform sein):\n---\n"
        f"{schema_str}\n---\n\n"
        "BEISPIEL-JSON:\n---\n"
        f"{example_str}\n---\n\n"
        "AUFGABE:\n"
        "Analysiere die im Bild/PDF sichtbare handgeschriebene Schülerarbeit anhand des Rasters.\n"
        + criteria_instruction
        + "PUNKTE-SKALA: Vergib für jedes Kriterium GENAU eine ganzzahlige Punktzahl von 1 bis 5 "
        "(1 = nicht erfüllt, 2 = schwach/unzureichend, 3 = ausreichend/befriedigend, "
        "4 = gut, 5 = sehr gut). Orientiere dich strikt an den Stufenbeschreibungen im Raster. "
        "Verwende KEINE Dezimalzahlen, keine Werte außerhalb 1–5.\n"
        + "Beziehe sichtbare Lehrermarkierungen/-korrekturen in die Bewertung ein.\n"
        "Erstelle eine Bewertung für jedes Kriterium mit Stufe, Punkten, Stärken, Schwächen und Vorschlägen.\n"
        + noten_instruction
        + "Erstelle außerdem:\n"
        "- 'transkription': Vollständige wortgetreue Abschrift des handgeschriebenen Schülertexts "
        "(inklusive aller Absätze, Überschrift falls vorhanden, und Zeilenumbrüche als \\n). "
        "Handschriftliche Lehrermarkierungen NICHT einschließen — nur den Schülertext selbst.\n"
        "- 'zusammenfassung': 2-3 Sätze holistische Gesamteinschätzung der Arbeit.\n"
        "- 'staerken_global': Liste von genau 3 globalen Stärken der Arbeit (kurze Strings).\n"
        "- 'verbesserungsbereiche': Liste von genau 3 konkreten Verbesserungsbereichen (kurze Strings).\n"
        "- 'hinweise': Liste von 2–4 pädagogischen Randnotizen auf Wortebene (Stilbeobachtungen,\n"
        "  Stärken, Registerfragen) — NICHT für Rechtschreib-/Grammatikfehler.\n"
        "  Jeder Eintrag ist ein Objekt mit:\n"
        "  'zitat': GENAU 1–5 Wörter aus dem Schülertext (wortgetreu, kein ganzer Satz!)\n"
        "  'kommentar': kurze direkte Anmerkung dazu, wie eine Randnotiz einer Lehrerin\n"
        "  Beispiel: {\"zitat\": \"mega geil\", \"kommentar\": \"Register überdenken — zu umgangssprachlich\"}\n"
        "  WICHTIG: 'zitat' darf NIE länger als 5 Wörter sein. Keine ganzen Sätze als Zitat!\n"
        + _fehler_anweisungen(fach)
        + "ACHTUNG: Trage Sprachfehler AUSSCHLIESSLICH im Top-Level-Array 'fehler' ein. "
        "Verwende NICHT 'fehler_detail' oder 'fehlerschwerpunkte' innerhalb der Kriterien. "
        "Diese Felder sind veraltet und werden vom System ignoriert.\n"
        "KEINE KI-Floskeln ('Es wäre ratsam', 'Man könnte', 'Es lässt sich feststellen'). "
        "KEINE Formulierungen die auf KI hinweisen ('Als KI', 'Basierend auf meiner Analyse').\n\n"
        "WICHTIG: Antworte NUR mit validem JSON. Kein Markdown, kein Erklärtext, keine ```json-Blöcke.\n"
        "Das JSON muss dem obigen Schema entsprechen.\n"
        "SPRACHE: Schreibe ausschließlich auf Deutsch mit korrekten Sonderzeichen (ä, ö, ü, Ä, Ö, Ü, ß).\n"
        "Verwende niemals ae/oe/ue als Ersatz für Umlaute.\n"
    )


def extract_json_from_llm(text: str) -> dict[str, Any]:
    fenced = re.search(r"```(?:json)?\s*\n?(.*?)```", text, re.DOTALL)
    candidate = fenced.group(1).strip() if fenced else text.strip()
    start = candidate.find("{")
    if start != -1:
        depth = 0
        for i, ch in enumerate(candidate[start:], start):
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    candidate = candidate[start : i + 1]
                    break
    return json.loads(candidate)


def _inline_schema_refs(schema: dict[str, Any]) -> dict[str, Any]:
    """Löst $ref-Referenzen auf, damit Anthropic tool_use die Struktur korrekt versteht.

    Claude interpretiert unaufgelöste $ref-Referenzen nicht korrekt und serialisiert
    dann verschachtelte Felder (z. B. 'bewertung') als JSON-String statt als Objekt.
    """
    defs = schema.get("$defs", {})
    if not defs:
        return schema

    def resolve(obj: Any) -> Any:
        if isinstance(obj, dict):
            if "$ref" in obj and len(obj) == 1:
                ref = obj["$ref"]
                if ref.startswith("#/$defs/"):
                    def_name = ref[len("#/$defs/"):]
                    return resolve(dict(defs.get(def_name, obj)))
                return obj
            return {k: resolve(v) for k, v in obj.items() if k not in ("$defs", "$schema")}
        if isinstance(obj, list):
            return [resolve(item) for item in obj]
        return obj

    return resolve(dict(schema))


def validate_against_schema(data: dict[str, Any], schema: dict[str, Any]) -> list[str]:
    if _jsmod is None or not schema:
        return []
    errors: list[str] = []
    for err in sorted(
        _jsmod.Draft202012Validator(schema).iter_errors(data), key=lambda e: e.message
    ):
        errors.append(err.message)
    return errors


# ---------------------------------------------------------------------------
# Robuste LLM-Analyse-Pipeline: Prompt → API → JSON → Validierung → Retry
# ---------------------------------------------------------------------------

_MAX_RETRIES = 3
_RETRY_BACKOFF_SECONDS = 2


def _build_retry_prompt(
    original_prompt: str,
    error_message: str,
    llm_raw_response: str,
    attempt: int,
) -> str:
    """Erzeugt einen Retry-Prompt, der das LLM zur JSON-Korrektur auffordert."""
    extra_hint = ""
    if "is not of type 'object'" in error_message:
        extra_hint = (
            "\nKRITISCHER FEHLER: Ein Feld (z. B. 'bewertung' oder 'notenempfehlung') "
            "wurde als JSON-String zurückgegeben statt als direktes JSON-Objekt.\n"
            "FALSCH:  \"bewertung\": \"{\\\"inhalt\\\": {...}}\"\n"
            "RICHTIG: \"bewertung\": {\"inhalt\": {...}}\n"
            "Das Feld muss ein verschachteltes JSON-Objekt sein, KEIN String.\n"
        )
    return (
        f"DEINE VORHERIGE ANTWORT WAR UNGÜLTIG (Versuch {attempt}).\n\n"
        f"FEHLER: {error_message}{extra_hint}\n\n"
        f"DEINE ANTWORT:\n---\n{llm_raw_response[:2000]}\n---\n\n"
        f"ORIGINAL-AUFGABE (zur Erinnerung):\n---\n{original_prompt[:3000]}\n---\n\n"
        f"ANTWORTE JETZT AUSSCHLIESSLICH MIT GÜLTIGEM JSON. "
        f"Kein Markdown, kein Code-Block, kein Erklärtext. "
        f"NUR das JSON-Objekt, beginnend mit {{ und endend mit }}."
    )


def run_llm_analysis(
    docx_text: str,
    rubric_content: str,
    fach: str,
    schulstufe: str,
    textsorte: str,
    config: dict[str, Any],
    schueler: str = "",
    cancel_event: threading.Event | None = None,
    max_retries: int = _MAX_RETRIES,
    file_path: Path | None = None,
    bewertungsmodus: str = "benotet",
    ausgangstext_path: Path | None = None,
    klasse: str = "",
    aufgabe: str = "",
    erwartungshorizont_name: str = "",
    rubric_name: str = "",
    pseudonymisierung: bool = True,
    db_path_override: Path | None = None,
    bestaetigte_schueler_id: int | None = None,
    unterrichtseinsatz_id: str | None = None,
    material_id: str | None = None,
) -> tuple[dict[str, Any] | None, list[str]]:
    """
    Fuehrt die vollstaendige LLM-Analyse durch: Prompt bauen, API aufrufen,
    JSON extrahieren, gegen Schema validieren. Bei Fehlern wird bis zu
    *max_retries*-mal wiederholt.

    Wenn file_path auf eine PDF- oder Bilddatei zeigt, wird Vision-Modus aktiviert:
    Die Datei wird base64-enkodiert und als multimodaler Content-Block übertragen.

    Wenn ausgangstext_path gesetzt ist, wird der Ausgangstext (DOCX → Text,
    PDF/Bild → Vision-Block) vor dem Schüler-Content in den Prompt/Content eingebettet.

    Returns:
        (data, errors) – data ist das validierte JSON-Dict oder None,
        errors ist eine Liste aller aufgetretenen Fehlermeldungen.
    """
    schema = load_schema_for_mode(config, bewertungsmodus)

    # Duplikat-Erkennung (Hash-Check)
    if ndb is not None and file_path is not None and file_path.exists():
        try:
            db_path = ndb.get_db_path(config)
            ndb.init_db(db_path)
            file_hash = ndb._file_hash(file_path)
            existing = ndb.get_abgabe_by_hash(db_path, file_hash)
            if existing:
                # Versuche, gespeichertes JSON zu laden
                json_path_str = existing.get("feedback_json_path", "")
                if json_path_str and Path(json_path_str).exists():
                    try:
                        cached = json.loads(Path(json_path_str).read_text(encoding="utf-8"))
                        cached["_abgabe_id"] = existing.get("id")
                        return cached, ["Hinweis: Bereits analysiert (Duplikat)."]
                    except Exception:
                        logging.debug("Duplikat-Cache nicht lesbar", exc_info=True)
                # Minimal-Dict mit Abgabe-ID, damit Lehrer-Feedback möglich ist
                minimal = {
                    "_abgabe_id": existing.get("id"),
                    "datei": file_path.name,
                    "notenempfehlung": {
                        "note": existing.get("note"),
                        "bezeichnung": "",
                        "durchschnitt": existing.get("gesamtstufe"),
                    },
                }
                return minimal, [
                    f"Duplikat: '{file_path.name}' wurde am {existing.get('datum')}"
                    f" bereits analysiert (Note: {existing.get('note')})."
                ]
        except Exception:
            # DB-Check darf Analyse nie blockieren
            pass

    # Ausgangstext vorverarbeiten
    ausgangstext_text: str | None = None
    ausgangstext_vision_block: dict | None = None
    ausgangstext_warnings: list[str] = []

    if ausgangstext_path is not None:
        suf = ausgangstext_path.suffix.lower()
        if suf == ".docx":
            try:
                ausgangstext_text = read_docx_text(ausgangstext_path)
            except Exception as e:
                ausgangstext_warnings.append(f"Ausgangstext nicht lesbar: {e}")
        elif suf in _VISION_EXTENSIONS:
            provider = config.get("api", {}).get("provider", "anthropic").lower()
            if provider == "mistral":
                ausgangstext_warnings.append(
                    "Ausgangstext wird mit Mistral nicht als Vision verarbeitet; "
                    "bitte DOCX/Text verwenden."
                )
            elif provider == "openai" and suf == ".pdf":
                ausgangstext_warnings.append(
                    "Ausgangstext PDF wird mit OpenAI übersprungen "
                    "(kein nativer PDF-Support)."
                )
            else:
                try:
                    ausgangstext_vision_block, _ = encode_file_for_vision(ausgangstext_path)
                except (ValueError, OSError) as e:
                    ausgangstext_warnings.append(f"Ausgangstext Enkodierung fehlgeschlagen: {e}")

    vision_mode = file_path is not None and is_vision_file(file_path)

    # Pseudonymisierung: Namen aus der Klassenliste werden vor dem Versand an
    # den externen Anbieter durch Aliasse ersetzt (Datenminimierung, P0).
    # Der Original-docx_text bleibt unangetastet — er ist der rohtext in der DB;
    # nur der LLM-Pfad (Prompt, Zitat-Filter, SRDP) sieht die Alias-Fassung.
    docx_text_llm = docx_text
    schueler_llm = schueler
    pseudo_funde: list[dict[str, Any]] = []
    if pseudonymisierung:
        if vision_mode:
            ausgangstext_warnings.append(
                "Hinweis: PDF/Bild-Abgabe — Personenangaben können darin nicht "
                "ersetzt werden; die Datei geht unverändert an den Anbieter."
            )
        elif not klasse:
            # Ohne Klassenbezug existiert keine Klassenliste — es gibt nichts,
            # wogegen erkannt werden könnte. Kein Hinweis-Rauschen.
            pass
        elif pseu is None or ndb is None:
            ausgangstext_warnings.append(
                "Hinweis: Pseudonymisierung übersprungen (keine Klassenliste "
                "verfügbar) — Personenangaben wurden NICHT ersetzt."
            )
        else:
            try:
                # --db-path-Override der CLI respektieren (gemeinsame DB mit LUA),
                # sonst würde das Roster aus der Config-DB gelesen.
                _db_path = db_path_override or ndb.get_db_path(config)
                ndb.init_db(_db_path)
                _roster = ndb.get_schueler_by_klasse(_db_path, klasse)
                pseudo_funde = pseu.erkenne_personenangaben(
                    docx_text, file_path.name if file_path else "", schueler, _roster
                )
                if pseudo_funde:
                    docx_text_llm = pseu.ersetze_personenangaben(docx_text, pseudo_funde)
                    schueler_llm = pseu.alias_fuer_schuelerangabe(schueler, pseudo_funde) or ""
                    _anzahl = sum(f["vorkommen_text"] for f in pseudo_funde)
                    ausgangstext_warnings.append(
                        f"Hinweis: {_anzahl} Personenangabe(n) aus der Klassenliste "
                        "vor dem LLM-Versand durch Aliasse ersetzt."
                    )
            except Exception as e:
                ausgangstext_warnings.append(
                    f"Hinweis: Pseudonymisierung fehlgeschlagen ({e}) — "
                    "Personenangaben wurden NICHT ersetzt."
                )

    if vision_mode:
        if config.get("api", {}).get("provider", "anthropic").lower() == "mistral":
            return None, ausgangstext_warnings + [
                "Mistral unterstützt in dieser Integration nur Textdateien."
            ]
        try:
            vision_block, _ = encode_file_for_vision(file_path)  # type: ignore[arg-type]
            vision_content: list[dict] | None = (
                [ausgangstext_vision_block, vision_block]
                if ausgangstext_vision_block is not None
                else [vision_block]
            )
        except (ValueError, OSError) as e:
            return None, ausgangstext_warnings + [f"Datei konnte nicht enkodiert werden: {e}"]
        original_prompt = build_vision_prompt(
            rubric_content, fach, schulstufe, textsorte, config, schueler_llm, bewertungsmodus,
            ausgangstext_text=ausgangstext_text,
            has_vision_ausgangstext=(ausgangstext_vision_block is not None),
            klasse=klasse, aufgabe=aufgabe,
            erwartungshorizont_name=erwartungshorizont_name,
        )
        effective_max_retries = 1  # Retries ohne Bild weniger hilfreich
    else:
        if ausgangstext_vision_block is not None:
            # Student-DOCX + Vision-Ausgangstext: Vision-Block als Prefix, Text im Prompt
            vision_content = [ausgangstext_vision_block]
            ausgangstext_note = (
                "AUSGANGSTEXT / ARBEITSAUFTRAG: Das beigefügte Dokument/Bild ist der Ausgangstext "
                "(Vorlagetext + Aufgabenstellung). Beziehe ihn in die Inhaltsbewertung ein. "
                "Prüfe, ob der Schülertext inhaltlich korrekt auf den Ausgangstext eingeht "
                "und die Aufgabenstellung erfüllt.\n\n"
            )
            original_prompt = ausgangstext_note + build_analysis_prompt(
                docx_text_llm, rubric_content, fach, schulstufe, textsorte,
                config, schueler_llm, bewertungsmodus, ausgangstext_text=None,
                klasse=klasse, aufgabe=aufgabe,
                erwartungshorizont_name=erwartungshorizont_name,
            )
            effective_max_retries = max_retries
        else:
            vision_content = None
            original_prompt = build_analysis_prompt(
                docx_text_llm, rubric_content, fach, schulstufe, textsorte, config, schueler_llm,
                bewertungsmodus, ausgangstext_text=ausgangstext_text,
                klasse=klasse, aufgabe=aufgabe,
                erwartungshorizont_name=erwartungshorizont_name,
            )
            effective_max_retries = max_retries

    errors: list[str] = list(ausgangstext_warnings)
    current_prompt = original_prompt

    for attempt in range(1, effective_max_retries + 1):
        if cancel_event is not None and cancel_event.is_set():
            errors.append("Analyse abgebrochen")
            return None, errors

        # Vision-Content nur beim ersten Versuch; Retries korrigieren nur JSON
        vc = vision_content if attempt == 1 else None
        raw_response = run_llm_api(
            current_prompt, config, cancel_event=cancel_event, schema=schema, vision_content=vc
        )

        if cancel_event is not None and cancel_event.is_set():
            errors.append("Analyse abgebrochen")
            return None, errors

        if raw_response.startswith("FEHLER"):
            errors.append(f"API-Fehler (Versuch {attempt}): {raw_response}")
            # Bei API-Fehlern lohnt sich kein Retry mit gleichem Prompt
            return None, errors

        # Versuch, JSON zu extrahieren
        try:
            data = extract_json_from_llm(raw_response)
        except (json.JSONDecodeError, AttributeError) as e:
            errors.append(f"JSON-Extraktion fehlgeschlagen (Versuch {attempt}): {e}")
            current_prompt = _build_retry_prompt(
                original_prompt, f"Kein gueltiges JSON gefunden: {e}", raw_response, attempt
            )
            time.sleep(_RETRY_BACKOFF_SECONDS * attempt)
            continue

        # Verschachtelte Felder normalisieren: Tool-Use gibt sie gelegentlich als JSON-String zurück
        for _key in ("bewertung", "notenempfehlung"):
            if isinstance(data.get(_key), str):
                try:
                    data[_key] = json.loads(data[_key])
                except (json.JSONDecodeError, TypeError) as _repair_err:
                    logging.warning(
                        "JSON-Repair fehlgeschlagen für '%s' (Versuch %d): %s",
                        _key, attempt, _repair_err,
                    )

        # Gegen Schema validieren
        validation_errors = validate_against_schema(data, schema)
        if validation_errors:
            errors.append(
                f"Schema-Validierung fehlgeschlagen (Versuch {attempt}): "
                f"{'; '.join(validation_errors[:3])}"
            )
            current_prompt = _build_retry_prompt(
                original_prompt,
                f"Schema-Verletzung: {'; '.join(validation_errors[:3])}",
                raw_response,
                attempt,
            )
            time.sleep(_RETRY_BACKOFF_SECONDS * attempt)
            continue

        # Metadaten-Override: bekannte Werte aus Ordnerstruktur/Config haben Vorrang vor LLM
        if file_path is not None:
            data["datei"] = file_path.name
        if klasse:
            data["klasse"] = klasse
        if fach:
            data["fach"] = fach
        if schulstufe:
            data["schulstufe"] = schulstufe
        if textsorte:
            data["textsorte"] = textsorte
        if rubric_name:
            data["rubrik"] = rubric_name
        if schueler:
            data["schueler"] = schueler
        # Ausgangstext für die LUA-Brücke mitführen (Quelltext-Vorbefüllung der Übung).
        if ausgangstext_text:
            data["ausgangstext"] = ausgangstext_text

        # Nullnummern-Filter: Einträge ohne sichtbare Korrektur raus (auch Vision).
        if data.get("fehler"):
            data["fehler"] = drop_unbrauchbare_fehler(data["fehler"])

        # Halluzinationsfilter: Fehler-Zitate gegen den Text prüfen, den das
        # LLM gesehen hat (Alias-Fassung) — Zitate enthalten ggf. Aliasse.
        if data.get("fehler") and docx_text_llm and not vision_mode:
            data["fehler"] = verify_fehler_against_text(data["fehler"], docx_text_llm)
            data["fehler"] = drop_satzzeichen_anhaengsel(data["fehler"], docx_text_llm)
            data["fehler"] = filter_title_false_positives(data["fehler"], docx_text_llm)

        # Konsistenzcheck: Fehleranzahl vs. Sprachrichtigkeit-Note
        fehler_count = len(data.get("fehler", []))
        sprach_key = next(
            (k for k in data.get("bewertung", {}) if "sprach" in k.lower()), None
        )
        if sprach_key:
            sprachnote = data["bewertung"][sprach_key].get("punkte", 3)
            if fehler_count > 15 and sprachnote >= 4:
                logging.warning(
                    "Inkonsistenz: %d Fehler aber Sprachrichtigkeit-Note %s",
                    fehler_count,
                    sprachnote,
                )

        # SRDP-Detailbewertung (zweiter LLM-Call) — vor Notenberechnung, damit Note daraus folgt.
        # Auch hier die Alias-Fassung, sonst gingen die Namen im zweiten Call doch raus.
        if bewertungsmodus == "benotet" and docx_text_llm and not vision_mode:
            srdp = generate_srdp_detail(
                docx_text_llm, data, config, cancel_event=cancel_event, textsorte=textsorte
            )
            if srdp:
                data["srdp_detail"] = srdp

        # App-Notenberechnung — aus SRDP-Detail wenn vorhanden, sonst Fallback auf Hauptanalyse
        if bewertungsmodus == "benotet" and data.get("bewertung"):
            if "notenempfehlung" in data:
                data["notenempfehlung_llm"] = data["notenempfehlung"]
            if schulstufe.lower() in ("oberstufe", "ahs-oberstufe"):
                app_note = berechne_note_srdp(
                    data["bewertung"], data.get("srdp_detail")
                )
            else:
                gew = parse_gewichtung(rubric_content) if rubric_content else None
                app_note = berechne_note_unterstufe(data["bewertung"], gew)
            data["notenempfehlung"] = {
                "durchschnitt": app_note["durchschnitt"],
                "note": app_note["note"],
                "bezeichnung": app_note["bezeichnung"],
                "begruendung": app_note["begruendung"],
            }
            data["notendetail"] = app_note

        # Aliasse in der Antwort zurücksetzen, BEVOR gespeichert wird: Zitate
        # müssen zum Original-rohtext passen (Annotation), Feedback-Texte sollen
        # die Lehrkraft mit echtem Namen erreichen. Nur die Übertragung war
        # pseudonymisiert.
        if pseudo_funde and pseu is not None:
            data = pseu.ruecksetze_personenangaben(data, pseudo_funde)

        # In Datenbank speichern (Tracking, Heatmap, Duplikat-Erkennung)
        if ndb is not None and file_path is not None and file_path.exists():
            try:
                db_path = ndb.get_db_path(config)
                ndb.init_db(db_path)
                wortanzahl = count_words(file_path)

                # Analyse-JSON auf Platte schreiben, damit feedback-docx
                # darauf zugreifen kann (feedback_json_path in abgabe-Tabelle)
                json_path_str = ""
                if klasse and aufgabe:
                    auf_cfg = get_aufgabe_cfg(config, klasse, aufgabe)
                    auf_out = auf_cfg.get("output", "")
                    fb_dir = (
                        PROJECT_ROOT / auf_out / "feedback_data"
                        if auf_out
                        else resolve_path(config, "output") / "feedback_data"
                    )
                    try:
                        import json as _json_mod
                        fb_dir.mkdir(parents=True, exist_ok=True)
                        json_file = fb_dir / f"{file_path.stem}_analysis.json"
                        json_file.write_text(
                            _json_mod.dumps(data, ensure_ascii=False, indent=2, default=str),
                            encoding="utf-8",
                        )
                        json_path_str = str(json_file)
                    except Exception:
                        pass

                abgabe_id = ndb.save_analysis_to_db(
                    db_path=db_path,
                    data=data,
                    file_path=file_path,
                    klasse=klasse,
                    aufgabe=aufgabe,
                    rohtext=docx_text,
                    wortanzahl=wortanzahl,
                    feedback_json_path=json_path_str,
                    bestaetigte_schueler_id=bestaetigte_schueler_id,
                    unterrichtseinsatz_id=unterrichtseinsatz_id,
                    material_id=material_id,
                )
                if abgabe_id and abgabe_id > 0:
                    data["_abgabe_id"] = abgabe_id
                else:
                    # save gab -1 (Duplikat): bestehende Abgabe-ID nachschlagen,
                    # damit Lehrer-Feedback trotzdem zugeordnet werden kann
                    existing = ndb.get_abgabe_by_hash(
                        db_path, ndb._file_hash(file_path)
                    )
                    if existing and existing.get("id"):
                        data["_abgabe_id"] = existing["id"]
            except Exception:
                # DB-Fehler duerfen Analyse nicht blockieren
                pass

        # Warnhinweise (Ausgangstext, Pseudonymisierung) auch bei Erfolg
        # zurückgeben — die Lehrkraft soll sehen, was übertragen wurde.
        # Retry-Fehlermeldungen früherer Versuche bleiben bewusst draußen.
        return data, list(ausgangstext_warnings)

    errors.append(
        f"Analyse nach {effective_max_retries} Versuchen fehlgeschlagen. "
        "Siehe fehlerlog.txt fuer Details."
    )
    return None, errors


def drop_unbrauchbare_fehler(fehler_list: list[dict]) -> list[dict]:
    """Entfernt Fehler-Einträge ohne sichtbare Korrektur (zitat == korrektur).

    Live-Eval P2b (2026-07-04): Modelle (v. a. DeepSeek) liefern trotz
    Prompt-Verbot Einträge, deren Korrektur identisch mit dem Zitat ist —
    unbrauchbar fürs Feedback und Gift für die Heatmap. Case-SENSITIV
    vergleichen: "lesen"→"Lesen" ist eine echte Korrektur.
    """
    if not fehler_list:
        return []
    behalten: list[dict] = []
    entfernt = 0
    for fehler in fehler_list:
        zitat = " ".join((fehler.get("zitat") or "").split())
        korrektur = " ".join((fehler.get("korrektur") or "").split())
        if not zitat or zitat == korrektur:
            entfernt += 1
            continue
        behalten.append(fehler)
    if entfernt:
        logging.warning("Fehler-Einträge ohne sichtbare Korrektur entfernt: %d", entfernt)
    return behalten


_SATZZEICHEN = ".,;:!?…"


def drop_satzzeichen_anhaengsel(fehler_list: list[dict], schuelertext: str) -> list[dict]:
    """Entfernt Pseudo-Korrekturen der Form korrektur == zitat + Satzzeichen,
    wenn im Originaltext direkt nach dem Zitat bereits ein Satzzeichen steht.

    Live-Eval P2c-Finale (2026-07-04): Modelle zitieren einen Satz OHNE sein
    Schluss-Satzzeichen und "korrigieren" durch Anhängen ("leer ist" →
    "leer ist.", obwohl der Punkt im Text steht). Folgt im Text dagegen ein
    Wort, kann das angehängte Komma legitim sein → Eintrag bleibt.
    """
    if not fehler_list or not schuelertext:
        return fehler_list or []
    norm_text = " ".join(schuelertext.lower().split())
    behalten: list[dict] = []
    entfernt = 0
    for fehler in fehler_list:
        zitat = " ".join((fehler.get("zitat") or "").split())
        korrektur = " ".join((fehler.get("korrektur") or "").split())
        angehaengt = (
            len(korrektur) == len(zitat) + 1
            and korrektur[:-1] == zitat
            and korrektur[-1] in _SATZZEICHEN
        )
        if angehaengt:
            nz = zitat.lower()
            pos = norm_text.find(nz)
            folgt_satzzeichen = False
            while pos != -1:
                nach = norm_text[pos + len(nz):pos + len(nz) + 1]
                if nach and nach in _SATZZEICHEN:
                    folgt_satzzeichen = True
                    break
                pos = norm_text.find(nz, pos + 1)
            if folgt_satzzeichen:
                entfernt += 1
                continue
        behalten.append(fehler)
    if entfernt:
        logging.warning("Satzzeichen-Anhängsel-Halluzinationen entfernt: %d", entfernt)
    return behalten


def verify_fehler_against_text(fehler_list: list[dict], schuelertext: str) -> list[dict]:
    """Entfernt halluzinierte Fehler deren Zitat nicht im Text vorkommt."""
    if not fehler_list or not schuelertext:
        return fehler_list or []

    normalized_text = " ".join(schuelertext.lower().split())

    verified: list[dict] = []
    removed: list[str] = []

    for fehler in fehler_list:
        zitat = fehler.get("zitat", "").strip()
        if not zitat:
            continue

        normalized_zitat = " ".join(zitat.lower().split())

        if normalized_zitat in normalized_text:
            verified.append(fehler)
            continue

        # Toleranter Fallback ohne Satzzeichen — aber NICHT für Zeichensetzungs-
        # Fehler: bei typ=Z ist das Satzzeichen die Substanz des Zitats. Ein
        # Z-Zitat, das nur ohne Satzzeichen matcht, ist eine Halluzination
        # (Live-Eval P2b: Mistral erfand fehlende Punkte, die im Text stehen).
        if fehler.get("typ") != "Z":
            stripped_zitat = re.sub(r"[^\w\s]", "", normalized_zitat)
            stripped_text = re.sub(r"[^\w\s]", "", normalized_text)
            if stripped_zitat and stripped_zitat in stripped_text:
                verified.append(fehler)
                continue

        removed.append(zitat)

    if removed:
        logging.warning(
            "Halluzinierte Fehler entfernt (%d): %s",
            len(removed),
            ", ".join(f"'{z}'" for z in removed[:5]),
        )

    return verified


def filter_title_false_positives(fehler_list: list, schuelertext: str) -> list:
    """Entfernt falsche R-Fehler bei Großschreibung am Absatz-/Titelanfang."""
    paragraphs = [p.strip() for p in schuelertext.split("\n") if p.strip()]
    first_words = {p.split()[0].lower() for p in paragraphs if p.split()}
    filtered = []
    for f in fehler_list:
        if f.get("typ") == "R" and "klein" in f.get("erklaerung", "").lower():
            zitat_first = f.get("zitat", "").split()[0].lower() if f.get("zitat") else ""
            if zitat_first in first_words:
                continue
        filtered.append(f)
    return filtered


def generate_srdp_detail(
    schuelertext: str,
    hauptanalyse: dict[str, Any],
    config: dict[str, Any],
    cancel_event: threading.Event | None = None,
    textsorte: str = "",
) -> dict[str, Any] | None:
    """Zweiter LLM-Call: SRDP-Detailbewertung mit 15 Unterkriterien (Option A)."""
    bewertung_json = json.dumps(hauptanalyse.get("bewertung", {}), ensure_ascii=False, indent=2)
    textsorte_upper = textsorte.upper() if textsorte else "TEXTSORTE"
    textsorte_hint = (
        f"TEXTSORTE: {textsorte}\n\n"
        f"WICHTIG: Bewerte die Arbeit ALS {textsorte_upper}. "
        f'Verwende in allen Begründungen die korrekte Textsortenbezeichnung "{textsorte}", '
        f"NICHT \"Erörterung\" oder andere Textsorten. "
        f"Beurteile die Schreibhandlung im Sinne der Textsorte {textsorte}.\n\n"
        if textsorte else ""
    )
    prompt = (
        "Du bist eine AHS-Deutschlehrerin. Vor dir liegt eine Schülerarbeit "
        "und die Hauptanalyse dazu.\n\n"
        f"{textsorte_hint}"
        f"SCHÜLERTEXT:\n---\n{schuelertext}\n---\n\n"
        f"HAUPTANALYSE (Zusammenfassung):\n{bewertung_json}\n\n"
        "Erstelle eine Detailbewertung nach dem SRDP-Beurteilungsraster. "
        "Für jedes der folgenden 15 Kriterien:\n"
        "- Eine Stufe (0-4)\n"
        "- Eine Begründung (2-4 Sätze mit konkreten Textbelegen)\n"
        "SKALEN-HINWEIS: Die Detailskala 0-4 ist die offizielle SRDP-Subkriterien-Skala\n"
        "(0 = nicht erfüllt … 4 = weit über das Wesentliche hinausgehend). Sie ist NICHT\n"
        "identisch mit den 1-5-Stufen des Bewertungsrasters aus der Hauptanalyse —\n"
        "übernimm deren Werte nicht.\n\n"
        "Antworte als JSON mit diesem Schema:\n"
        '{\n'
        '  "gesamteindruck": "3-5 Sätze Gesamteinschätzung",\n'
        '  "k1_inhalt": {\n'
        '    "schreibhandlung": {"stufe": 0, "begruendung": "..."},\n'
        '    "arbeitsauftraege": {"stufe": 0, "begruendung": "..."},\n'
        '    "textbeilage": {"stufe": 0, "begruendung": "..."},\n'
        '    "sachlich": {"stufe": 0, "begruendung": "..."},\n'
        '    "qualitaet": {"stufe": 0, "begruendung": "..."}\n'
        '  },\n'
        '  "k1_textstruktur": {\n'
        '    "kohaerenz": {"stufe": 0, "begruendung": "..."},\n'
        '    "bezugnahme": {"stufe": 0, "begruendung": "..."},\n'
        '    "kohaesion": {"stufe": 0, "begruendung": "..."}\n'
        '  },\n'
        '  "k3_stil": {\n'
        '    "situationsadaequat": {"stufe": 0, "begruendung": "..."},\n'
        '    "wortwahl": {"stufe": 0, "begruendung": "..."},\n'
        '    "satzstrukturen": {"stufe": 0, "begruendung": "..."},\n'
        '    "eigenstaendigkeit": {"stufe": 0, "begruendung": "..."}\n'
        '  },\n'
        '  "k3_sprachnormen": {\n'
        '    "orthografie": {"stufe": 0, "begruendung": "..."},\n'
        '    "zeichensetzung": {"stufe": 0, "begruendung": "..."},\n'
        '    "grammatik": {"stufe": 0, "begruendung": "..."}\n'
        '  },\n'
        '  "verbesserung_inhaltlich": "2-3 Sätze",\n'
        '  "verbesserung_strukturell": "2-3 Sätze",\n'
        '  "verbesserung_sprachlich": "2-3 Sätze"\n'
        '}\n\n'
        "REGELN:\n"
        "- Begründungen MÜSSEN konkrete Textzitate enthalten.\n"
        "- Verwende den SRDP-Stufenwortlaut in der Begründung.\n"
        "- Bewerte STRENG aber FAIR. Stufe 3-4 nur bei erkennbar guter Leistung.\n"
        "- Antworte NUR mit dem JSON, kein Text davor oder danach."
    )

    try:
        raw = run_llm_api(prompt, config, cancel_event=cancel_event)
    except Exception as e:
        logging.warning("SRDP-Detailbewertung fehlgeschlagen (API): %s", e)
        return None

    if raw.startswith("FEHLER"):
        logging.warning("SRDP-Detailbewertung: API-Fehler: %s", raw[:200])
        return None

    try:
        return extract_json_from_llm(raw)
    except (json.JSONDecodeError, AttributeError) as e:
        logging.warning("SRDP-Detailbewertung: JSON-Extraktion fehlgeschlagen: %s", e)
        return None


# ---------------------------------------------------------------------------
# App-seitige Notenberechnung (SRDP + Unterstufe)
# ---------------------------------------------------------------------------

_BEZ = {1: "Sehr gut", 2: "Gut", 3: "Befriedigend", 4: "Genügend", 5: "Nicht genügend"}

# Kriteriumsnamen variieren je nach Rubrik (Deutsch/Englisch). Diese Zuordnung mappt die
# vier SRDP-Hauptkriterien (kanonisch) auf die möglichen Schlüssel in `bewertung`/
# `kriterium_historie`. Single Source of Truth — wird von berechne_note_srdp() und der
# Längsschnitt-Aggregation (natascha_db.get_schueler_laengsschnitt) gemeinsam genutzt.
KRITERIUM_KEY_VARIANTS: dict[str, tuple[str, ...]] = {
    "inhalt": ("inhalt", "task_achievement", "analyse", "interpretation"),
    "textstruktur": ("textstruktur", "aufbau", "organisation_layout", "einleitung_aufbau"),
    "ausdruck": ("ausdruck", "stil_ausdruck", "lexical_range_accuracy"),
    "sprachrichtigkeit": (
        "sprachrichtigkeit",
        "normative_sprachrichtigkeit",
        "grammatical_range_accuracy",
    ),
}


def _srdp_result(
    note: int,
    bezeichnung: str,
    *,
    k1_note: int,
    k3_note: int,
    k1_stufe: float,
    k3_stufe: float,
    sonderregel: str | None,
    begruendung: str,
    quelle: str = "app",
) -> dict:
    """Einheitliches Ergebnis-Dict für die SRDP-Notenberechnung."""
    return {
        "note": note,
        "bezeichnung": bezeichnung,
        "durchschnitt": round((k1_stufe + k3_stufe) / 2, 2),
        "k1_note": k1_note,
        "k3_note": k3_note,
        "k1_schnitt": round(k1_stufe, 2),
        "k3_schnitt": round(k3_stufe, 2),
        "sonderregel": sonderregel,
        "begruendung": begruendung,
        "quelle": quelle,
    }


def parse_gewichtung(rubric_text: str) -> dict[str, float]:
    """Liest die Gewichtung aus dem '## Gewichtung'-Abschnitt eines Rubric-MD."""
    gewichtung: dict[str, float] = {}
    in_section = False
    for line in rubric_text.split("\n"):
        if line.strip().startswith("## Gewichtung"):
            in_section = True
            continue
        if in_section and line.strip().startswith("- "):
            match = re.match(r"-\s*([\w\s]+?):\s*(\d+)\s*%", line)
            if match:
                key = match.group(1).strip().lower()
                for canon in ("inhalt", "textstruktur", "ausdruck", "sprachrichtigkeit"):
                    if canon in key:
                        key = canon
                        break
                pct = int(match.group(2)) / 100
                gewichtung[key] = pct
        elif in_section and line.strip().startswith("##"):
            break
    return gewichtung or {
        "inhalt": 0.25,
        "textstruktur": 0.25,
        "ausdruck": 0.25,
        "sprachrichtigkeit": 0.25,
    }


def berechne_note_srdp(
    bewertung: dict, srdp_detail: dict | None = None
) -> dict:
    """SRDP-konforme Notenberechnung für Oberstufe.

    Skala: 1-5 (1=nicht erfüllt, 5=sehr gut).
    K1 = Inhalt + Textstruktur, K3 = Stil + Sprachnormen.
    Note = 6 - Stufe (invertiert: Stufe 5 → Note 1).
    """
    if srdp_detail:
        k1_vals: list[float] = []
        for section in ("k1_inhalt", "k1_textstruktur"):
            for entry in srdp_detail.get(section, {}).values():
                if isinstance(entry, dict) and "stufe" in entry:
                    k1_vals.append(float(entry["stufe"]))

        k3_vals: list[float] = []
        for section in ("k3_stil", "k3_sprachnormen"):
            for entry in srdp_detail.get(section, {}).values():
                if isinstance(entry, dict) and "stufe" in entry:
                    k3_vals.append(float(entry["stufe"]))

        k1_stufe = sum(k1_vals) / len(k1_vals) if k1_vals else 3.0
        k3_stufe = sum(k3_vals) / len(k3_vals) if k3_vals else 3.0
    else:
        def _get(variants: tuple) -> float:
            for k in variants:
                if k in bewertung:
                    val = bewertung[k]
                    return float(val.get("punkte", 3) if isinstance(val, dict) else val)
            return 3.0

        inhalt = _get(KRITERIUM_KEY_VARIANTS["inhalt"])
        struktur = _get(KRITERIUM_KEY_VARIANTS["textstruktur"])
        ausdruck = _get(KRITERIUM_KEY_VARIANTS["ausdruck"])
        sprache = _get(KRITERIUM_KEY_VARIANTS["sprachrichtigkeit"])
        k1_stufe = (inhalt + struktur) / 2
        k3_stufe = (ausdruck + sprache) / 2

    k1_note = max(1, min(5, round(6 - k1_stufe)))
    k3_note = max(1, min(5, round(6 - k3_stufe)))

    if k1_stufe <= 1.5:
        return _srdp_result(
            5,
            "Nicht genügend",
            k1_note=5,
            k3_note=k3_note,
            k1_stufe=k1_stufe,
            k3_stufe=k3_stufe,
            sonderregel="K1_NICHT_ERFUELLT",
            begruendung=(
                f"K1 nicht erfüllt (Stufe {k1_stufe:.2f}). "
                "Automatisch Nicht genügend gemäß SRDP."
            ),
        )

    if k3_stufe <= 1.5:
        return _srdp_result(
            5,
            "Nicht genügend",
            k1_note=k1_note,
            k3_note=5,
            k1_stufe=k1_stufe,
            k3_stufe=k3_stufe,
            sonderregel="K3_NICHT_ERFUELLT",
            begruendung=(
                f"K3/1 nicht erfüllt (Stufe {k3_stufe:.2f}). "
                "Automatisch Nicht genügend gemäß SRDP."
            ),
        )

    gesamt = round((k1_note + k3_note) / 2)
    gesamt = max(1, min(5, gesamt))

    return _srdp_result(
        gesamt,
        _BEZ[gesamt],
        k1_note=k1_note,
        k3_note=k3_note,
        k1_stufe=k1_stufe,
        k3_stufe=k3_stufe,
        sonderregel=None,
        begruendung=(
            f"K1: Note {k1_note} (Stufe {k1_stufe:.2f}), "
            f"K3/1: Note {k3_note} (Stufe {k3_stufe:.2f})."
        ),
    )


def berechne_note_unterstufe(
    bewertung: dict, gewichtung: dict | None = None
) -> dict:
    """Notenberechnung für Unterstufe. Gewichteter Durchschnitt der 4 Kriterien."""
    if not gewichtung:
        gewichtung = {
            "inhalt": 0.25,
            "textstruktur": 0.25,
            "ausdruck": 0.25,
            "sprachrichtigkeit": 0.25,
        }

    total = 0.0
    details: dict[str, float] = {}
    for key, weight in gewichtung.items():
        val = bewertung.get(key, {})
        stufe = float(val.get("punkte", 3) if isinstance(val, dict) else 3)
        total += stufe * weight
        details[key] = stufe

    note = max(1, min(5, round(6 - total)))
    detail_str = ", ".join(f"{k}: Stufe {v:.0f}" for k, v in details.items())

    return {
        "note": note,
        "bezeichnung": _BEZ[note],
        "durchschnitt": round(total, 2),
        "begruendung": f"Gewichteter Durchschnitt {total:.2f} ({detail_str}).",
        "quelle": "app",
    }


def _model_slug(provider: str, model: str) -> str:
    """Kurzes Modellkürzel für Dateinamen (max 12 Zeichen)."""
    slug = model.lower()
    slug = re.sub(r"claude-?", "", slug)
    slug = re.sub(r"[-_.]", "", slug)
    slug = re.sub(r"\d{8,}", "", slug)
    slug = slug[:12]
    return slug or provider[:8]


def archive_existing_analysis(analysis_path: "Path") -> "Path | None":
    """Verschiebt eine bestehende Analyse-JSON in .history/ bevor sie überschrieben wird."""
    from pathlib import Path as _Path
    analysis_path = _Path(analysis_path)
    if not analysis_path.exists():
        return None
    history_dir = analysis_path.parent / ".history"
    history_dir.mkdir(exist_ok=True)
    try:
        data = json.loads(analysis_path.read_text(encoding="utf-8"))
        provider = data.get("provider", "unknown")
        model = data.get("modell", "unknown")
        slug = _model_slug(provider, model)
    except Exception:
        slug = "unknown"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    archive_name = f"{analysis_path.stem}_{timestamp}_{slug}.json"
    archive_path = history_dir / archive_name
    shutil.move(str(analysis_path), str(archive_path))
    return archive_path


def build_project_paths(
    config: dict[str, Any],
    klasse: str | None = None,
    aufgabe: str | None = None,
) -> gf.ProjectPaths:
    """Baut Projektpfade für die angegebene Klasse + Aufgabe.

    Hierarchie: Aufgabe-Pfade > Klassen-Pfade > [paths]-Fallback.
    Wenn klasse/aufgabe None sind, werden aktive Werte aus config gelesen.
    """
    classes_cfg = config.get("classes", {})
    class_names = [k for k in classes_cfg if k != "active" and isinstance(classes_cfg[k], dict)]

    if class_names:
        if klasse is None:
            klasse = classes_cfg.get("active", class_names[0])
        cls = classes_cfg.get(klasse) if klasse else None
        if isinstance(cls, dict):
            # Aufgaben-Pfade ermitteln
            if aufgabe is None:
                aufgabe = cls.get("active_aufgabe")
            auf_cfg: dict[str, Any] = {}
            if aufgabe:
                auf_cfg = cls.get("aufgaben", {}).get(aufgabe, {})

            input_rel = auf_cfg.get("input") or cls.get("input", "input")
            output_rel = auf_cfg.get("output") or cls.get("output", "output")
            output_dir = PROJECT_ROOT / output_rel
            return gf.ProjectPaths(
                root=PROJECT_ROOT,
                input_dir=PROJECT_ROOT / input_rel,
                output_dir=output_dir,
                feedback_data_dir=output_dir / "feedback_data",
                fehlerlog=output_dir / "fehlerlog.txt",
            )

    # Fallback: klassischer [paths]-Block
    output_dir = resolve_path(config, "output")
    return gf.ProjectPaths(
        root=PROJECT_ROOT,
        input_dir=resolve_path(config, "input"),
        output_dir=output_dir,
        feedback_data_dir=resolve_path(config, "feedback_data"),
        fehlerlog=resolve_path(config, "fehlerlog"),
    )


_AUSGANGSTEXT_EXTENSIONS: frozenset[str] = frozenset({".docx", ".pdf", ".jpg", ".jpeg", ".png"})


def detect_ausgangstext(config: dict[str, Any], klasse: str, aufgabe: str) -> Path | None:
    """Sucht die Ausgangstext-Datei im Unterordner ausgangstext/ der Aufgabe.

    Unterstützte Formate: DOCX, PDF, JPG, JPEG, PNG.
    Gibt die erste Datei (alphabetisch) zurück, oder None wenn nicht gefunden.
    Fehlende Ordner werden stillschweigend ignoriert.
    """
    paths = build_project_paths(config, klasse, aufgabe)
    ausgangstext_dir = paths.input_dir / "ausgangstext"
    if not ausgangstext_dir.is_dir():
        return None
    candidates = sorted(
        f for f in ausgangstext_dir.iterdir()
        if f.is_file() and f.suffix.lower() in _AUSGANGSTEXT_EXTENSIONS
    )
    return candidates[0] if candidates else None


def list_classes(config: dict[str, Any]) -> list[str]:
    """Gibt alle konfigurierten Klassennamen zurück (ohne 'active')."""
    classes_cfg = config.get("classes", {})
    return [k for k in classes_cfg if k != "active" and isinstance(classes_cfg[k], dict)]


def active_klasse(config: dict[str, Any]) -> str | None:
    """Gibt den aktiven Klassennamen zurück, oder None wenn keine Klassen konfiguriert."""
    names = list_classes(config)
    if not names:
        return None
    return config.get("classes", {}).get("active", names[0])


def save_active_klasse(klasse: str) -> None:
    """Schreibt die aktive Klasse in natascha_config.toml."""
    doc = _load_toml_doc()
    doc["classes"]["active"] = klasse
    _save_toml_doc(doc)


def add_class_to_config(name: str, input_rel: str, output_rel: str) -> None:
    """Fügt eine neue Klasse in natascha_config.toml ein und setzt sie als aktiv."""
    doc = _load_toml_doc()
    doc["classes"]["active"] = name
    new_cls = tomlkit.table()
    new_cls.add("input", input_rel)
    new_cls.add("output", output_rel)
    doc["classes"].add(name, new_cls)
    _save_toml_doc(doc)


def list_aufgaben(config: dict[str, Any], klasse: str) -> list[str]:
    """Gibt Slug-Liste aller Aufgaben einer Klasse zurück (Reihenfolge wie in config)."""
    cls = config.get("classes", {}).get(klasse, {})
    aufgaben = cls.get("aufgaben", {})
    return [k for k in aufgaben if isinstance(aufgaben[k], dict)]


def active_aufgabe(config: dict[str, Any], klasse: str) -> str | None:
    """Gibt die aktive Aufgabe einer Klasse zurück, oder None wenn keine vorhanden."""
    names = list_aufgaben(config, klasse)
    if not names:
        return None
    cls = config.get("classes", {}).get(klasse, {})
    return cls.get("active_aufgabe", names[-1])


def get_aufgabe_cfg(config: dict[str, Any], klasse: str, aufgabe: str) -> dict[str, Any]:
    """Gibt die Konfigurations-Dict einer Aufgabe zurück (oder {})."""
    cls = config.get("classes", {}).get(klasse, {})
    return cls.get("aufgaben", {}).get(aufgabe, {})


def load_erwartungshorizont(config: dict[str, Any], klasse: str, aufgabe: str) -> str | None:
    """Lädt den Erwartungshorizont für eine Aufgabe, falls konfiguriert."""
    auf_cfg = get_aufgabe_cfg(config, klasse, aufgabe)
    eh_name = auf_cfg.get("erwartungshorizont", "")
    if not eh_name:
        return None
    eh_path = resolve_path(config, "rubrics") / eh_name
    if not eh_path.exists():
        return None
    return eh_path.read_text(encoding="utf-8")


def _erwartungshorizont_block(
    config: dict[str, Any], klasse: str, aufgabe: str, override_name: str = ""
) -> str:
    """Gibt den Erwartungshorizont-Block für den Prompt zurück, oder '' wenn nicht konfiguriert."""
    if override_name:
        eh_path = resolve_path(config, "rubrics") / override_name
        erwartung = eh_path.read_text(encoding="utf-8") if eh_path.exists() else None
    elif klasse and aufgabe:
        erwartung = load_erwartungshorizont(config, klasse, aufgabe)
    else:
        return ""
    if not erwartung:
        return ""
    return (
        "ERWARTUNGSHORIZONT (aufgabenspezifisch):\n---\n"
        f"{erwartung}\n---\n\n"
        "HINWEIS: Der Erwartungshorizont definiert konkrete inhaltliche Erwartungen pro Operator. "
        "Prüfe den Schülertext gegen diese Erwartungen. Bewerte insbesondere:\n"
        "- Welche der erwarteten Inhalte pro Operator vorhanden sind\n"
        "- Ob typische Fehler aus dem Erwartungshorizont auftreten\n"
        "- Ob die Argumentationslinien nachvollziehbar und im Rahmen des Akzeptablen sind\n\n"
    )


def save_erwartungshorizont_to_config(klasse: str, aufgabe: str, eh_filename: str) -> None:
    """Trägt den EH-Dateinamen in natascha_config.toml ein."""
    doc = _load_toml_doc()
    doc["classes"][klasse]["aufgaben"][aufgabe]["erwartungshorizont"] = eh_filename
    _save_toml_doc(doc)


def generate_erwartungshorizont(
    config: dict[str, Any],
    klasse: str,
    aufgabe: str,
    provider: str = "",
    model: str = "",
    cancel_event: threading.Event | None = None,
) -> str:
    """Generiert einen Erwartungshorizont aus Ausgangstext + Rubrik + Textsorte.

    Gibt das generierte Markdown zurück. Wirft ValueError wenn Voraussetzungen fehlen,
    RuntimeError wenn der LLM-Call fehlschlägt.
    """
    auf_cfg = get_aufgabe_cfg(config, klasse, aufgabe)
    textsorte = auf_cfg.get("textsorte", "")
    situation = auf_cfg.get("situation", "")

    ausgangstext_path = detect_ausgangstext(config, klasse, aufgabe)
    if not ausgangstext_path:
        raise ValueError(
            "Kein Ausgangstext gefunden. Bitte zuerst in ausgangstext/ ablegen."
        )

    suf = ausgangstext_path.suffix.lower()
    if suf == ".docx":
        ausgangstext_text = read_docx_text(ausgangstext_path)
    elif suf in (".txt", ".md"):
        ausgangstext_text = ausgangstext_path.read_text(encoding="utf-8")
    else:
        raise ValueError(
            f"Ausgangstext-Format nicht unterstützt: {suf}. "
            "Bitte als .docx oder .txt ablegen."
        )

    rubric = load_rubric_for_aufgabe(config, klasse, aufgabe)

    prompt_template_path = PROJECT_ROOT / "prompts" / "PROMPT_ERWARTUNGSHORIZONT.md"
    prompt_template = (
        prompt_template_path.read_text(encoding="utf-8")
        if prompt_template_path.exists()
        else ""
    )

    prompt = (
        "Du bist eine erfahrene AHS-Deutschlehrerin in Österreich.\n"
        "Erstelle einen Erwartungshorizont im Markdown-Format.\n\n"
        f"TEXTSORTE: {textsorte}\n"
    )
    if situation:
        prompt += f"SITUATION: {situation}\n"
    prompt += (
        f"\nBEWERTUNGSRASTER:\n---\n{rubric}\n---\n\n"
        f"AUSGANGSTEXT + AUFGABENSTELLUNG:\n---\n{ausgangstext_text}\n---\n\n"
        "AUSGABEFORMAT:\n"
        "Antworte ausschließlich im Markdown-Format.\n"
        "Leite die erwarteten Inhalte AUS DER TEXTBEILAGE ab.\n"
        "Nenne Pro- UND Kontra-Argumente.\n"
        "Nenne typische Fehler pro Operator.\n"
        "Verwende korrekte Umlaute (ä, ö, ü, ß).\n"
    )
    if prompt_template:
        prompt += f"\nVORLAGE (Struktur):\n---\n{prompt_template}\n---\n"

    eh_cfg = config.get("erwartungshorizont", {})
    prov = provider or eh_cfg.get("default_provider", "")
    mod = model or eh_cfg.get("default_model", "")

    api_cfg = dict(config.get("api", {}))
    if prov:
        api_cfg["provider"] = prov
    if mod:
        api_cfg["model"] = mod
    eh_config = {**config, "api": api_cfg}

    response = run_llm_api(prompt, eh_config, cancel_event=cancel_event)
    if response.startswith("FEHLER"):
        raise RuntimeError(response)
    return response


def aufgabe_defaults(
    config: dict[str, Any], klasse: str | None, aufgabe: str | None
) -> dict[str, str]:
    """Gibt fach/schulstufe/textsorte/rubric der Aufgabe zurück.

    Leere Strings wenn nicht konfiguriert — Caller fällt dann auf globale Defaults zurück.
    """
    if not klasse or not aufgabe:
        return {}
    auf_cfg = get_aufgabe_cfg(config, klasse, aufgabe)
    return {
        "fach": auf_cfg.get("fach", ""),
        "schulstufe": auf_cfg.get("schulstufe", ""),
        "textsorte": auf_cfg.get("textsorte", ""),
        "rubric": auf_cfg.get("rubric", ""),
        "bewertungsmodus": auf_cfg.get("bewertungsmodus", "benotet"),
        "erwartungshorizont": auf_cfg.get("erwartungshorizont", ""),
    }


def save_active_aufgabe(klasse: str, aufgabe: str) -> None:
    """Schreibt active_aufgabe für die Klasse in natascha_config.toml."""
    doc = _load_toml_doc()
    doc["classes"][klasse]["active_aufgabe"] = aufgabe
    _save_toml_doc(doc)


def add_aufgabe_to_config(
    klasse: str,
    slug: str,
    label: str,
    fach: str,
    schulstufe: str,
    textsorte: str,
    rubric: str,
) -> None:
    """Fügt eine neue Aufgabe zur Klasse in natascha_config.toml hinzu.

    Setzt die neue Aufgabe sofort als active_aufgabe und erstellt die Ordner.
    """
    doc = _load_toml_doc()
    cls_doc = doc["classes"][klasse]
    cls_doc["active_aufgabe"] = slug

    base_input = str(cls_doc.get("input", f"input/{klasse}"))
    base_output = str(cls_doc.get("output", f"output/{klasse}"))
    input_rel = f"{base_input}/{slug}"
    output_rel = f"{base_output}/{slug}"

    if "aufgaben" not in cls_doc:
        cls_doc.add("aufgaben", tomlkit.table())

    auf_tbl = tomlkit.table()
    auf_tbl.add("label", label)
    auf_tbl.add("fach", fach)
    auf_tbl.add("schulstufe", schulstufe)
    auf_tbl.add("textsorte", textsorte)
    auf_tbl.add("rubric", rubric)
    auf_tbl.add("input", input_rel)
    auf_tbl.add("output", output_rel)
    cls_doc["aufgaben"].add(slug, auf_tbl)

    _save_toml_doc(doc)

    # Ordner anlegen
    (PROJECT_ROOT / input_rel).mkdir(parents=True, exist_ok=True)
    (PROJECT_ROOT / output_rel).mkdir(parents=True, exist_ok=True)


def log_tui_error(paths: gf.ProjectPaths, message: str) -> None:
    paths.fehlerlog.parent.mkdir(parents=True, exist_ok=True)
    with paths.fehlerlog.open("a", encoding="utf-8") as f:
        f.write(f"[TUI] {message}\n")


def rubric_options_for(
    fach: str,
    schulstufe: str,
    config: dict[str, Any],
    current_rubric: str = "",
) -> list[str]:
    """Gibt alle verfügbaren Rubrik-Dateien zurück, gefiltert nach Schulstufe.

    Schulstufen-spezifische Rubriken werden getrennt; neutrale/custom Rubriken
    bleiben in beiden Stufen auswählbar. Die aktuell zugewiesene Rubrik
    erscheint immer in der Liste, auch wenn sie nicht zur Schulstufe passt.
    """
    rubric_dir = resolve_path(config, "rubrics")
    _preferred: dict[tuple[str, str], list[str]] = {
        ("Deutsch", "Oberstufe"): ["srdp_deutsch_oberstufe.md"],
        ("Deutsch", "Unterstufe"): ["deutsch_unterstufe.md"],
        ("Englisch", "Unterstufe"): ["englisch_a2.md"],
        ("Englisch", "Oberstufe"): ["srdp_englisch_b2.md", "srdp_englisch_b1.md"],
    }
    preferred = _preferred.get((fach, schulstufe), [])
    all_rubrics = sorted(
        f.name
        for f in rubric_dir.glob("*.md")
        if not f.name.upper().startswith("README")
        and not f.name.startswith("erwartungshorizont_")
    )

    def _stage(filename: str) -> str:
        name = filename.lower()
        if name == "englisch_a2.md" or "_unterstufe" in name:
            return "unterstufe"
        if name.startswith("srdp_") or "_oberstufe" in name:
            return "oberstufe"
        return "generic"

    stufe_lower = schulstufe.lower()
    if stufe_lower == "unterstufe":
        filtered = [f for f in all_rubrics if _stage(f) in ("unterstufe", "generic")]
    elif stufe_lower == "oberstufe":
        filtered = [f for f in all_rubrics if _stage(f) in ("oberstufe", "generic")]
    else:
        filtered = all_rubrics

    result = [f for f in preferred if f in filtered]
    result += [f for f in filtered if f not in result]

    if current_rubric and current_rubric not in result:
        rubric_path = rubric_dir / current_rubric
        if rubric_path.exists():
            result.insert(0, current_rubric)

    return result


def default_rubric_for(fach: str, schulstufe: str, config: dict[str, Any]) -> str:
    mapping = config.get("rubric_mapping", {})
    configured = mapping.get(f"{fach}+{schulstufe}")
    options = rubric_options_for(fach, schulstufe, config)
    if configured in options:
        return configured
    return options[0] if options else ""


def copy_to_clipboard(text: str) -> bool:
    for cmd in [
        "clip.exe",
        "xclip -selection clipboard",
        "xsel --clipboard",
        "wl-copy",
    ]:
        parts = cmd.split()
        try:
            proc = subprocess.run(
                parts,
                input=text.encode(),
                timeout=5,
                capture_output=True,
            )
            if proc.returncode == 0:
                return True
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    return False


def save_settings(fach: str, schulstufe: str, provider: str, model: str) -> None:
    """Schreibt API- und Default-Einstellungen in natascha_config.toml."""
    doc = _load_toml_doc()
    doc["defaults"]["fach"] = fach
    doc["defaults"]["schulstufe"] = schulstufe
    doc["api"]["provider"] = provider
    doc["api"]["model"] = model
    _save_toml_doc(doc)


def open_file(path: Path) -> bool:
    """Öffnet eine Datei oder einen Ordner mit der passenden Windows-App.

    Konvertiert WSL-Pfade nach Windows und nutzt:
    - Sublime Text für .md-Dateien
    - explorer.exe für Verzeichnisse
    - cmd.exe /c start für alle anderen Dateien (Windows-Standardprogramm)
    Fällt auf wslview / xdg-open zurück falls wslpath nicht verfügbar.
    """
    path = path.resolve()

    # WSL → Windows-Pfad
    try:
        result = subprocess.run(
            ["wslpath", "-w", str(path)],
            capture_output=True, text=True, timeout=5,
        )
        win_path = result.stdout.strip()
    except (FileNotFoundError, subprocess.TimeoutExpired):
        win_path = None

    if win_path:
        try:
            if path.is_dir():
                subprocess.Popen(
                    ["explorer.exe", win_path],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                )
            elif path.suffix.lower() == ".md":
                sublime = r"C:\Program Files\Sublime Text\subl.exe"
                subprocess.Popen(
                    ["cmd.exe", "/c", sublime, win_path],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                )
            else:
                subprocess.Popen(
                    ["cmd.exe", "/c", "start", "", win_path],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                )
            return True
        except (FileNotFoundError, OSError):
            pass

    # Fallback: wslview / xdg-open / open
    for cmd in [["wslview", str(path)], ["xdg-open", str(path)], ["open", str(path)]]:
        try:
            subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            return True
        except FileNotFoundError:
            continue
    return False


def run_agent_sync(cmd_template: str, prompt: str, timeout: int) -> str:
    parts = cmd_template.split()
    if os.name == "nt" and parts == ["cat"]:
        parts = ["cmd", "/C", "more"]
    try:
        result = subprocess.run(
            parts,
            input=prompt,
            capture_output=True,
            text=True,
            timeout=timeout,
        )
        if result.returncode != 0:
            return f"FEHLER (Exit {result.returncode}): {result.stderr[:500]}"
        return result.stdout
    except subprocess.TimeoutExpired:
        return f"FEHLER: Timeout nach {timeout}s"
    except FileNotFoundError as e:
        return f"FEHLER: Befehl nicht gefunden: {e}"


def _call_ollama_native(
    base_url: str,
    model: str,
    prompt: str,
    timeout: int,
    cancel_event: threading.Event | None = None,
) -> str:
    """Ruft Ollama über die native /api/chat-Schnittstelle auf."""
    if cancel_event is not None and cancel_event.is_set():
        return "FEHLER: Abgebrochen"
    url = base_url.rstrip("/") + "/api/chat"
    payload = json.dumps(
        {
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "think": False,
            "stream": False,
            "options": {"num_ctx": 32768},
        }
    ).encode("utf-8")
    req = urllib.request.Request(
        url, data=payload, headers={"Content-Type": "application/json"}, method="POST"
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            data = json.loads(resp.read())
        return data["message"]["content"]
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")
        return f"FEHLER: HTTP {e.code}: {body[:400]}"
    except Exception as e:
        if cancel_event is not None and cancel_event.is_set():
            return "FEHLER: Abgebrochen"
        return f"FEHLER: Ollama-Aufruf fehlgeschlagen: {e}"


def _call_openai_compat(
    base_url: str,
    api_key: str,
    model: str,
    prompt: str,
    timeout: int,
    extra_body: dict | None = None,
    cancel_event: threading.Event | None = None,
    vision_content: list[dict] | None = None,
    max_tokens: int = 4096,
) -> str:
    """Ruft eine OpenAI-kompatible Chat-API auf (GLM, Kimi, OpenAI, Ollama).

    vision_content: Anthropic-formatierte Content-Blöcke (image-Typ).
    Wird in OpenAI image_url-Format konvertiert. PDF-Blöcke werden übersprungen
    (OpenAI unterstützt kein natives PDF).
    """
    if cancel_event is not None and cancel_event.is_set():
        return "FEHLER: Abgebrochen"
    url = base_url.rstrip("/") + "/chat/completions"

    if vision_content:
        openai_parts: list[dict] = []
        for block in vision_content:
            if block.get("type") == "image":
                src = block["source"]
                data_url = f"data:{src['media_type']};base64,{src['data']}"
                openai_parts.append({"type": "image_url", "image_url": {"url": data_url}})
            # PDF-Blöcke (type == "document") werden nicht konvertiert
        openai_parts.append({"type": "text", "text": prompt})
        message_content: str | list = openai_parts
    else:
        message_content = prompt

    body: dict = {
        "model": model,
        "messages": [{"role": "user", "content": message_content}],
    }
    if base_url.rstrip("/") == "https://api.openai.com/v1" and (
        model.startswith("gpt-5") or model.startswith("o")
    ):
        body["max_completion_tokens"] = max_tokens
    else:
        body["max_tokens"] = max_tokens
    if extra_body:
        body.update(extra_body)
    payload = json.dumps(body).encode("utf-8")
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    req = urllib.request.Request(url, data=payload, headers=headers, method="POST")
    retry_status = {429, 503}
    max_retries = 3
    for attempt in range(max_retries + 1):
        try:
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                resp_data = json.loads(resp.read())
            choice = (resp_data.get("choices") or [{}])[0]
            message = choice.get("message", {}) or {}
            finish_reason = choice.get("finish_reason", "")
            content = message.get("content", "") or ""
            if not str(content).strip() and finish_reason == "length":
                return (
                    "FEHLER: Antwort abgeschnitten (finish_reason=length) — "
                    "max_tokens zu niedrig (Reasoning verbraucht das Budget)."
                )
            if not str(content).strip():
                reasoning_content = message.get("reasoning_content", "") or ""
                if str(reasoning_content).strip():
                    return reasoning_content
                return "FEHLER: Leere Antwort von OpenAI-kompatibler API"
            return content
        except urllib.error.HTTPError as e:
            err_body = e.read().decode("utf-8", errors="replace")
            should_retry = e.code in retry_status and attempt < max_retries
            if should_retry:
                time.sleep(2**(attempt + 1))
                continue
            return f"FEHLER: HTTP {e.code}: {err_body[:400]}"
        except Exception as e:
            if cancel_event is not None and cancel_event.is_set():
                return "FEHLER: Abgebrochen"
            return f"FEHLER: API-Aufruf fehlgeschlagen: {e}"


def _should_retry_with_stable_model(response: str) -> bool:
    """Erkennt typische Provider-Antworten fuer nicht verfuegbare Modellnamen."""
    if not response.startswith("FEHLER:"):
        return False
    lowered = response.lower()
    markers = (
        "model_not_found",
        "model not found",
        "does not exist",
        "not found",
        "invalid model",
        "model_not_supported",
        "unsupported model",
        "no access",
        "not available",
    )
    return any(marker in lowered for marker in markers)


def _with_model_fallback(
    provider: str,
    model: str,
    fallback_model: str,
    call: Any,
) -> str:
    response = call(model)
    if model != fallback_model and _should_retry_with_stable_model(response):
        fallback_response = call(fallback_model)
        if not fallback_response.startswith("FEHLER:"):
            logging.warning(
                "Modell '%s' (%s) nicht verfuegbar; automatisch mit '%s' wiederholt.",
                model,
                provider,
                fallback_model,
            )
            return fallback_response
    return response


def api_key_available(config: dict[str, Any]) -> bool:
    """Prüft ob der API-Key für den konfigurierten Provider gesetzt ist."""
    provider = config.get("api", {}).get("provider", "anthropic").lower()
    key_map = {
        "anthropic": "ANTHROPIC_API_KEY",
        "deepseek": "DEEPSEEK_API_KEY",
        "qwen": "QWEN_API_KEY",
        "kimi": "KIMI_API_KEY",
        "mistral": "MISTRAL_API_KEY",
        "openai": "OPENAI_API_KEY",
        "ollama": None,  # kein Key nötig
    }
    env_key = key_map.get(provider)
    if env_key is None:
        return True
    if provider == "qwen":
        return bool(os.environ.get("QWEN_API_KEY", "") or os.environ.get("DASHSCOPE_API_KEY", ""))
    return bool(os.environ.get(env_key, ""))


def run_llm_api(
    prompt: str,
    config: dict[str, Any],
    cancel_event: threading.Event | None = None,
    schema: dict[str, Any] | None = None,
    vision_content: list[dict] | None = None,
) -> str:
    """Dispatcht den LLM-Aufruf je nach konfiguriertem Provider.

    vision_content: Liste von Anthropic-Content-Blöcken (image/document).
    Wird an vision-fähige Provider weitergereicht; nicht-vision-fähige ignorieren es.
    """
    if cancel_event is not None and cancel_event.is_set():
        return "FEHLER: Abgebrochen"
    api_cfg = config.get("api", {})
    provider = api_cfg.get("provider", "anthropic").lower()
    timeout = config.get("agent", {}).get("timeout_seconds", 120)
    model = api_cfg.get("model", "")

    if provider == "anthropic":
        return run_anthropic_api(
            prompt, model or "claude-sonnet-4-6", timeout, cancel_event=cancel_event,
            schema=schema, vision_content=vision_content,
        )

    if provider == "deepseek":
        api_key = os.environ.get("DEEPSEEK_API_KEY", "")
        if not api_key:
            return "FEHLER: DEEPSEEK_API_KEY nicht gesetzt (.env prüfen)"
        selected_model = model or "deepseek-chat"
        return _with_model_fallback(
            provider,
            selected_model,
            "deepseek-chat",
            lambda active_model: _call_openai_compat(
                "https://api.deepseek.com/v1",
                api_key,
                active_model,
                prompt,
                timeout,
                cancel_event=cancel_event,
                extra_body=None if "reasoner" in active_model.lower() else {
                    "response_format": {"type": "json_object"}
                },
                max_tokens=8192 if "reasoner" in active_model.lower() else 4096,
            ),
        )

    if provider == "qwen":
        api_key = os.environ.get("QWEN_API_KEY", "") or os.environ.get("DASHSCOPE_API_KEY", "")
        if not api_key:
            return "FEHLER: QWEN_API_KEY oder DASHSCOPE_API_KEY nicht gesetzt (.env prüfen)"
        base_url = os.environ.get(
            "QWEN_BASE_URL",
            os.environ.get("DASHSCOPE_BASE_URL", "https://dashscope-intl.aliyuncs.com/compatible-mode/v1"),
        )
        selected_model = model or "qwen-plus"
        return _with_model_fallback(
            provider,
            selected_model,
            "qwen-plus",
            lambda active_model: _call_openai_compat(
                base_url,
                api_key,
                active_model,
                prompt,
                timeout,
                extra_body={"enable_thinking": False},
                cancel_event=cancel_event,
                vision_content=vision_content,
            ),
        )

    if provider == "mistral":
        api_key = os.environ.get("MISTRAL_API_KEY", "")
        if not api_key:
            return "FEHLER: MISTRAL_API_KEY nicht gesetzt (.env prüfen)"
        base_url = os.environ.get("MISTRAL_BASE_URL", "https://api.mistral.ai/v1")
        selected_model = model or "mistral-small-latest"
        return _with_model_fallback(
            provider,
            selected_model,
            "mistral-small-latest",
            lambda active_model: _call_openai_compat(
                base_url,
                api_key,
                active_model,
                prompt,
                timeout,
                extra_body={"response_format": {"type": "json_object"}},
                cancel_event=cancel_event,
            ),
        )

    if provider == "kimi":
        api_key = os.environ.get("KIMI_API_KEY", "")
        if not api_key:
            return "FEHLER: KIMI_API_KEY nicht gesetzt (.env prüfen)"
        base_url = os.environ.get("KIMI_BASE_URL", "https://api.moonshot.ai/v1")
        return _call_openai_compat(
            base_url, api_key, model or "moonshot-v1-128k", prompt, timeout,
            cancel_event=cancel_event,
        )

    if provider == "openai":
        api_key = os.environ.get("OPENAI_API_KEY", "")
        if not api_key:
            return "FEHLER: OPENAI_API_KEY nicht gesetzt (.env prüfen)"
        selected_model = model or "gpt-4o"
        return _with_model_fallback(
            provider,
            selected_model,
            "gpt-4o",
            lambda active_model: _call_openai_compat(
                "https://api.openai.com/v1", api_key, active_model, prompt, timeout,
                cancel_event=cancel_event, vision_content=vision_content,
            ),
        )

    if provider == "ollama":
        base_url = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")
        return _call_ollama_native(
            base_url, model or "qwen3.5:27b", prompt, timeout, cancel_event=cancel_event,
        )

    return (
        "FEHLER: Unbekannter Provider '"
        f"{provider}' — erlaubt: anthropic, deepseek, qwen, mistral, kimi, openai, ollama"
    )


def run_anthropic_api(
    prompt: str,
    model: str,
    timeout: int = 120,
    cancel_event: threading.Event | None = None,
    schema: dict[str, Any] | None = None,
    vision_content: list[dict] | None = None,
) -> str:
    """Ruft die Anthropic Messages API auf.

    vision_content: Liste von Anthropic-Content-Blöcken (image/document).
    Bei PDF-Blöcken wird betas=["pdfs-2024-09-25"] automatisch gesetzt.
    """
    if cancel_event is not None and cancel_event.is_set():
        return "FEHLER: Abgebrochen"
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return "FEHLER: ANTHROPIC_API_KEY nicht gesetzt"
    try:
        import anthropic
    except ImportError:
        return "FEHLER: anthropic-Paket nicht installiert (pip install anthropic)"
    client = anthropic.Anthropic(api_key=api_key)

    needs_pdf_beta = vision_content and any(b.get("type") == "document" for b in vision_content)
    if vision_content:
        message_content: str | list = list(vision_content) + [{"type": "text", "text": prompt}]
    else:
        message_content = prompt

    call_kwargs: dict[str, Any] = dict(
        model=model,
        max_tokens=4096,
        messages=[{"role": "user", "content": message_content}],
        timeout=timeout,
    )
    if needs_pdf_beta:
        call_kwargs["extra_headers"] = {"anthropic-beta": "pdfs-2024-09-25"}
    if schema:
        tool_schema = _inline_schema_refs(schema)
        call_kwargs["tools"] = [{
            "name": "feedback_result",
            "description": "Bewertungsergebnis der Schülerarbeit als strukturiertes JSON",
            "input_schema": tool_schema,
        }]
        call_kwargs["tool_choice"] = {"type": "tool", "name": "feedback_result"}

    try:
        response = client.messages.create(**call_kwargs)
        if schema:
            for block in response.content:
                if block.type == "tool_use":
                    result = dict(block.input)
                    # Claude gibt verschachtelte Objekte gelegentlich als JSON-String zurück
                    for key in ("bewertung", "notenempfehlung"):
                        if isinstance(result.get(key), str):
                            try:
                                result[key] = json.loads(result[key])
                            except (json.JSONDecodeError, TypeError) as _repair_err:
                                logging.warning(
                                    "tool_use Repair fehlgeschlagen für '%s': %s",
                                    key, _repair_err,
                                )
                    return json.dumps(result, ensure_ascii=False)
            return "FEHLER: Kein Tool-Use-Block in Antwort"
        if response.content and len(response.content) > 0:
            return response.content[0].text
        return "FEHLER: Leere Antwort"
    except Exception as e:
        if cancel_event is not None and cancel_event.is_set():
            return "FEHLER: Abgebrochen"
        return f"FEHLER: API-Aufruf fehlgeschlagen: {e}"


def check_agent_availability(config: dict[str, Any]) -> dict[str, bool]:
    commands = config.get("agent", {}).get("commands", {})
    availability: dict[str, bool] = {}
    for name, cmd in commands.items():
        binary = cmd.split()[0]
        try:
            result = subprocess.run(
                ["which", binary],
                capture_output=True,
                timeout=2,
            )
            availability[name] = result.returncode == 0
        except (FileNotFoundError, subprocess.TimeoutExpired):
            availability[name] = False
    return availability


def humanize_agent_error(error: str, agent_name: str) -> str:
    if "nicht gefunden" in error or "not found" in error.lower():
        return (
            f"Der Agent {agent_name!r} ist nicht installiert.\n"
            f"  Tipp: Nutze einen anderen Agent oder den Zwischenablage-Modus."
        )
    if "Timeout" in error or "timeout" in error.lower():
        return (
            f"Der Agent {agent_name!r} hat zu lange gebraucht.\n"
            f"  Tipp: Erhoehe den Timeout in den Einstellungen oder nutze einen anderen Agent."
        )
    if "API" in error and "key" in error.lower():
        return (
            "API-Schluessel fehlt oder ist ungueltig.\n"
            "  Tipp: Setze ANTHROPIC_API_KEY und pruefe die Verbindung im Einstellungsmenue."
        )
    return error


def compute_statistics(analyses: list[dict[str, Any]]) -> dict[str, Any]:
    """Berechnet Klassen-Statistiken aus einer Liste von Analyse-Dicts.

    Returns:
        {
            "total": int,                         # Anzahl ausgewerteter Dateien
            "grade_distribution": {1..5: int},    # Anzahl pro Note
            "grade_average": float,               # Gewichteter Gesamtdurchschnitt
            "criteria_averages": {                # Pro Kriterium
                key: {"avg": float, "count": int, "min": float, "max": float}
            },
            "weakest_criterion": str | None,      # Kriterium mit niedrigstem Ø
            "strongest_criterion": str | None,    # Kriterium mit höchstem Ø
        }
    """
    grade_dist: dict[int, int] = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
    criteria_scores: dict[str, list[float]] = {}

    for data in analyses:
        note_data = data.get("notenempfehlung", {})
        note = note_data.get("note")
        if isinstance(note, (int, float)) and 1 <= int(note) <= 5:
            grade_dist[int(note)] += 1

        bewertung = data.get("bewertung", {})
        for key, crit in bewertung.items():
            if isinstance(crit, dict):
                punkte = crit.get("punkte")
                if isinstance(punkte, (int, float)):
                    criteria_scores.setdefault(key, []).append(float(punkte))

    total = sum(grade_dist.values())
    grade_average = (
        round(sum(n * c for n, c in grade_dist.items()) / total, 2) if total > 0 else 0.0
    )

    criteria_averages = {
        key: {
            "avg": round(sum(scores) / len(scores), 2),
            "count": len(scores),
            "min": min(scores),
            "max": max(scores),
        }
        for key, scores in criteria_scores.items()
    }

    weakest = (
        min(criteria_averages, key=lambda k: criteria_averages[k]["avg"])
        if criteria_averages
        else None
    )
    strongest = (
        max(criteria_averages, key=lambda k: criteria_averages[k]["avg"])
        if criteria_averages
        else None
    )

    return {
        "total": total,
        "grade_distribution": grade_dist,
        "grade_average": grade_average,
        "criteria_averages": criteria_averages,
        "weakest_criterion": weakest,
        "strongest_criterion": strongest,
    }


def compute_class_progress(config: dict[str, Any], klasse: str) -> list[dict[str, Any]]:
    """Berechnet Lernfortschritt ueber alle Aufgaben einer Klasse hinweg.

    Returns:
        Liste von Dicts (eins pro Aufgabe mit Analysen), sortiert nach
        Config-Reihenfolge:
            [{"aufgabe": slug, "label": ..., "avg_note": float,
              "avg_criteria": {key: float}, "n": int}, ...]
    """
    aufgaben = list_aufgaben(config, klasse)
    progress: list[dict[str, Any]] = []

    for slug in aufgaben:
        auf_cfg = get_aufgabe_cfg(config, klasse, slug)
        label = auf_cfg.get("label", slug)
        paths = build_project_paths(config, klasse, slug)
        fb_dir = paths.feedback_data_dir

        if not fb_dir.exists():
            continue

        analyses: list[dict[str, Any]] = []
        for json_path in sorted(fb_dir.glob("*.json")):
            try:
                data = json.loads(json_path.read_text(encoding="utf-8"))
                if isinstance(data, dict) and "notenempfehlung" in data:
                    analyses.append(data)
            except (json.JSONDecodeError, OSError):
                continue

        if not analyses:
            continue

        stats = compute_statistics(analyses)

        criteria_avgs = {key: vals["avg"] for key, vals in stats["criteria_averages"].items()}

        progress.append(
            {
                "aufgabe": slug,
                "label": label,
                "avg_note": stats["grade_average"],
                "avg_criteria": criteria_avgs,
                "n": stats["total"],
            }
        )

    return progress


def build_schueler_profil_prompt(laengsschnitt: dict[str, Any]) -> str:
    """Baut einen DATENMINIMIERTEN Prompt für ein LLM-Schülerprofil.

    ENTWURF — Schicht 3 des dreischichtigen Längsschnitt-Systems. Diese Funktion baut
    NUR den Prompt-String und gibt ihn zurück: KEIN API-Call, KEINE UI-Verdrahtung, kein
    automatischer Aufruf. Sie dient als Diskussionsgrundlage.

    DSGVO-KRITISCH: Dieser Prompt darf NIEMALS enthalten:
    - Vorname, Nachname, Klasse, Schüler-ID, Dateinamen
    - zusammenhängende Schülertexte oder ganze Sätze
    Erlaubt sind NUR: aggregierte Stufen/Noten, Fehlertyp-Häufigkeiten und die kurzen
    isolierten Fehlerzitate (max. 6 Wörter), die ohnehin schon in der DB stehen. Der Name
    wird erst NACH der LLM-Antwort lokal wieder ans Profil geheftet.

    Diese Funktion liest bewusst NUR `verlauf`, `trend` und `fehlerschwerpunkte` aus dem
    Aggregat — niemals `laengsschnitt["schueler"]`. Ein DSGVO-Regressionstest sichert ab,
    dass kein Name/keine Klasse im erzeugten Prompt landet.

    Erwarteter LLM-Output (im Prompt spezifiziert): JSON mit
        staerken[], schwaechen[], foerderbereiche[] (je mit konkreter, auf die
        SRDP-Kategorien bezogener Übungsempfehlung), maturabezug (kurzer Text, nur Oberstufe).
    """
    verlauf = laengsschnitt.get("verlauf", [])
    trend = laengsschnitt.get("trend", {})
    fehlerschwerpunkte = laengsschnitt.get("fehlerschwerpunkte", [])
    anzahl = laengsschnitt.get("anzahl_abgaben", len(verlauf))

    def _fmt(x: object) -> str:
        return "—" if x is None else (f"{x:.1f}" if isinstance(x, float) else str(x))

    def _trendzeile(key: str, titel: str) -> str:
        t = trend.get(key, {})
        return (
            f"- {titel}: {_fmt(t.get('start'))} → {_fmt(t.get('ende'))} "
            f"(Tendenz: {t.get('richtung', 'stabil')})"
        )

    zeilen: list[str] = []
    zeilen.append(
        "Du bist eine erfahrene AHS-Deutschlehrerin in Österreich und schreibst eine "
        "Lernverlaufs-Einschätzung für eine Kollegin, die diesen Schüler übernommen hat. "
        "Schreibe so, wie du es einer geschätzten Kollegin im Konferenzzimmer sagen würdest: "
        "fachlich genau, ehrlich, aber wohlwollend. Du-Anrede an die Kollegin, nie an den "
        "Schüler."
    )
    zeilen.append("")
    zeilen.append(
        "Datenbasis sind aggregierte Kennzahlen aus mehreren korrigierten Arbeiten desselben "
        "Schülers. Es liegen bewusst keine Namen und keine Texte vor, nur Zahlen und kurze "
        "Fehlerbeispiele."
    )
    zeilen.append("")
    zeilen.append(f"Anzahl ausgewerteter Arbeiten: {anzahl}")
    zeilen.append("")
    zeilen.append("VERLAUF (chronologisch, je Arbeit Stufen der vier Hauptkriterien):")
    for i, v in enumerate(verlauf, start=1):
        k = v.get("kriterien", {})
        zeilen.append(
            f"  Arbeit {i}: Inhalt {_fmt(k.get('inhalt'))}, "
            f"Textstruktur {_fmt(k.get('textstruktur'))}, "
            f"Ausdruck {_fmt(k.get('ausdruck'))}, "
            f"Sprachrichtigkeit {_fmt(k.get('sprachrichtigkeit'))} "
            f"| K1 {_fmt(v.get('k1'))}, K3 {_fmt(v.get('k3'))} "
            f"| Note(App) {_fmt(v.get('note_app'))}, Note(Lehrkraft) {_fmt(v.get('note_lehrer'))}"
        )
    zeilen.append("")
    zeilen.append("TREND (erste → letzte Arbeit; 'steigt' = Verbesserung):")
    zeilen.append(_trendzeile("note_app", "Note (App)"))
    zeilen.append(_trendzeile("note_lehrer", "Note (Lehrkraft)"))
    zeilen.append(_trendzeile("inhalt", "Inhalt"))
    zeilen.append(_trendzeile("textstruktur", "Textstruktur"))
    zeilen.append(_trendzeile("ausdruck", "Ausdruck"))
    zeilen.append(_trendzeile("sprachrichtigkeit", "Sprachrichtigkeit"))
    zeilen.append(_trendzeile("k1", "K1 (Inhalt + Textstruktur)"))
    zeilen.append(_trendzeile("k3", "K3 (Ausdruck + Sprachnormen)"))
    zeilen.append("")
    zeilen.append("FEHLERSCHWERPUNKTE (häufigste Fehlertypen, mit kurzen Beispielzitaten):")
    if fehlerschwerpunkte:
        for f in fehlerschwerpunkte:
            beispiele = "; ".join(
                # defensiv auf 6 Wörter kürzen — es dürfen nur kurze isolierte Zitate rein
                " ".join((b.get("zitat", "")).split()[:6])
                for b in f.get("beispiele", [])
            )
            zusatz = f" (z. B. {beispiele})" if beispiele else ""
            zeilen.append(f"  - {f.get('label', f.get('typ'))}: {f.get('anzahl', 0)}×{zusatz}")
    else:
        zeilen.append("  (keine Fehlerdaten vorhanden)")
    zeilen.append("")
    zeilen.append("SO LIEST DU DIE DATEN (wichtig, halte dich daran):")
    zeilen.append(
        "- Eine einzelne schwache Arbeit zwischen besseren ist meist ein Ausreißer, keine "
        "Tendenz. Benenne den Trend über mehrere Arbeiten, nicht einen einzelnen Einbruch."
    )
    zeilen.append(
        "- Eine Stagnation auf hoher Stufe (4 bis 5) ist gefestigtes Können, keine Schwäche. "
        "Eine Stagnation auf niedriger Stufe (1 bis 2) ist ein Förderauftrag. Unterscheide das."
    )
    zeilen.append(
        "- Die Spreizung zwischen K1 (Inhalt und Struktur) und K3 (Ausdruck und Sprachnormen) "
        "erzählt das eigentliche Profil: Ein Schüler mit starkem K1 und schwachem K3 hat etwas "
        "zu sagen, aber Mühe mit der Form. Umgekehrt deutet starkes K3 bei schwachem K1 auf "
        "sprachliche Sicherheit ohne inhaltliche Tiefe. Mache diese Unterscheidung explizit."
    )
    zeilen.append(
        "- Wenn App-Note und Lehrer-Note auseinanderliegen, deute es vorsichtig: Eine mildere "
        "Lehrer-Note kann auf sichtbare Anstrengung, mündliche Stärke oder Kontext hindeuten, "
        "den die reine Textanalyse nicht sieht. Behaupte nichts, biete es als Lesart an."
    )
    zeilen.append(
        "- Bei den Fehlerschwerpunkten zählt das Muster, nicht die Einzelzahl: Wenn ein Fehlertyp "
        "über mehrere Arbeiten dominiert, ist das der lohnendste Übungsansatz."
    )
    zeilen.append("")
    zeilen.append("DEINE EINSCHÄTZUNG:")
    zeilen.append(
        "- Beginne mit echten Stärken, konkret an der Entwicklung festgemacht, nicht pauschal."
    )
    zeilen.append(
        "- Benenne dann die Förderbereiche ehrlich, jeweils mit einem konkreten, in der Klasse "
        "umsetzbaren Übungsvorschlag, der auf die schwächste SRDP-Kategorie zielt."
    )
    zeilen.append(
        "- Bei Oberstufe: Verorte den Stand mit Blick auf die Matura. Sage konkret, welche "
        "Kategorie bis zur Reifeprüfung noch tragen muss und wo der Hebel liegt. Sei dabei "
        "realistisch, nicht beschönigend und nicht entmutigend."
    )
    zeilen.append("")
    zeilen.append(
        "Keine KI-Floskeln ('Es wäre ratsam', 'Man könnte', 'Es lässt sich feststellen'). "
        "Schreibe wie ein Mensch mit Unterrichtserfahrung, auf Deutsch mit korrekten Umlauten."
    )
    zeilen.append("")
    zeilen.append(
        "Antworte NUR mit validem JSON in dieser Form:\n"
        "{\n"
        '  "kurzbild": "2-3 Sätze Gesamtbild des Schülers, wie eine mündliche Einschätzung",\n'
        '  "staerken": ["konkrete Stärke mit Bezug zur Entwicklung", ...],\n'
        '  "foerderbereiche": [\n'
        '    {"kategorie": "Sprachrichtigkeit", "befund": "was die Daten zeigen",\n'
        '     "uebung": "konkreter, in der Klasse umsetzbarer Übungsvorschlag"}\n'
        "  ],\n"
        '  "maturabezug": "nur Oberstufe: realistische Verortung mit Blick auf die Matura, sonst leer"\n'
        "}"
    )
    return "\n".join(zeilen)


def build_klassen_briefing_prompt(aggregat: dict[str, Any]) -> str:
    """Baut einen DATENMINIMIERTEN Prompt für ein LLM-Klassen-Briefing.

    DSGVO-KRITISCH: Dieser Prompt darf NIEMALS enthalten:
    - Schülernamen, Klassenname (nur "Klasse X" als Platzhalter)
    - Dateinamen, zusammenhängende Schülertexte
    - Vollständige Beispielzitate (nur gekürzte Fragmente, max. 6 Wörter)

    Erlaubt sind NUR: aggregierte Stufen/Noten, Fehlertyp-Häufigkeiten,
    Prozentwerte, kurze isolierte Fehlerfragmente ohne Namen.
    """
    k1k3 = aggregat.get("k1k3", {})
    noten = aggregat.get("notenverteilung", {})
    kalibrierung = aggregat.get("kalibrierung", {})
    heatmap = aggregat.get("heatmap", [])
    trend = aggregat.get("trend", [])
    anzahl = aggregat.get("anzahl_abgaben", 0)
    fehler_gesamt = aggregat.get("gesamt_fehler", 0)
    kriterien = aggregat.get("kriterien", {})
    aufgabe = aggregat.get("aufgabe", "")

    def _fmt(x: object) -> str:
        return "—" if x is None else (f"{x:.2f}" if isinstance(x, (int, float)) else str(x))

    zeilen: list[str] = []
    zeilen.append(
        "Du bist eine erfahrene AHS-Deutschlehrerin und Fachkoordinatorin in Österreich. "
        "Du schreibst ein kurzes, aber fundiertes Klassen-Briefing für eine Kollegin, die "
        "gerade eine Gesamtkorrektur abgeschlossen hat. Schreibe so, wie du es einer "
        "geschätzten Kollegin im Konferenzzimmer sagen würdest: fachlich genau, ehrlich, "
        "aber lösungsorientiert. Du-Anrede an die Kollegin."
    )
    zeilen.append("")
    zeilen.append(
        "Datenbasis sind anonymisierte Kennzahlen einer komplett korrigierten Klassenarbeit. "
        "Es liegen bewusst keine Namen und keine Texte vor, nur Aggregate und kurze Fehlerfragmente."
    )
    zeilen.append("")
    zeilen.append(f"Aufgabe: {aufgabe or 'alle Aufgaben'}")
    zeilen.append(f"Anzahl ausgewerteter Abgaben: {anzahl}")
    zeilen.append(f"Gesamtanzahl erfasster Fehler: {fehler_gesamt}")
    zeilen.append("")

    zeilen.append("K1 / INHALT & STRUKTUR (Ø über alle Abgaben):")
    k1 = k1k3.get("k1", {})
    zeilen.append(f"  Ø Stufe: {_fmt(k1.get('avg'))} (n={k1.get('count', 0)})")
    zeilen.append("")

    zeilen.append("K3 / SPRACHE & AUSDRUCK (Ø über alle Abgaben):")
    k3 = k1k3.get("k3", {})
    zeilen.append(f"  Ø Stufe: {_fmt(k3.get('avg'))} (n={k3.get('count', 0)})")
    zeilen.append("")

    zeilen.append("NOTENVERTEILUNG (1=sehr gut, 5=nicht genügend):")
    for note in range(1, 6):
        zeilen.append(f"  Note {note}: {noten.get(note, 0)} Abgaben")
    zeilen.append("")

    n_mit = kalibrierung.get("n_mit_feedback", 0)
    n_ges = kalibrierung.get("n_gesamt", 0)
    if n_mit > 0:
        zeilen.append("KALIBRIERUNG (App-Note vs. Lehrer-Note):")
        zeilen.append(f"  App-Ø: {_fmt(kalibrierung.get('app_avg'))}")
        zeilen.append(f"  Lehrer-Ø: {_fmt(kalibrierung.get('lehrer_avg'))}")
        zeilen.append(f"  Differenz: {_fmt(kalibrierung.get('diff'))}")
        zeilen.append(f"  Tendenz: {kalibrierung.get('tendenz', 'n/a')}")
        zeilen.append(f"  Vergleichspaare: {n_mit} von {n_ges} Abgaben")
        zeilen.append("")

    if trend:
        zeilen.append("TREND ÜBER AUFGABEN (chronologisch):")
        for t in trend:
            lehrer_str = f", Lehrer-Ø {_fmt(t.get('avg_note_lehrer'))}" if t.get("avg_note_lehrer") is not None else ""
            zeilen.append(
                f"  {t.get('aufgabe', '?')}: App-Ø {_fmt(t.get('avg_note_app'))}{lehrer_str} "
                f"(n={t.get('n', 0)})"
            )
        zeilen.append("")

    zeilen.append("FEHLER-HEATMAP (dominante Typen):")
    for h in heatmap:
        zeilen.append(f"  {h.get('typ', '?')}: {h.get('anzahl', 0)} ({h.get('prozent', 0):.1f}%)")
    zeilen.append("")

    # Beispielzitate werden bewusst NICHT in den Prompt aufgenommen,
    # da sie potenziell Schülernamen enthalten könnten (DSGVO).
    # Die Heatmap (Fehlertypen + Prozente) reicht als Datenbasis.

    if kriterien:
        zeilen.append("KRITERIEN-ROHDURCHSCHNITTE (zur Orientierung):")
        for k_name, k_data in sorted(kriterien.items()):
            avg = k_data.get("avg")
            if avg is not None:
                zeilen.append(f"  {k_name}: Stufe {avg:.2f} (n={k_data.get('count', 0)})")
        zeilen.append("")

    zeilen.append("SO LIEST DU DIE DATEN:")
    zeilen.append(
        "- Die K1/K3-Spaltung ist der entscheidende Hebel, ABER: Sprich von einer Stärke-Schwäche-"
        "Spreizung nur, wenn der Abstand zwischen K1 und K3 mindestens 0.5 Stufen beträgt. "
        "Liegen K1 und K3 näher beieinander (z. B. 2.8 vs. 2.7), beschreibe das Profil als "
        "ausgewogen, nicht als gespalten."
    )
    zeilen.append(
        "- Die Notenverteilung zeigt das Leistungsprofil der Gruppe: Eine breite Streuung "
        "bedeutet differenzierten Unterricht, eine Konzentration auf 3–4 deutet auf einen "
        "gemeinsamen Wiederholungsbedarf hin."
    )
    kali_warn = (
        "- Wenn App-Note und Lehrer-Note auseinanderliegen, deute es vorsichtig: "
    )
    if n_mit < 5:
        kali_warn += (
            f"Es liegen nur {n_mit} Vergleichspaare vor (von {n_ges} Abgaben). "
            "Das ist eine Momentaufnahme, keine systematische Aussage. Sprich von 'ersten Hinweisen', "
            "nicht von 'die Klasse wird systematisch zu milde/streng bewertet'. "
        )
    else:
        kali_warn += (
            f"Es liegen {n_mit} Vergleichspaare vor — das ist eine solide Basis. "
        )
    kali_warn += (
        "Eine Abweichung kann auf mündliche Stärken oder Kontext hindeuten, den die "
        "Textanalyse nicht sieht. Nicht als Fehler der App lesen, sondern als Hinweis auf "
        "nicht-textliche Leistungsanteile."
    )
    zeilen.append(kali_warn)
    zeilen.append(
        "- Bei den Fehlern zählt das Muster, nicht die Einzelzahl: Wenn ein Typ über mehrere "
        "Arbeiten dominiert, ist das der lohnendste Übungsansatz für die ganze Klasse."
    )
    zeilen.append("")
    zeilen.append("DEINE EINSCHÄTZUNG:")
    zeilen.append(
        "- Beginne mit einer knappen, aber ehrlichen Gesamteinschätzung der Klasse."
    )
    zeilen.append(
        "- Benenne dann 2–3 konkrete Schwerpunkte mit Bezug zu den Daten, jeweils mit einer "
        "konkreten, im Unterricht umsetzbaren Unterrichtsempfehlung. Die Empfehlungen müssen "
        "Stundenideen enthalten (was genau in der nächsten Stunde passiert), nicht nur allgemeine "
        "Ratschläge. Koppelle jede Empfehlung an die schwächste SRDP-Kategorie."
    )
    zeilen.append(
        "- Bei Oberstufe: Verorte den Stand mit Blick auf die Matura. Sage konkret, welche "
        "Kategorie bis zur Reifeprüfung noch tragen muss und wo der Hebel liegt. Sei realistisch, "
        "nicht beschönigend und nicht entmutigend."
    )
    zeilen.append("")
    zeilen.append(
        "Keine KI-Floskeln ('Es wäre ratsam', 'Man könnte', 'Es lässt sich feststellen', "
        "'Es ist wichtig, dass'). Schreibe wie eine erfahrene Kollegin mit 15 Jahren "
        "Unterrichtspraxis, auf Deutsch mit korrekten Umlauten."
    )
    zeilen.append("")
    zeilen.append(
        "Antworte NUR mit validem JSON in dieser Form:\n"
        "{\n"
        '  "kurzbild": "3–4 Sätze Gesamteinschätzung der Klasse, wie eine mündliche Einschätzung",\n'
        '  "schwerpunkte": [\n'
        '    {"bereich": "K1 / Inhalt & Struktur", "befund": "was die Daten zeigen",\n'
        '     "empfehlung": "konkrete Unterrichtsempfehlung mit Stundenidee"},\n'
        '    {"bereich": "K3 / Sprache & Ausdruck", "befund": "...", "empfehlung": "..."}\n'
        "  ],\n"
        '  "unterrichtsempfehlungen": [\n'
        '    {"fokus": "konkrete Schwäche", "stundenidee": "Was genau in der nächsten Stunde passiert",\n'
        '     "material": "Arbeitsblatt / Übung / Text", "zielgruppe": "schwache Hälfte / ganze Klasse"}\n'
        "  ],\n"
        '  "matura_fokus": "nur Oberstufe: realistische Verortung mit Blick auf die Matura, sonst leer"\n'
        "}"
    )
    return "\n".join(zeilen)


def docx_to_pdf(docx_path: Path, out_path: Path | None = None) -> Path | None:
    """Konvertiert DOCX nach PDF via LibreOffice headless oder unoconv.

    Args:
        docx_path: Pfad zur Quelldatei (.docx)
        out_path: Gewuenschter Ausgabepfad (optional; Standard: gleicher Ordner, .pdf-Endung)

    Returns:
        Pfad zur erzeugten PDF-Datei, oder None wenn kein Konverter verfuegbar.
    """
    candidates = [
        (shutil.which("libreoffice"), ["libreoffice", "--headless", "--convert-to", "pdf"]),
        (shutil.which("soffice"), ["soffice", "--headless", "--convert-to", "pdf"]),
        (shutil.which("unoconv"), ["unoconv", "-f", "pdf"]),
    ]

    cmd_base = None
    for found, base in candidates:
        if found:
            cmd_base = base
            break

    if cmd_base is None:
        return None

    if out_path is None:
        out_path = docx_path.with_suffix(".pdf")

    out_dir = str(out_path.parent)

    if cmd_base[0] in ("libreoffice", "soffice"):
        cmd = cmd_base + [f"--outdir={out_dir}", str(docx_path)]
    else:
        cmd = cmd_base + ["-o", str(out_path), str(docx_path)]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode == 0 and out_path.exists():
            return out_path
    except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
        pass

    return None
